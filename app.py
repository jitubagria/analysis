import streamlit as st
import pandas as pd
import numpy as np
import io
import base64
import plotly.express as px
import plotly.graph_objects as go
import random
import time
from data_utils import load_data, get_descriptive_stats, generate_age_groups, analyze_master_parameters, calculate_paired_differences, calculate_manual_paired_differences
from stats_utils import perform_ttest, perform_anova, perform_correlation_analysis, perform_pvalue_analysis
from visualization_utils import (create_histogram, create_boxplot, create_scatterplot, 
                                create_correlation_heatmap, create_distribution_plot)
from report_generator import generate_report
from ml_utils import (analyze_data_structure, prepare_data_for_ml, train_model, 
                     perform_automated_analysis)
from drag_drop_ui import display_drag_drop_ui  # Using the new drag and drop UI
from drag_drop_config import display_config_tool  # Keep the old one as backup

# Import AI utilities
from gemini_utils import set_gemini_api_key, get_gemini_summary, check_gemini_api_status
from ai_data_preparation_clean import get_available_analysis_results, format_data_for_ai

# Define emoji dictionary for different analysis types
EMOJI_DICT = {
    "System": "🖥️",
    "Gender": "👤",
    "Age": "👶👧👩👵",
    "Pre/Post": "⏮️⏭️",
    "Statistics": "📊",
    "Correlation": "🔄",
    "Upload": "📤",
    "Configuration": "⚙️",
    "Analysis": "🧮",
    "Tables": "📋",
    "Charts": "📈",
    "Output": "📁",
    "Research": "🔬",
    "Success": ["🎉", "🥳", "🚀", "💯", "🏆"],
    "Loading": ["⏳", "🔄", "⌛", "🔍", "🧠"]
}

# Confetti JS code
CONFETTI_JS = """
<script src="https://cdn.jsdelivr.net/npm/canvas-confetti@1.5.1/dist/confetti.browser.min.js"></script>
<script>
function launchConfetti() {
    const count = 200;
    const defaults = {
        origin: { y: 0.7 },
        zIndex: 1000
    };
    
    function fire(particleRatio, opts) {
        confetti({
            ...defaults,
            ...opts,
            particleCount: Math.floor(count * particleRatio)
        });
    }
    
    fire(0.25, {
        spread: 26,
        startVelocity: 55,
    });
    fire(0.2, {
        spread: 60,
    });
    fire(0.35, {
        spread: 100,
        decay: 0.91,
        scalar: 0.8
    });
    fire(0.1, {
        spread: 120,
        startVelocity: 25,
        decay: 0.92,
        scalar: 1.2
    });
    fire(0.1, {
        spread: 120,
        startVelocity: 45,
    });
}
launchConfetti();
</script>
"""

# Initialize session state for chart favorites
if 'chart_favorites' not in st.session_state:
    st.session_state.chart_favorites = {}

# Function to toggle a chart in the favorites (chart cart)
def toggle_chart_favorite(chart_id, fig, title, chart_type):
    """Toggle a chart in the favorites (add/remove)"""
    # Initialize dictionary to track favorites if it doesn't exist
    if 'chart_favorites' not in st.session_state:
        st.session_state.chart_favorites = {}
    
    # If the chart is already in favorites, remove it
    if chart_id in st.session_state.chart_favorites:
        del st.session_state.chart_favorites[chart_id]
        st.toast(f"Removed '{title}' from favorites", icon="💔")
        return False
    # Otherwise add it to favorites
    else:
        st.session_state.chart_favorites[chart_id] = {
            'figure': fig,
            'title': title,
            'chart_type': chart_type,
            'added_time': time.time()
        }
        st.toast(f"Added '{title}' to favorites!", icon="❤️")
        return True

# Function to remove a chart from favorites by chart_id
def remove_from_favorites(chart_id):
    """Remove a chart from favorites"""
    if 'chart_favorites' in st.session_state and chart_id in st.session_state.chart_favorites:
        removed_item = st.session_state.chart_favorites.pop(chart_id)
        st.toast(f"Removed '{removed_item['title']}' from favorites", icon="💔")

# Function to export favorite charts to Microsoft Word
def create_pvalue_chart(pvalue_data, title):
    """Create a chart for p-value analysis results"""
    import plotly.express as px
    import plotly.graph_objects as go
    
    if len(pvalue_data) == 0:
        # Create empty chart
        fig = go.Figure()
        fig.add_annotation(text="No data available", xref="paper", yref="paper",
                          x=0.5, y=0.5, showarrow=False, font=dict(size=16))
        fig.update_layout(title=title, showlegend=False)
        return fig
    
    # Create bar chart for p-values
    fig = px.bar(
        pvalue_data, 
        x='Parameter', 
        y='P_Value',
        color='Significant',
        color_discrete_map={'Yes': '#E74C3C', 'No': '#27AE60'},
        title=title,
        labels={'P_Value': 'P-Value', 'Parameter': 'Parameters'}
    )
    
    # Add significance line at alpha=0.05
    fig.add_hline(y=0.05, line_dash="dash", line_color="red", 
                  annotation_text="α = 0.05", annotation_position="bottom right")
    
    # Update layout for professional appearance
    fig.update_layout(
        font_family="Times New Roman",
        plot_bgcolor='white',
        paper_bgcolor='white',
        title_font_size=16,
        title_x=0.5,
        xaxis=dict(title_font_size=12, tickfont_size=10),
        yaxis=dict(title_font_size=12, tickfont_size=10),
        legend=dict(font_size=10)
    )
    
    return fig

def export_charts_to_word():
    """Export all favorited charts to a Microsoft Word document"""
    from docx import Document
    from docx.shared import Inches
    import io
    import base64
    from PIL import Image
    from datetime import datetime
    
    # Create a new Word document
    doc = Document()
    doc.add_heading('Statistical Analysis Charts', 0)
    
    # Add creation date
    now = datetime.now()
    doc.add_paragraph(f"Generated on: {now.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()
    
    # Check if we have favorites
    if 'chart_favorites' not in st.session_state or not st.session_state.chart_favorites:
        doc.add_paragraph("No charts have been favorited. Please mark charts with the heart icon to include them in the export.")
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    
    # Add each favorited chart to the document
    for chart_id, chart_item in st.session_state.chart_favorites.items():
        # Add chart title
        doc.add_heading(chart_item['title'], level=1)
        
        try:
            # Create a clean copy of the figure without UI elements for export
            # Use the original figure directly instead of trying to copy it
            export_fig = chart_item['figure']
            
            # Convert the figure to an image
            img_bytes = io.BytesIO()
            export_fig.write_image(img_bytes, format="png", scale=2)  # Higher scale for better quality
            img_bytes.seek(0)
            
            # Add image to document
            doc.add_picture(img_bytes, width=Inches(6))
            
            # Add chart type as metadata
            chart_type = chart_item.get('chart_type', 'Chart')
            doc.add_paragraph(f"Chart type: {chart_type}")
            
        except Exception as e:
            # If direct image conversion fails, add a placeholder and note
            doc.add_paragraph(f"Chart could not be rendered: {str(e)}")
        
        # Add some spacing
        doc.add_paragraph()
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = f"Created with Statistical Analysis Portal • {now.strftime('%Y-%m-%d')}"
    
    # Save the document to a BytesIO object
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Function to display confetti
def show_confetti():
    st.markdown(CONFETTI_JS, unsafe_allow_html=True)

# Function to display spinning emoji
def show_spinner_emoji(emoji_type="Loading"):
    emojis = EMOJI_DICT.get(emoji_type, ["⏳"])
    emoji = random.choice(emojis) if isinstance(emojis, list) else emojis
    st.markdown(
        f"""
        <style>
        .emoji-spinner {{
            display: inline-block;
            animation: spin 2s linear infinite;
            font-size: 2em;
        }}
        @keyframes spin {{
            0% {{ transform: rotate(0deg); }}
            100% {{ transform: rotate(360deg); }}
        }}
        </style>
        <div class="emoji-spinner">{emoji}</div>
        """,
        unsafe_allow_html=True
    )
    return emoji

# Function to add decorative emoji to section headers
def emoji_header(title, emoji_type):
    emoji = random.choice(EMOJI_DICT.get(emoji_type, ["📊"])) if isinstance(EMOJI_DICT.get(emoji_type, ["📊"]), list) else EMOJI_DICT.get(emoji_type, "📊")
    return f"{emoji} {title} {emoji}"

# Function to create a playful progress bar
def fun_progress_bar(percent, label="Loading"):
    # Random emojis to use for the progress bar
    progress_emojis = ["🔵", "🟢", "🟡", "🟣", "🔴", "⚪"]
    empty_emoji = "⚫"
    
    # Number of segments in the progress bar
    segments = 20
    filled = int(segments * percent / 100)
    
    # Create the progress bar
    bar = ""
    for i in range(segments):
        if i < filled:
            bar += random.choice(progress_emojis)
        else:
            bar += empty_emoji
    
    # Display the progress bar
    st.markdown(f"**{label}**: {bar} {percent}%", unsafe_allow_html=True)
    
# Function to display mobile-friendly dataframes
def display_mobile_friendly_dataframe(df, title="Dataset", use_container_width=True, height=None):
    """Display a dataframe in a mobile-friendly way with column selection options"""
    st.subheader(title)
    
    # Add mobile-friendly table controls
    if st.checkbox(f"Enable mobile-friendly view for {title}", value=False, key=f"mobile_view_{title}"):
        # For mobile, limit the number of columns shown and add a column selector
        col1, col2 = st.columns([1, 1])
        with col1:
            num_cols = st.slider(f"Number of columns ({title})", 
                                min_value=2, 
                                max_value=min(10, len(df.columns)), 
                                value=min(5, len(df.columns)),
                                key=f"num_cols_{title}")
        
        with col2:
            # Allow users to select which columns to view on mobile
            selected_cols = st.multiselect(
                f"Select columns ({title})",
                options=list(df.columns),
                default=list(df.columns)[:num_cols],
                key=f"selected_cols_{title}"
            )
            
            if not selected_cols:  # If no columns selected, use default
                selected_cols = list(df.columns)[:num_cols]
        
        # Show only selected columns
        display_df = df[selected_cols]
        st.dataframe(display_df, use_container_width=use_container_width, height=height)
    else:
        # Show all columns (desktop view)
        st.dataframe(df, use_container_width=use_container_width, height=height)

# Set page configuration with mobile-responsive settings
st.set_page_config(
    page_title="Statistical Analysis Portal",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items={
        'About': "Statistical Analysis Portal - A comprehensive data analysis tool"
    }
)

# Add meta tags for mobile responsiveness
st.markdown("""
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="mobile-web-app-capable" content="yes">
""", unsafe_allow_html=True)

# Add comprehensive business-class CSS styling
st.markdown("""
<style>
/* Business Class Theme - Professional Styling */
:root {
    --primary-color: #2E86AB;
    --secondary-color: #A23B72;
    --accent-color: #F18F01;
    --success-color: #27AE60;
    --warning-color: #F39C12;
    --error-color: #E74C3C;
    --text-primary: #2C3E50;
    --text-secondary: #7F8C8D;
    --background-primary: #FAFAFA;
    --background-secondary: #F5F7FA;
    --background-card: #FFFFFF;
    --border-color: #E8ECEF;
    --shadow-light: 0 2px 4px rgba(0,0,0,0.05);
    --shadow-medium: 0 4px 12px rgba(0,0,0,0.1);
    --shadow-strong: 0 8px 24px rgba(0,0,0,0.15);
    --border-radius: 8px;
    --border-radius-large: 12px;
}

/* Global improvements */
.main .block-container {
    background: var(--background-primary);
    min-height: 100vh;
    padding: 2rem 1rem;
}

.stApp {
    background: var(--background-primary);
}

/* Enhanced typography */
h1, h2, h3, h4, h5, h6 {
    color: var(--text-primary) !important;
    font-weight: 600 !important;
    letter-spacing: -0.025em !important;
}

h1 {
    font-size: 2.5rem !important;
    color: var(--text-primary) !important;
    margin-bottom: 2rem !important;
}

/* Professional card styling */
.stExpander, .stContainer, .element-container {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
    padding: 1.5rem !important;
    margin-bottom: 1rem !important;
}

/* Enhanced buttons */
.stButton > button {
    background: var(--primary-color) !important;
    color: white !important;
    border: none !important;
    border-radius: var(--border-radius) !important;
    padding: 0.75rem 1.5rem !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.025em !important;
    box-shadow: var(--shadow-light) !important;
    cursor: pointer !important;
}

.stButton > button:hover {
    background: var(--secondary-color) !important;
}

/* Heart button styling */
.heart-button {
    background: transparent !important;
    border: 2px solid var(--border-color) !important;
    color: var(--text-secondary) !important;
    font-size: 1.5rem !important;
    cursor: pointer !important;
    transition: all 0.3s ease !important;
    padding: 0.5rem !important;
    border-radius: 50% !important;
    width: 44px !important;
    height: 44px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}

.heart-button:hover {
    color: var(--error-color) !important;
    border-color: var(--error-color) !important;
    transform: scale(1.1) !important;
    background: rgba(231, 76, 60, 0.1) !important;
}

.heart-button.active {
    color: var(--error-color) !important;
    border-color: var(--error-color) !important;
    background: rgba(231, 76, 60, 0.1) !important;
}

/* Enhanced forms and inputs */
.stSelectbox, .stMultiselect, .stTextInput, .stTextArea, .stNumberInput {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
}

.stSelectbox > div > div {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    min-height: 44px !important;
}

/* Enhanced dataframes */
.stDataFrame {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
    overflow: hidden !important;
}

.stDataFrame table {
    background: var(--background-card) !important;
}

.stDataFrame th {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color)) !important;
    color: white !important;
    font-weight: 600 !important;
    padding: 1rem !important;
    border: none !important;
}

.stDataFrame td {
    padding: 0.75rem 1rem !important;
    border-bottom: 1px solid var(--border-color) !important;
}

/* Enhanced charts */
.stPlotlyChart {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius-large) !important;
    box-shadow: var(--shadow-medium) !important;
    padding: 1rem !important;
    margin: 1rem 0 !important;
}

/* Enhanced sidebar */
.css-1d391kg {
    background: var(--background-card) !important;
    border-right: 1px solid var(--border-color) !important;
    box-shadow: var(--shadow-medium) !important;
}

/* Navigation improvements */
.desktop-nav-container {
    background: var(--background-card) !important;
    padding: 1rem !important;
    border-radius: var(--border-radius-large) !important;
    box-shadow: var(--shadow-medium) !important;
    margin-bottom: 2rem !important;
    border: 1px solid var(--border-color) !important;
}

.mobile-nav-container {
    background: var(--background-card) !important;
    padding: 1rem !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
    margin-bottom: 1rem !important;
    border: 1px solid var(--border-color) !important;
}

/* Enhanced metrics */
.metric-container {
    background: var(--background-card) !important;
    padding: 1.5rem !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
    border: 1px solid var(--border-color) !important;
    text-align: center !important;
}

/* Progress indicators */
.stProgress {
    background: var(--border-color) !important;
    border-radius: 1rem !important;
    overflow: hidden !important;
}

.stProgress > div > div {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color)) !important;
}

/* Enhanced alerts and messages */
.stSuccess {
    background: rgba(39, 174, 96, 0.1) !important;
    border: 1px solid var(--success-color) !important;
    border-radius: var(--border-radius) !important;
    color: var(--success-color) !important;
}

.stError {
    background: rgba(231, 76, 60, 0.1) !important;
    border: 1px solid var(--error-color) !important;
    border-radius: var(--border-radius) !important;
    color: var(--error-color) !important;
}

.stWarning {
    background: rgba(243, 156, 18, 0.1) !important;
    border: 1px solid var(--warning-color) !important;
    border-radius: var(--border-radius) !important;
    color: var(--warning-color) !important;
}

.stInfo {
    background: rgba(46, 134, 171, 0.1) !important;
    border: 1px solid var(--primary-color) !important;
    border-radius: var(--border-radius) !important;
    color: var(--primary-color) !important;
}

/* Mobile responsiveness */
@media (max-width: 768px) {
    .main .block-container {
        padding: 1rem 0.5rem !important;
    }
    
    h1 {
        font-size: 2rem !important;
    }
    
    .stButton > button {
        width: 100% !important;
        font-size: 16px !important;
        padding: 12px 16px !important;
    }
    
    .stSelectbox > div > div,
    input, select {
        min-height: 44px !important;
        font-size: 16px !important;
    }
    
    .heart-button {
        width: 48px !important;
        height: 48px !important;
        font-size: 1.25rem !important;
    }
    
    .stDataFrame, .stExpander {
        overflow-x: auto !important;
    }
    
    .desktop-nav-container {
        display: none !important;
    }
}

/* Tablet responsiveness */
@media (min-width: 769px) and (max-width: 1024px) {
    .main .block-container {
        padding: 1.5rem !important;
    }
    
    h1 {
        font-size: 2.25rem !important;
    }
}

/* Desktop enhancements */
@media (min-width: 1025px) {
    .main .block-container {
        max-width: 1200px !important;
        margin: 0 auto !important;
    }
    
    .mobile-nav-container {
        display: none !important;
    }
}

/* Remove all animations */

/* Enhanced scrollbars */
::-webkit-scrollbar {
    width: 8px;
    height: 8px;
}

::-webkit-scrollbar-track {
    background: var(--background-secondary);
    border-radius: 4px;
}

::-webkit-scrollbar-thumb {
    background: var(--text-secondary);
    border-radius: 4px;
}

::-webkit-scrollbar-thumb:hover {
    background: var(--primary-color);
}

/* Business section enhancements */
.business-section-header {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white !important;
    padding: 2rem !important;
    border-radius: var(--border-radius-large) !important;
    margin-bottom: 2rem !important;
    text-align: center !important;
    box-shadow: var(--shadow-medium) !important;
}

.business-section-header h2 {
    color: white !important;
    margin-bottom: 0.5rem !important;
    font-size: 2rem !important;
}

.business-section-header p {
    color: rgba(255, 255, 255, 0.9) !important;
    font-size: 1.1rem !important;
    margin-bottom: 0 !important;
}

.metrics-card {
    background: var(--background-card) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: var(--border-radius) !important;
    box-shadow: var(--shadow-light) !important;
    padding: 1.5rem !important;
    margin-bottom: 2rem !important;
    text-align: center !important;
}

.metric-item {
    display: inline-block !important;
}

.metric-number {
    font-size: 3rem !important;
    font-weight: 700 !important;
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1 !important;
}

.metric-label {
    font-size: 1.1rem !important;
    color: var(--text-secondary) !important;
    font-weight: 500 !important;
    margin-top: 0.5rem !important;
}

.download-container {
    text-align: center !important;
    padding: 1rem !important;
    background: rgba(46, 134, 171, 0.05) !important;
    border-radius: var(--border-radius) !important;
    border: 1px solid var(--primary-color) !important;
    margin: 1rem 0 !important;
}

.download-button {
    display: inline-block !important;
    background: linear-gradient(135deg, var(--success-color), #2ECC71) !important;
    color: white !important;
    padding: 1rem 2rem !important;
    border-radius: var(--border-radius) !important;
    text-decoration: none !important;
    font-weight: 600 !important;
    font-size: 1.1rem !important;
    box-shadow: var(--shadow-medium) !important;
    transition: all 0.3s ease !important;
    text-transform: uppercase !important;
    letter-spacing: 0.025em !important;
}

.download-button:hover {
    transform: translateY(-2px) !important;
    box-shadow: var(--shadow-strong) !important;
    background: linear-gradient(135deg, #27AE60, #2ECC71) !important;
    text-decoration: none !important;
    color: white !important;
}

.chart-preview-item {
    background: var(--background-secondary) !important;
    padding: 1rem !important;
    border-radius: var(--border-radius) !important;
    margin-bottom: 1rem !important;
    border-left: 4px solid var(--primary-color) !important;
}

.chart-preview-item h4 {
    margin: 0 0 0.5rem 0 !important;
    color: var(--text-primary) !important;
}

.chart-type-badge {
    background: linear-gradient(135deg, var(--accent-color), #E67E22) !important;
    color: white !important;
    padding: 0.25rem 0.75rem !important;
    border-radius: 1rem !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.025em !important;
}

.empty-state {
    text-align: center !important;
    padding: 4rem 2rem !important;
    background: var(--background-card) !important;
    border: 2px dashed var(--border-color) !important;
    border-radius: var(--border-radius-large) !important;
    margin: 2rem 0 !important;
}

.empty-state-icon {
    font-size: 4rem !important;
    margin-bottom: 1rem !important;
    opacity: 0.6 !important;
}

.empty-state h3 {
    color: var(--text-primary) !important;
    margin-bottom: 1rem !important;
    font-size: 1.5rem !important;
}

.empty-state p {
    color: var(--text-secondary) !important;
    font-size: 1.1rem !important;
    line-height: 1.6 !important;
    max-width: 500px !important;
    margin: 0 auto !important;
}

/* Print styles */
@media print {
    .main .block-container {
        background: white !important;
        box-shadow: none !important;
    }
    
    .stButton, .heart-button {
        display: none !important;
    }
    
    .business-section-header {
        background: var(--text-primary) !important;
        color: white !important;
    }
}
</style>
""", unsafe_allow_html=True)

# Initialize session state variables if they don't exist
# Initialize core session state variables
if 'data' not in st.session_state:
    st.session_state.data = None
if 'filename' not in st.session_state:
    st.session_state.filename = None
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'previous_step' not in st.session_state:
    st.session_state.previous_step = None
if 'selected_columns' not in st.session_state:
    st.session_state.selected_columns = []
if 'analysis_type' not in st.session_state:
    st.session_state.analysis_type = None
if 'visualization_type' not in st.session_state:
    st.session_state.visualization_type = None

# Initialize dropdown selections that need to persist
if 'auto_analyses' not in st.session_state:
    st.session_state.auto_analyses = []
if 'group_col' not in st.session_state:
    st.session_state.group_col = None
if 'target_col' not in st.session_state:
    st.session_state.target_col = None
if 'predictor_cols' not in st.session_state:
    st.session_state.predictor_cols = []
if 'pre_col' not in st.session_state:
    st.session_state.pre_col = None
if 'post_col' not in st.session_state:
    st.session_state.post_col = None

def main():
    # Initialize AI-related session state variables
    if 'gemini_api_key' not in st.session_state:
        st.session_state.gemini_api_key = ""
    if 'gemini_api_key_configured' not in st.session_state:
        st.session_state.gemini_api_key_configured = False
    if 'ai_generated_summary' not in st.session_state:
        st.session_state.ai_generated_summary = None
    
    # Fun title with random success emoji
    success_emoji = random.choice(EMOJI_DICT["Success"])
    st.title(f"Statistical Analysis Portal {success_emoji} 📊 {success_emoji}")
    
    # Sidebar for AI Configuration
    with st.sidebar:
        st.header("🤖 AI Analysis Configuration")
        st.markdown("Configure AI providers for comprehensive analysis summaries")
        
        # AI Provider Selection
        ai_provider = st.selectbox(
            "AI Provider",
            ["Google Gemini", "OpenAI (Coming Soon)", "Anthropic (Coming Soon)"],
            index=0,
            help="Select your preferred AI provider for analysis summaries"
        )
        
        if ai_provider == "Google Gemini":
            # Gemini Model Selection
            gemini_model = st.selectbox(
                "Gemini Model",
                ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-1.0-pro"],
                index=0,
                help="Flash: Faster and cost-effective | Pro: Enhanced reasoning capabilities"
            )
            
            # Store selected model in session state
            st.session_state.selected_gemini_model = gemini_model
            
            # Model Information
            if gemini_model == "gemini-1.5-flash":
                st.info("Flash model: Fast responses, optimized for speed and efficiency")
            elif gemini_model == "gemini-1.5-pro":
                st.info("Pro model: Advanced reasoning, ideal for complex analyses")
            else:
                st.info("Standard model: Balanced performance for general use")
            
            # API Key input
            api_key = st.text_input(
                "Gemini API Key:",
                value=st.session_state.get('gemini_api_key', ''),
                type="password",
                help="Get your free API key from Google AI Studio"
            )
            
            # Quick link to get API key
            st.markdown("[Get API Key](https://aistudio.google.com/app/apikey) | [Documentation](https://ai.google.dev/docs)")
            
            if api_key and api_key != st.session_state.get('gemini_api_key', ''):
                # Configure the API key
                if set_gemini_api_key(api_key):
                    st.success("API key configured successfully!")
                else:
                    st.error("Failed to configure API key. Please check your key and try again.")
            
            # Display current status
            status = check_gemini_api_status()
            if status['ready']:
                st.success(f"AI Ready - {gemini_model}")
            elif status['api_key_present']:
                st.warning("API Key needs validation")
            else:
                st.info("Enter API key to enable AI summaries")
        
        else:
            st.info("Additional AI providers will be available soon!")
    
    # Create a horizontal navigation bar with buttons to jump to steps
    steps = [
        {"title": "1. Upload Excel", "step": 1}, 
        {"title": "2. Data Management", "step": 2}, 
        {"title": "3. Data Analysis", "step": 3}, 
        {"title": "4. Data Tables", "step": 4}, 
        {"title": "5. Visualizations", "step": 5}, 
        {"title": "6. Reference Data", "step": 6},
        {"title": "7. Output", "step": 7},
        {"title": "8. AI-Powered Analysis", "step": 8}
    ]
    
    # Mobile-friendly navigation
    # For small screens (mobile), display a dropdown instead of buttons
    st.markdown("""
    <style>
    /* Hide the dropdown on larger screens */
    @media (min-width: 768px) {
        .mobile-nav-container {
            display: none;
        }
    }
    /* Hide the button navigation on small screens */
    @media (max-width: 767px) {
        .desktop-nav-container {
            display: none;
        }
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Mobile navigation with dropdown
    st.markdown('<div class="mobile-nav-container">', unsafe_allow_html=True)
    selected_step = st.selectbox(
        "Go to step:",
        options=[f"{step['step']}. {step['title'].split('. ')[1]}" for step in steps],
        index=st.session_state.current_step - 1
    )
    
    # Handle the mobile navigation selection
    if selected_step:
        selected_step_num = int(selected_step.split(".")[0])
        if selected_step_num != st.session_state.current_step:
            if selected_step_num == 1 or st.session_state.data is not None:
                st.session_state.previous_step = st.session_state.current_step
                st.session_state.current_step = selected_step_num
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Desktop navigation with buttons
    st.markdown('<div class="desktop-nav-container">', unsafe_allow_html=True)
    cols = st.columns(7)
    for i, (col, step) in enumerate(zip(cols, steps)):
        step_num = i + 1
        with col:
            if step_num == st.session_state.current_step:
                # Highlight current step with blue background
                st.markdown(f'<div style="background-color: #e6f2ff; padding: 10px; border-radius: 5px; text-align: center; font-weight: bold;">{step["title"]}</div>', unsafe_allow_html=True)
            else:
                # Create a button for each step
                if st.button(step["title"], key=f"nav_{step['step']}"):
                    # Only allow jumping to steps if data is loaded (except for the first step)
                    if step["step"] == 1 or st.session_state.data is not None:
                        # Store the current step before changing it
                        st.session_state.previous_step = st.session_state.current_step
                        
                        # Update the current step
                        st.session_state.current_step = step["step"]
                        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Add a little space after the navigation
    st.write("")
    
    # Original welcome text
    st.write("""
    Welcome to the Statistical Analysis Portal! Follow the steps below to analyze your data:
    1. Upload your data file
    2. Manage and clean your data
    3. Select your desired analysis
    4. Review data tables
    5. Create visualizations
    6. Generate final output
    """)
    
    # Sidebar for navigation and steps
    with st.sidebar:
        st.header("Navigation")
        
        # Show current step indicator
        st.markdown(f"**Current Step: {st.session_state.current_step}/7**")
        
        # Step descriptions
        steps = {
            1: "Upload Excel",
            2: "Data Management",
            3: "Desired Data Analysis",
            4: "Tables of Data",
            5: "Graph/Chart Preparation",
            6: "Reference Data",
            7: "Output"
        }
        
        for step_num, step_desc in steps.items():
            if step_num == st.session_state.current_step:
                st.markdown(f"**→ {step_num}. {step_desc}**")
            else:
                st.markdown(f"{step_num}. {step_desc}")
        
        # Navigation buttons
        col1, col2 = st.columns(2)
        
        with col1:
            if st.session_state.current_step > 1 and st.button("Previous Step"):
                st.session_state.current_step -= 1
                st.rerun()
        
        with col2:
            if st.session_state.current_step < 6 and st.session_state.data is not None and st.button("Next Step"):
                st.session_state.current_step += 1
                st.rerun()
                
        # Help information
        with st.expander("How to Use"):
            st.write("""
            **Excel Analysis Workflow:**
            
            1. **Upload Excel** - Upload your raw data (first sheet only)
            2. **Data Management** - Clean and prepare your data
            3. **Desired Analysis** - Choose the type of analysis to perform
            4. **Tables of Data** - View and customize data tables
            5. **Graph/Chart Preparation** - Create visualizations
            6. **Output** - Generate reports and export results
            
            This application replaces manual Excel analysis with an automated workflow.
            """)
    
    # Main content area based on current step
    if st.session_state.current_step == 1:
        # Step 1: Upload Excel
        st.header(emoji_header("1. Upload Excel", "Upload"))
        
        upload_col, sample_col = st.columns(2)
        
        with upload_col:
            st.subheader("Upload Your Data")
            # Simple file uploader
            uploaded_file = st.file_uploader("Upload your data file", type=['csv', 'xlsx', 'xls'])
            
            if uploaded_file is not None:
                try:
                    # Show loading animation and progress bar
                    with st.spinner(""):
                        # Display fun loading spinner with emoji
                        spinner_container = st.empty()
                        spinner_container.markdown(
                            f"""
                            <div style="text-align: center; margin-bottom: 20px;">
                                <div class="emoji-spinner">{show_spinner_emoji('Loading')}</div>
                                <p>Analyzing your data...</p>
                            </div>
                            """, 
                            unsafe_allow_html=True
                        )
                        
                        # Show playful progress bar
                        progress_bar = st.empty()
                        for i in range(0, 101, 10):
                            fun_progress_bar(i, "Preparing your data")
                            time.sleep(0.05)
                        
                        # Load data
                        df, filename = load_data(uploaded_file)
                        
                        # Finish loading animation
                        spinner_container.empty()
                        progress_bar.empty()
                        
                        # Show success with confetti
                        show_confetti()
                        success_emoji = random.choice(EMOJI_DICT["Success"])
                        st.success(f"{success_emoji} Data loaded successfully! Found {df.shape[0]} rows and {df.shape[1]} columns")
                        

                    
                    # DEBUG: Check if the file already has Age Group columns
                    age_group_cols = [col for col in df.columns if 'age group' in col.lower() or 'Age group' in col]
                    if age_group_cols:
                        st.warning(f"Found existing Age Group columns in the file: {age_group_cols}")
                    
                    # Store the original file for multi-sheet analysis
                    if filename.endswith('.xlsx') or filename.endswith('.xls'):
                        st.session_state.excel_file = uploaded_file
                    
                    # Store the original columns from the Excel file
                    st.session_state.original_columns = df.columns.tolist()
                    
                    # Store the dataframe
                    st.session_state.data = df
                    st.session_state.filename = filename
                    # Success message is already shown above with emoji and confetti
                    
                    # Show the complete raw data table with fixed width
                    st.subheader("🧮 Raw Data Table")
                    
                    # Create a fun data type display with emoji indicators
                    dt_col1, dt_col2, dt_col3, dt_col4 = st.columns(4)
                    with dt_col1:
                        st.markdown("**🔢 Numeric columns:** " + str(len(df.select_dtypes(include=['int64', 'float64']).columns)))
                    with dt_col2:
                        st.markdown("**🔤 Text columns:** " + str(len(df.select_dtypes(include=['object']).columns)))
                    with dt_col3:
                        st.markdown("**📅 Date columns:** " + str(len(df.select_dtypes(include=['datetime64']).columns)))
                    with dt_col4:
                        st.markdown("**✓✗ Boolean columns:** " + str(len(df.select_dtypes(include=['bool']).columns)))
                    
                    # Display dataframe in a mobile-friendly way
                    display_mobile_friendly_dataframe(df, title="Uploaded Data Preview", height=600)
                    
                    # Display data info in tabular manner
                    st.subheader("📊 Data Information")
                    
                    # Create a more structured table for data information
                    col_types = df.dtypes.value_counts()
                    missing_data = df.isnull().sum().sum()
                    
                    # Data info table
                    info_data = {
                        "Attribute": [
                            "Rows", 
                            "Columns", 
                            "Total Data Points",
                            "Missing Values",
                            "Numeric Columns",
                            "Text/Categorical Columns",
                            "Date/Time Columns",
                            "Boolean Columns"
                        ],
                        "Value": [
                            df.shape[0],
                            df.shape[1],
                            df.shape[0] * df.shape[1],
                            missing_data,
                            len(df.select_dtypes(include=['int64', 'float64']).columns),
                            len(df.select_dtypes(include=['object']).columns),
                            len(df.select_dtypes(include=['datetime64']).columns),
                            len(df.select_dtypes(include=['bool']).columns)
                        ]
                    }
                    
                    st.table(pd.DataFrame(info_data))
                    
                    if st.button("Proceed to Data Management"):
                        st.session_state.current_step = 2
                        # Add JavaScript to scroll to top after rerun
                        st.markdown(
                            """
                            <script>
                                window.scrollTo(0, 0);
                            </script>
                            """,
                            unsafe_allow_html=True
                        )
                        st.rerun()
                    
                except Exception as e:
                    st.error(f"Error loading data: {str(e)}")
        
        with sample_col:
            st.subheader("Use Sample Data")
            st.write("Don't have a dataset? Use our sample data to explore the application.")
            
            # Add a sample data option for demonstration
            if st.button("Use Sample Data (Iris Dataset)"):
                from sklearn.datasets import load_iris
                iris = load_iris()
                iris_df = pd.DataFrame(data=iris.data, columns=iris.feature_names)
                iris_df['target'] = iris.target
                iris_df['species'] = iris_df['target'].map({
                    0: 'setosa',
                    1: 'versicolor',
                    2: 'virginica'
                })
                # Reset index to start from 1 instead of 0
                iris_df.index = iris_df.index + 1
                
                st.session_state.data = iris_df
                st.session_state.filename = "iris_dataset.csv"
                st.success("Loaded sample Iris dataset")
                
                # Show the complete raw data table with fixed width
                # Display sample iris data in a mobile-friendly way
                display_mobile_friendly_dataframe(iris_df, title="Raw Data Table (Iris Sample)", height=600)
                
                # Display data info in tabular manner
                st.subheader("Data Information")
                
                # Create a more structured table for data information
                col_types = iris_df.dtypes.value_counts()
                missing_data = iris_df.isnull().sum().sum()
                
                # Data info table
                info_data = {
                    "Attribute": [
                        "Rows", 
                        "Columns", 
                        "Total Data Points",
                        "Missing Values",
                        "Numeric Columns",
                        "Text/Categorical Columns",
                        "Date/Time Columns",
                        "Boolean Columns"
                    ],
                    "Value": [
                        iris_df.shape[0],
                        iris_df.shape[1],
                        iris_df.shape[0] * iris_df.shape[1],
                        missing_data,
                        len(iris_df.select_dtypes(include=['int64', 'float64']).columns),
                        len(iris_df.select_dtypes(include=['object']).columns),
                        len(iris_df.select_dtypes(include=['datetime64']).columns),
                        len(iris_df.select_dtypes(include=['bool']).columns)
                    ]
                }
                
                st.table(pd.DataFrame(info_data))
                
                if st.button("Proceed to Data Management", key="proceed_sample"):
                    st.session_state.current_step = 2
                    # Add JavaScript to scroll to top after rerun
                    st.markdown(
                        """
                        <script>
                            window.scrollTo(0, 0);
                        </script>
                        """,
                        unsafe_allow_html=True
                    )
                    st.rerun()
    
    elif st.session_state.current_step == 2:
        # Step 2: Data Management
        st.header(emoji_header("2. Data Management", "Configuration"))
        
        if st.session_state.data is not None:
            df = st.session_state.data
            
            # Create tabs for different data management approaches
            tab1, tab2 = st.tabs(["Excel Drag & Drop Config", "Basic Column Selection"])
            
            # Tab 1: Excel Drag & Drop Config
            with tab1:
                # Use the drag & drop UI without introductory text
                display_drag_drop_ui()
                
                # Check if configuration is complete
                if 'configuration_complete' in st.session_state and st.session_state.configuration_complete:
                    # When configuration is complete, update selected columns with configured columns
                    
                    # Add age groups if age column is specified
                    if 'age_column' in st.session_state and st.session_state.age_column:
                        age_column = st.session_state.age_column['name']
                        age_type = st.session_state.age_type if 'age_type' in st.session_state else "years"
                        
                        # Add age column to selected columns
                        if age_column not in st.session_state.selected_columns:
                            st.session_state.selected_columns.append(age_column)
                        
                        # Age groups are now generated in drag_drop_ui.py when clicking Complete Configuration
                        # This section is intentionally disabled to prevent premature age group generation
                        pass
                    
                    # Add gender column
                    if 'gender_column' in st.session_state and st.session_state.gender_column:
                        if st.session_state.gender_column['name'] not in st.session_state.selected_columns:
                            st.session_state.selected_columns.append(st.session_state.gender_column['name'])
                    
                    # Add master columns
                    if 'master_columns' in st.session_state:
                        for col in st.session_state.master_columns:
                            if col['name'] not in st.session_state.selected_columns:
                                st.session_state.selected_columns.append(col['name'])
                    
                    # Add pair columns
                    if 'pair_areas' in st.session_state:
                        for pair in st.session_state.pair_areas:
                            for area_type in ['pre', 'post']:
                                for col in pair[area_type]:
                                    if col['name'] not in st.session_state.selected_columns:
                                        st.session_state.selected_columns.append(col['name'])
            
            # Tab 2: Basic Column Selection
            with tab2:
                # Display full dataframe with increased height to show all rows
                st.subheader("Full Dataset")
                st.dataframe(df, height=600)
                
                # Column selection for further analysis
                st.subheader("Select Columns for Analysis")
                all_cols = df.columns.tolist()
                st.session_state.selected_columns = st.multiselect(
                    "Select columns you want to include in your analysis:",
                    all_cols,
                    default=st.session_state.selected_columns if st.session_state.selected_columns else all_cols[:5]
                )
                
                if st.session_state.selected_columns:
                    # Show preview of selected columns
                    st.subheader("Preview of Selected Data")
                    st.dataframe(df[st.session_state.selected_columns].head())
                    
                    # Display column information
                    st.subheader("Column Information")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("Numerical Columns:")
                        numerical_cols = [col for col in st.session_state.selected_columns if pd.api.types.is_numeric_dtype(df[col])]
                        if numerical_cols:
                            st.write(", ".join(numerical_cols))
                        else:
                            st.write("No numerical columns selected")
                    
                    with col2:
                        st.write("Categorical Columns:")
                        categorical_cols = [col for col in st.session_state.selected_columns if not pd.api.types.is_numeric_dtype(df[col])]
                        if categorical_cols:
                            st.write(", ".join(categorical_cols))
                        else:
                            st.write("No categorical columns selected")
                    
                    # Check for missing values
                    st.subheader("Missing Values")
                    missing_values = df[st.session_state.selected_columns].isnull().sum()
                    if missing_values.sum() > 0:
                        st.write("Columns with missing values:")
                        missing_df = pd.DataFrame({
                            'Column': missing_values.index,
                            'Missing Values': missing_values.values,
                            'Percentage': (missing_values.values / len(df) * 100).round(2)
                        })
                        missing_df = missing_df[missing_df['Missing Values'] > 0].sort_values('Missing Values', ascending=False)
                        st.table(missing_df)
                        
                        # Option to handle missing values
                        st.subheader("Handle Missing Values")
                        missing_strategy = st.selectbox(
                            "Select strategy for handling missing values:",
                            ["None (Keep as is)", "Drop rows with missing values", "Fill with mean/mode"]
                        )
                        
                        if missing_strategy != "None (Keep as is)" and st.button("Apply Missing Value Strategy"):
                            if missing_strategy == "Drop rows with missing values":
                                df_clean = df.dropna(subset=st.session_state.selected_columns)
                                st.session_state.data = df_clean
                                st.success(f"Dropped rows with missing values. Remaining rows: {len(df_clean)}")
                                st.dataframe(df_clean[st.session_state.selected_columns].head())
                            
                            elif missing_strategy == "Fill with mean/mode":
                                df_filled = df.copy()
                                for col in st.session_state.selected_columns:
                                    if pd.api.types.is_numeric_dtype(df[col]):
                                        # Fill numeric columns with mean
                                        fill_value = df[col].mean()
                                        df_filled[col] = df_filled[col].fillna(fill_value)
                                        st.info(f"Filled '{col}' with mean: {fill_value:.2f}")
                                    else:
                                        # Fill categorical columns with mode
                                        fill_value = df[col].mode()[0] if not df[col].mode().empty else "Unknown"
                                        df_filled[col] = df_filled[col].fillna(fill_value)
                                        st.info(f"Filled '{col}' with mode: {fill_value}")
                                
                                st.session_state.data = df_filled
                                st.success("Filled missing values with mean/mode.")
                                st.dataframe(df_filled[st.session_state.selected_columns].head())
                    else:
                        st.write("No missing values found in the selected columns.")
                    
                    # Data types
                    st.subheader("Data Types")
                    dtypes_df = pd.DataFrame({
                        'Column': [col for col in st.session_state.selected_columns],
                        'Data Type': [str(df[col].dtype) for col in st.session_state.selected_columns]  # Convert dtype to string to prevent Arrow conversion issues
                    })
                    st.table(dtypes_df)
                else:
                    st.warning("Please select at least one column for analysis.")
            
            # Final dataset section is now moved into the drag & drop UI
            # and appears only after clicking "Complete Configuration"
            
            # Determine if we have master, age, age group, gender, pairs columns to show
            display_columns = []
            
            # Add master columns if available
            if 'master_columns' in st.session_state and st.session_state.master_columns:
                master_cols = [col['name'] for col in st.session_state.master_columns]
                display_columns.extend(master_cols)
                
            # Add age column if available
            if 'age_column' in st.session_state and st.session_state.age_column:
                age_col = st.session_state.age_column['name']
                if age_col not in display_columns:
                    display_columns.append(age_col)
                
                # Add only one Age Group column if any exists in the dataframe
                # Prioritize the Generated Age Group (our dynamically created one)
                age_group_cols = [col for col in df.columns if 'age group' in col.lower()]
                if age_group_cols:
                    # If we have our generated column, use that
                    if 'Generated Age Group' in age_group_cols:
                        display_columns.append('Generated Age Group')
                    # Otherwise use the first available age group column
                    else:
                        display_columns.append(age_group_cols[0])
            
            # Add gender column if available
            if 'gender_column' in st.session_state and st.session_state.gender_column:
                gender_col = st.session_state.gender_column['name']
                if gender_col not in display_columns:
                    display_columns.append(gender_col)
            
            # Add pair columns if available
            if 'pair_areas' in st.session_state and st.session_state.pair_areas:
                for pair in st.session_state.pair_areas:
                    for area_type in ['pre', 'post']:
                        for col in pair[area_type]:
                            if col['name'] not in display_columns:
                                display_columns.append(col['name'])
            
            # This section has been moved to drag_drop_ui.py and appears
            # only after the Complete Configuration button is clicked
            
            # Make sure we're using all selected columns
            all_display_columns = st.session_state.selected_columns.copy()
            
            # Add any display columns that aren't already in the selected columns
            existing_display_columns = [col for col in display_columns if col in df.columns]
            for col in existing_display_columns:
                if col not in all_display_columns:
                    all_display_columns.append(col)
            
            if all_display_columns:
                # Create a styled dataframe with background colors for different column types
                styled_df = df[all_display_columns].head(10).style
                
                # Apply background colors based on column types
                def highlight_columns(x):
                    styles = pd.DataFrame('', index=x.index, columns=x.columns)
                    
                    # Master columns (light blue)
                    if 'master_columns' in st.session_state and st.session_state.master_columns:
                        master_cols = [col['name'] for col in st.session_state.master_columns]
                        for col in master_cols:
                            if col in x.columns:
                                styles[col] = 'background-color: #e6f2ff'
                    
                    # Age column (light yellow)
                    if 'age_column' in st.session_state and st.session_state.age_column:
                        age_col = st.session_state.age_column['name']
                        if age_col in x.columns:
                            styles[age_col] = 'background-color: #ffffcc'
                    
                    # Age group columns (light orange)
                    age_group_cols = [col for col in x.columns if 'age group' in col.lower() or 'generated age group' in col.lower()]
                    for col in age_group_cols:
                        styles[col] = 'background-color: #ffebcc'
                    
                    # Gender column (light pink)
                    if 'gender_column' in st.session_state and st.session_state.gender_column:
                        gender_col = st.session_state.gender_column['name']
                        if gender_col in x.columns:
                            styles[gender_col] = 'background-color: #ffdddd'
                    
                    # Pair columns (light green shades for pre, light purple for post)
                    if 'pair_areas' in st.session_state and st.session_state.pair_areas:
                        for idx, pair in enumerate(st.session_state.pair_areas):
                            # Pre columns (light green with varying shades)
                            green_intensity = 200 - (idx * 10) % 50  # Vary the shade for different pairs
                            for col_obj in pair['pre']:
                                col = col_obj['name']
                                if col in x.columns:
                                    styles[col] = f'background-color: rgba(144, {green_intensity}, 144, 0.3)'
                            
                            # Post columns (light purple with varying shades)
                            purple_intensity = 200 - (idx * 10) % 50  # Vary the shade for different pairs
                            for col_obj in pair['post']:
                                col = col_obj['name']
                                if col in x.columns:
                                    styles[col] = f'background-color: rgba({purple_intensity}, 144, 240, 0.3)'
                    
                    return styles
                
                styled_df = styled_df.apply(highlight_columns, axis=None)
                
                # Display the styled dataframe
                st.dataframe(styled_df, height=400)
                
                # Show legend for colors
                st.write("**Column Legend:**")
                col1, col2, col3, col4, col5, col6 = st.columns(6)
                with col1:
                    st.markdown('<div style="background-color: #e6f2ff; padding: 5px;">Master Columns</div>', unsafe_allow_html=True)
                with col2:
                    st.markdown('<div style="background-color: #ffffcc; padding: 5px;">Age Column</div>', unsafe_allow_html=True)
                with col3:
                    st.markdown('<div style="background-color: #ffebcc; padding: 5px;">Age Group Column</div>', unsafe_allow_html=True)
                with col4:
                    st.markdown('<div style="background-color: #ffdddd; padding: 5px;">Gender Column</div>', unsafe_allow_html=True)
                with col5:
                    st.markdown('<div style="background-color: rgba(144, 200, 144, 0.3); padding: 5px;">Pre Pairs</div>', unsafe_allow_html=True)
                with col6:
                    st.markdown('<div style="background-color: rgba(200, 144, 240, 0.3); padding: 5px;">Post Pairs</div>', unsafe_allow_html=True)
            else:
                st.warning("No columns selected for analysis.")
            
            # Option to continue
            if st.button("Continue to Data Analysis"):
                if not st.session_state.selected_columns:
                    st.error("Please select at least one column for analysis before continuing.")
                else:
                    st.session_state.current_step = 3
                    # Add JavaScript to scroll to top after rerun
                    st.markdown(
                        """
                        <script>
                            window.scrollTo(0, 0);
                        </script>
                        """,
                        unsafe_allow_html=True
                    )
                    st.rerun()
    
    elif st.session_state.current_step == 3:
        # Step 3: Desired Data Analysis
        st.header(emoji_header("3. Desired Data Analysis", "Analysis"))
        
        if st.session_state.data is not None:
            df = st.session_state.data
            
            # If columns weren't already selected, default to ALL columns
            if not st.session_state.selected_columns:
                st.session_state.selected_columns = df.columns.tolist()
            
            # Create analysis selection options
            st.subheader("Select Analysis Type")
            
            analysis_options = [
                "Descriptive Statistics",
                "Statistical Tests", 
                "Correlation Analysis",
                "P-Value Analysis",
                "Automated Analysis"
            ]
            
            st.session_state.analysis_type = st.radio(
                "What type of analysis would you like to perform?",
                analysis_options,
                index=analysis_options.index(st.session_state.analysis_type) if st.session_state.analysis_type in analysis_options else 0
            )
            
            if st.session_state.analysis_type == "Descriptive Statistics":
                st.write("You'll get summary statistics for your selected columns, master parameter analysis, and pre/post analysis.")
            
            elif st.session_state.analysis_type == "Statistical Tests":
                st.write("You'll be able to perform hypothesis tests (t-tests, ANOVA) on your data.")
                
                # Allow user to select specific test
                test_options = ["t-Test", "ANOVA", "Chi-Square Test"]
                test_type = st.selectbox("Select specific test:", test_options)
                
                if test_type == "t-Test":
                    # t-test setup
                    # Use all columns for dropdown selection
                    all_cols = st.session_state.selected_columns
                    # Get only numeric columns for statistical testing
                    num_cols = [col for col in st.session_state.selected_columns if pd.api.types.is_numeric_dtype(df[col])]
                    ttest_settings = {}
                    
                    # Check if we have numeric columns for testing
                    if not num_cols:
                        st.error("No numeric columns available for t-test analysis. Please include numeric columns in your selection.")
                    elif all_cols:
                        ttest_settings['test_type'] = st.radio(
                            "Select t-test type:",
                            ["One-sample t-test", "Two-sample t-test"]
                        )
                        
                        if ttest_settings['test_type'] == "One-sample t-test":
                            # Use only numeric columns for the test variable
                            ttest_settings['column'] = st.selectbox(
                                "Select column for test:", 
                                num_cols,
                                help="Only numeric columns are shown as they are required for t-tests."
                            )
                            ttest_settings['mu'] = st.number_input("Population mean (μ₀):", value=0.0)
                            ttest_settings['alpha'] = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                        
                        else:  # Two-sample t-test
                            ttest_settings['method'] = st.radio(
                                "Select method:",
                                ["Compare two columns", "Compare groups within a column"]
                            )
                            
                            if ttest_settings['method'] == "Compare two columns":
                                # Use only numeric columns for column comparison
                                ttest_settings['col1'] = st.selectbox(
                                    "Select first column:", 
                                    num_cols,
                                    help="Only numeric columns are shown as they are required for t-tests."
                                )
                                # Filter out the first column from numeric columns
                                remaining_num_cols = [col for col in num_cols if col != ttest_settings['col1']]
                                if remaining_num_cols:
                                    ttest_settings['col2'] = st.selectbox(
                                        "Select second column:", 
                                        remaining_num_cols,
                                        help="Only numeric columns are shown as they are required for t-tests."
                                    )
                                    ttest_settings['equal_var'] = st.checkbox("Assume equal variances", value=False)
                                    ttest_settings['alpha'] = st.slider("Significance level (α):", 0.01, 0.10, 0.05, key="ttest_alpha2")
                                else:
                                    st.warning("Need at least two numeric columns for two-sample t-test.")
                            
                            else:  # Compare groups
                                # Use only numeric columns for the dependent variable
                                ttest_settings['num_col'] = st.selectbox(
                                    "Select numeric column to analyze:", 
                                    num_cols,
                                    help="Only numeric columns are shown as they are required for t-tests."
                                )
                                # Use all columns for grouping
                                ttest_settings['group_col'] = st.selectbox("Select grouping column:", all_cols)
                                # Get unique values for the selected grouping column
                                unique_groups = df[ttest_settings['group_col']].dropna().unique().tolist()
                                
                                if len(unique_groups) >= 2:
                                    ttest_settings['group1'] = st.selectbox("Select first group:", unique_groups)
                                    # Filter out the first group
                                    remaining_groups = [g for g in unique_groups if g != ttest_settings['group1']]
                                    ttest_settings['group2'] = st.selectbox("Select second group:", remaining_groups)
                                    ttest_settings['equal_var'] = st.checkbox("Assume equal variances", value=False, key="ttest_eq_var2")
                                    ttest_settings['alpha'] = st.slider("Significance level (α):", 0.01, 0.10, 0.05, key="ttest_alpha3")
                                else:
                                    st.warning(f"Need at least two groups in {ttest_settings['group_col']} for two-sample t-test.")
                    else:
                        st.warning("No columns available for analysis.")
                    
                    # Store the settings in session state for later use
                    st.session_state.ttest_settings = ttest_settings
                
                elif test_type == "ANOVA":
                    # ANOVA setup
                    all_cols = st.session_state.selected_columns
                    # Get only numeric columns for analysis (needed for dependent variable)
                    num_cols = [col for col in st.session_state.selected_columns if pd.api.types.is_numeric_dtype(df[col])]
                    anova_settings = {}
                    
                    # Check if we have numeric columns for testing
                    if not num_cols:
                        st.error("No numeric columns available for ANOVA analysis. Please include numeric columns in your selection.")
                    elif all_cols:
                        # Use only numeric columns for the dependent variable
                        anova_settings['num_col'] = st.selectbox(
                            "Select column for analysis (dependent variable):", 
                            num_cols,
                            help="Only numeric columns are shown as the dependent variable must be numeric for ANOVA."
                        )
                        
                        # All columns can be used as categorical factors
                        anova_settings['cat_col'] = st.selectbox(
                            "Select categorical column (factor):", 
                            all_cols,
                            help="The factor column should have distinct groups/categories."
                        )
                        
                        # Check if the categorical column has enough groups
                        unique_groups = df[anova_settings['cat_col']].dropna().unique()
                        if len(unique_groups) < 3:
                            st.warning(f"ANOVA typically requires at least 3 groups, but '{anova_settings['cat_col']}' only has {len(unique_groups)} unique values.")
                        
                        # Display the unique groups found
                        with st.expander("View factor levels"):
                            st.write(f"Factor '{anova_settings['cat_col']}' has these unique values:")
                            st.write(", ".join([str(g) for g in unique_groups]))
                        
                        anova_settings['alpha'] = st.slider("Significance level (α):", 0.01, 0.10, 0.05, key="anova_alpha")
                    else:
                        st.warning("No columns available for ANOVA.")
                    
                    # Store the settings in session state for later use
                    st.session_state.anova_settings = anova_settings
                    
                elif test_type == "Chi-Square Test":
                    # Chi-Square test setup - needs categorical columns
                    all_cols = st.session_state.selected_columns
                    
                    # Get categorical columns - any column that's not clearly numeric
                    cat_cols = [col for col in all_cols if not pd.api.types.is_numeric_dtype(df[col]) or df[col].nunique() < 10]
                    
                    chi2_settings = {}
                    
                    # Check if we have categorical columns for testing
                    if not cat_cols:
                        st.error("No categorical columns identified for Chi-Square Test. Please include categorical columns in your selection.")
                    elif len(cat_cols) < 2:
                        st.error("Chi-Square Test requires at least two categorical variables to test for association.")
                    else:
                        # Create a simple interface with dropdowns
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            chi2_settings['col1'] = st.selectbox(
                                "Select first categorical column:", 
                                cat_cols,
                                help="Select a categorical variable to test for association."
                            )
                            
                            # Show the unique values in this column
                            unique_vals1 = df[chi2_settings['col1']].dropna().unique()
                            st.write(f"**Unique values:** {len(unique_vals1)}")
                            if len(unique_vals1) <= 10:  # Only show if not too many values
                                st.write(", ".join([str(val) for val in unique_vals1]))
                        
                        # Filter to exclude the first column
                        remaining_cat_cols = [col for col in cat_cols if col != chi2_settings['col1']]
                        
                        with col2:
                            chi2_settings['col2'] = st.selectbox(
                                "Select second categorical column:", 
                                remaining_cat_cols,
                                help="Select another categorical variable to test for association with the first."
                            )
                            
                            # Show the unique values in this column
                            unique_vals2 = df[chi2_settings['col2']].dropna().unique()
                            st.write(f"**Unique values:** {len(unique_vals2)}")
                            if len(unique_vals2) <= 10:  # Only show if not too many values
                                st.write(", ".join([str(val) for val in unique_vals2]))
                        
                        # Preview contingency table
                        st.write("**Preview of Contingency Table:**")
                        preview_table = pd.crosstab(
                            df[chi2_settings['col1']].dropna(), 
                            df[chi2_settings['col2']].dropna()
                        )
                        
                        # If table is too large, show a message
                        if preview_table.shape[0] > 10 or preview_table.shape[1] > 10:
                            st.info(f"Table is large ({preview_table.shape[0]}×{preview_table.shape[1]}). Showing partial view.")
                            st.dataframe(preview_table.iloc[:10, :10])
                        else:
                            st.dataframe(preview_table)
                        
                        # Set significance level
                        chi2_settings['alpha'] = st.slider(
                            "Significance level (α):", 
                            0.01, 0.10, 0.05, 
                            key="chi2_alpha"
                        )
                        
                        # Store the settings in session state for later use
                        st.session_state.chi2_settings = chi2_settings
            
            elif st.session_state.analysis_type == "Correlation Analysis":
                st.write("You'll analyze relationships between numerical variables.")
                
                # Correlation settings
                all_cols = st.session_state.selected_columns
                corr_settings = {}
                
                if len(all_cols) >= 2:
                    corr_settings['columns'] = st.multiselect(
                        "Select columns for correlation analysis:",
                        all_cols,
                        default=all_cols  # Default to all columns 
                    )
                    
                    if len(corr_settings['columns']) >= 2:
                        corr_settings['method'] = st.selectbox(
                            "Correlation method:",
                            ["pearson", "spearman", "kendall"],
                            index=0
                        )
                        st.write("Pearson: Linear correlation, assumes normal distribution")
                        st.write("Spearman: Rank correlation, resistant to outliers")
                        st.write("Kendall: Another rank correlation, better for small samples")
                        
                        corr_settings['alpha'] = st.slider("Significance level (α):", 0.01, 0.10, 0.05, key="corr_alpha")
                    else:
                        st.warning("Please select at least two columns for correlation analysis.")
                else:
                    st.warning("Need at least two columns for correlation analysis.")
                
                # Store the settings in session state for later use
                st.session_state.corr_settings = corr_settings
                
            elif st.session_state.analysis_type == "P-Value Analysis":
                st.write("Calculate overall p-values for each parameter and system-wise p-values.")
                
                # P-value analysis settings
                # Use all available columns from the dataset, not just selected ones
                all_cols = df.columns.tolist()
                pvalue_settings = {}
                
                # Show all available columns for p-value analysis (will filter numeric internally)
                if len(all_cols) >= 1:
                    # Default to numeric columns for better user experience
                    numeric_cols = [col for col in all_cols if pd.api.types.is_numeric_dtype(df[col])]
                    default_cols = numeric_cols if numeric_cols else all_cols[:5]
                    
                    pvalue_settings['columns'] = st.multiselect(
                        "Select columns for p-value analysis:",
                        all_cols,
                        default=default_cols
                    )
                    
                    if len(pvalue_settings['columns']) >= 1:
                        # Test type selection for p-value analysis
                        pvalue_settings['test_type'] = st.selectbox(
                            "Select test type for p-value calculation:",
                            ["Normality Test (Shapiro-Wilk)", "One-Sample T-Test", "Two-Sample T-Test", "ANOVA"]
                        )
                        
                        # Additional settings based on test type
                        if pvalue_settings['test_type'] == "One-Sample T-Test":
                            pvalue_settings['test_value'] = st.number_input(
                                "Test value (μ₀):", 
                                value=0.0,
                                help="The hypothesized population mean to test against"
                            )
                        
                        elif pvalue_settings['test_type'] == "Two-Sample T-Test":
                            # Check if we have grouping columns
                            categorical_cols = [col for col in all_cols if df[col].dtype == 'object' or col.lower() in ['system', 'group', 'category']]
                            if categorical_cols:
                                pvalue_settings['group_column'] = st.selectbox(
                                    "Select grouping column:",
                                    categorical_cols
                                )
                            else:
                                st.warning("No categorical columns found for grouping. Consider using One-Sample T-Test instead.")
                        
                        elif pvalue_settings['test_type'] == "ANOVA":
                            # Check if we have grouping columns
                            categorical_cols = [col for col in all_cols if df[col].dtype == 'object' or col.lower() in ['system', 'group', 'category']]
                            if categorical_cols:
                                pvalue_settings['group_column'] = st.selectbox(
                                    "Select grouping column for ANOVA:",
                                    categorical_cols
                                )
                            else:
                                st.warning("No categorical columns found for grouping. ANOVA requires a grouping variable.")
                        
                        # Significance level
                        pvalue_settings['alpha'] = st.slider(
                            "Significance level (α):", 
                            0.01, 0.10, 0.05, 
                            key="pvalue_alpha"
                        )
                        
                        # Check for master columns for system-wise analysis
                        master_cols = []
                        
                        # Check if 'System' column exists
                        if 'System' in df.columns:
                            master_cols.append('System')
                        
                        # Check if age groups were generated
                        if 'Generated Age Group' in df.columns:
                            master_cols.append('Generated Age Group')
                        
                        # Check for any other potential master columns (categorical with reasonable number of groups)
                        for col in df.columns:
                            if col not in master_cols and df[col].dtype == 'object' and 2 <= df[col].nunique() <= 10:
                                if col.lower() in ['group', 'category', 'type', 'class', 'gender', 'treatment']:
                                    master_cols.append(col)
                        
                        if master_cols:
                            st.info(f"System-wise analysis will be performed using: {', '.join(master_cols)}")
                            pvalue_settings['master_columns'] = master_cols
                        else:
                            st.info("No master columns identified. Only overall p-values will be calculated.")
                            pvalue_settings['master_columns'] = []
                    else:
                        st.warning("Please select at least one column for p-value analysis.")
                else:
                    st.warning("No numeric columns available for p-value analysis.")
                
                # Store the settings in session state for later use
                st.session_state.pvalue_settings = pvalue_settings
                
            elif st.session_state.analysis_type == "Automated Analysis":
                st.write("We'll automatically analyze your data and suggest appropriate statistical tests.")
                
                # Show data structure analysis results
                data_analysis = analyze_data_structure(df)
                
                # Show data role detection (input/process/output)
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.subheader("Input Columns")
                    if 'input_columns' in data_analysis and data_analysis['input_columns']:
                        for col in data_analysis['input_columns'][:10]:
                            st.write(f"- {col}")
                        if len(data_analysis['input_columns']) > 10:
                            st.write(f"...and {len(data_analysis['input_columns']) - 10} more")
                    else:
                        st.write("No input columns detected")
                
                with col2:
                    st.subheader("Process Columns")
                    if 'process_columns' in data_analysis and data_analysis['process_columns']:
                        for col in data_analysis['process_columns'][:10]:
                            st.write(f"- {col}")
                        if len(data_analysis['process_columns']) > 10:
                            st.write(f"...and {len(data_analysis['process_columns']) - 10} more")
                    else:
                        st.write("No process columns detected")
                
                with col3:
                    st.subheader("Output Columns")
                    if 'output_columns' in data_analysis and data_analysis['output_columns']:
                        for col in data_analysis['output_columns'][:10]:
                            st.write(f"- {col}")
                        if len(data_analysis['output_columns']) > 10:
                            st.write(f"...and {len(data_analysis['output_columns']) - 10} more")
                    else:
                        st.write("No output columns detected")
                
                # Show detected calculation patterns
                if 'calculation_patterns' in data_analysis and data_analysis['calculation_patterns']:
                    st.subheader("Detected Calculation Patterns")
                    patterns = data_analysis['calculation_patterns']
                    
                    # Group patterns by type
                    pattern_types = {}
                    for pattern in patterns:
                        pattern_type = pattern['type']
                        if pattern_type not in pattern_types:
                            pattern_types[pattern_type] = []
                        pattern_types[pattern_type].append(pattern)
                    
                    # Show patterns by type
                    for pattern_type, type_patterns in pattern_types.items():
                        with st.expander(f"{pattern_type.replace('_', ' ').title()} ({len(type_patterns)} found)"):
                            for i, pattern in enumerate(type_patterns[:10]):
                                if pattern_type == 'pre_post_comparison':
                                    st.write(f"{i+1}. {pattern['pre_column']} → {pattern['post_column']}")
                                elif pattern_type == 'difference_calculation' or pattern_type == 'percentage_calculation':
                                    st.write(f"{i+1}. {' & '.join([str(col) for col in pattern['input_columns']])} → {pattern['output_column']}")
                                else:
                                    st.write(f"{i+1}. {pattern['output_column']}")
                            
                            if len(type_patterns) > 10:
                                st.write(f"...and {len(type_patterns) - 10} more")
                
                # Automatic master parameter analysis
                st.subheader("Master Parameter Analysis")
                
                # Detect potential master parameters
                master_candidates = []
                
                # Check if 'System' column exists
                if 'System' in df.columns:
                    master_candidates.append('System')
                
                # Only auto-detect Gender if specifically requested by the user
                # Gender auto-detection is now disabled to respect user's explicit configuration
                
                # Check if age groups were generated
                if 'Generated Age Group' in df.columns:
                    master_candidates.append('Generated Age Group')
                
                if master_candidates:
                    st.write(f"Automatically analyzing by master parameters: {', '.join(master_candidates)}")
                    
                    # For each master parameter, perform analysis
                    for master_param in master_candidates:
                        with st.expander(f"Analysis by {master_param}", expanded=False):
                            # Call analyze_master_parameters for this master parameter
                            master_results = analyze_master_parameters(df, [master_param])
                            
                            if not master_results or master_param not in master_results:
                                st.error(f"Error analyzing master parameter {master_param}.")
                            else:
                                results = master_results[master_param]
                                col_type = results.get('type', 'unknown')
                                st.write(f"Parameter type: {col_type}")
                                
                                # Check if there's only one group
                                groups = results.get('groups', {})
                                if len(groups) == 1:
                                    group_name = list(groups.keys())[0]
                                    st.info(f"⚠️ Note: The dataset only contains one value for {master_param} ({group_name}).")
                                
                                # Sort groups by age for age group columns, or alphabetically for other columns
                                sorted_group_items = []
                                
                                if master_param == 'Generated Age Group' or 'age group' in master_param.lower():
                                    # For age groups, sort them numerically by extracting the starting age
                                    def extract_start_age(group_name):
                                        try:
                                            # Extract the number before the dash
                                            if '-' in group_name:
                                                start_age = group_name.split('-')[0].strip()
                                                # Convert to number
                                                return int(start_age)
                                            else:
                                                return 999  # For any group that doesn't follow the pattern
                                        except:
                                            return 999  # Default high value for sorting
                                    
                                    # Sort groups by the starting age
                                    sorted_group_items = sorted(groups.items(), key=lambda x: extract_start_age(x[0]))
                                    print(f"Age groups sorted order: {[g[0] for g in sorted_group_items]}")
                                else:
                                    # For non-age groups, use regular dictionary items
                                    sorted_group_items = groups.items()
                                
                                # Display results for each group
                                for group_name, group_data in sorted_group_items:
                                    st.write(f"**{master_param} = {group_name} (n={group_data['count']})**")
                                    group_stats = group_data.get('stats')
                                    if group_stats is not None:
                                        # Display statistics
                                        st.dataframe(group_stats)
                                    else:
                                        st.write("No statistics available for this group.")
                
                # Pre/Post Analysis
                st.subheader("Pre/Post Comparisons")
                
                # MANUAL DESCRIPTIVE ANALYSIS: 
                # Only use manually configured pairs from the drag & drop UI
                # This is different from the automated ML analysis which detects pairs automatically
                pre_post_pairs = []
                if 'pair_areas' in st.session_state and st.session_state.pair_areas:
                    for pair_area in st.session_state.pair_areas:
                        pre_cols = [col['name'] for col in pair_area.get('pre', [])]
                        post_cols = [col['name'] for col in pair_area.get('post', [])]
                        
                        # Match each pre column with a post column if possible
                        for i in range(min(len(pre_cols), len(post_cols))):
                            if pre_cols[i] in df.columns and post_cols[i] in df.columns:
                                pre_post_pairs.append((pre_cols[i], post_cols[i]))
                
                if pre_post_pairs:
                    st.write(f"Analyzing {len(pre_post_pairs)} manually configured pre/post pairs...")
                    
                    # Create a dataframe to store the results
                    pre_post_results = []
                    
                    for pre_col, post_col in pre_post_pairs:
                        if pre_col in df.columns and post_col in df.columns:
                            # Ensure columns are numeric
                            if pd.api.types.is_numeric_dtype(df[pre_col]) and pd.api.types.is_numeric_dtype(df[post_col]):
                                # Calculate difference
                                diff = df[pre_col] - df[post_col]
                                pct_change = (diff / df[pre_col]) * 100 if df[pre_col].min() != 0 else None
                                
                                # Calculate statistics
                                mean_diff = diff.mean()
                                std_diff = diff.std()
                                min_diff = diff.min()
                                max_diff = diff.max()
                                
                                # Add to results
                                pre_post_results.append({
                                    'Pre Column': pre_col,
                                    'Post Column': post_col,
                                    'Mean Diff': round(mean_diff, 3),
                                    'Std Dev': round(std_diff, 3),
                                    'Min Diff': round(min_diff, 3),
                                    'Max Diff': round(max_diff, 3),
                                    'Mean % Change': round(pct_change.mean(), 2) if pct_change is not None else 'NA'
                                })
                    
                    if pre_post_results:
                        # Show pre/post analysis results
                        st.write("Pre/Post Difference Analysis:")
                        st.dataframe(pd.DataFrame(pre_post_results))
                        
                        # Show pre/post differences by master parameter if available
                        if master_candidates:
                            st.write("Pre/Post Differences by Master Parameter:")
                            
                            for master_param in master_candidates:
                                with st.expander(f"Pre/Post Analysis by {master_param}"):
                                    # Group by master parameter
                                    for group_name, group_df in df.groupby(master_param):
                                        st.write(f"**{master_param} = {group_name} (n={len(group_df)})**")
                                        
                                        # Create a dataframe for this group's pre/post results
                                        group_results = []
                                        
                                        for pre_col, post_col in pre_post_pairs:
                                            if pre_col in group_df.columns and post_col in group_df.columns:
                                                # Ensure columns are numeric
                                                if pd.api.types.is_numeric_dtype(group_df[pre_col]) and pd.api.types.is_numeric_dtype(group_df[post_col]):
                                                    # Calculate difference
                                                    diff = group_df[pre_col] - group_df[post_col]
                                                    pct_change = (diff / group_df[pre_col]) * 100 if group_df[pre_col].min() != 0 else None
                                                    
                                                    # Add to results
                                                    group_results.append({
                                                        'Pre Column': pre_col,
                                                        'Post Column': post_col,
                                                        'Mean Diff': round(diff.mean(), 3),
                                                        'Std Dev': round(diff.std(), 3),
                                                        'Min Diff': round(diff.min(), 3),
                                                        'Max Diff': round(diff.max(), 3),
                                                        'Mean % Change': round(pct_change.mean(), 2) if pct_change is not None else 'NA'
                                                    })
                                        
                                        if group_results:
                                            st.dataframe(pd.DataFrame(group_results))
                                        else:
                                            st.write("No pre/post results available for this group.")
                else:
                    st.info("No pre/post pairs detected in the data.")
                
                # Option to analyze all sheets if this is an Excel file
                if 'excel_file' in st.session_state:
                    st.subheader("Multi-Sheet Analysis")
                    analyze_all_sheets = st.checkbox("Analyze all sheets in the Excel file")
                    
                    if analyze_all_sheets:
                        try:
                            # Read the file with multiple sheets
                            uploaded_file = st.session_state.excel_file
                            from excel_analyzer import load_excel_file, analyze_excel_sheet
                            
                            with st.spinner("Loading and analyzing all sheets..."):
                                # Load all sheets from the Excel file
                                sheets_data, sheets_metadata = load_excel_file(uploaded_file)
                                
                                # Display sheet information
                                st.subheader("Excel File Structure")
                                st.write(f"Found {len(sheets_data)} sheets in the workbook.")
                                
                                # Analyze each sheet
                                all_sheets_analysis = {}
                                for sheet_name, sheet_df in sheets_data.items():
                                    with st.expander(f"Sheet: {sheet_name} ({sheet_df.shape[0]} rows, {sheet_df.shape[1]} columns)"):
                                        # Basic info about this sheet
                                        st.write(f"**Rows:** {sheet_df.shape[0]}, **Columns:** {sheet_df.shape[1]}")
                                        
                                        # Analyze data structure for this sheet
                                        sheet_analysis = analyze_data_structure(sheet_df)
                                        all_sheets_analysis[sheet_name] = sheet_analysis
                                        
                                        # Column summary for this sheet
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.write("**Numeric Columns:**")
                                            if sheet_analysis['numeric_columns']:
                                                for col in sheet_analysis['numeric_columns'][:10]:  # Limit display to first 10
                                                    st.write(f"- {col}")
                                                if len(sheet_analysis['numeric_columns']) > 10:
                                                    st.write(f"... and {len(sheet_analysis['numeric_columns']) - 10} more")
                                            else:
                                                st.write("None found")
                                        
                                        with col2:
                                            st.write("**Categorical Columns:**")
                                            if sheet_analysis['categorical_columns']:
                                                for col in sheet_analysis['categorical_columns'][:10]:  # Limit display to first 10
                                                    st.write(f"- {col}")
                                                if len(sheet_analysis['categorical_columns']) > 10:
                                                    st.write(f"... and {len(sheet_analysis['categorical_columns']) - 10} more")
                                            else:
                                                st.write("None found")
                                        
                                        # Display a preview of the data
                                        with st.expander("Data Preview"):
                                            st.dataframe(sheet_df.head(5))
                                        
                                        # Show suggested analyses for this sheet
                                        st.write("**Suggested Analyses:**")
                                        if 'suggested_analyses' in sheet_analysis:
                                            for analysis in sheet_analysis['suggested_analyses'][:5]:  # Limit to first 5 suggestions
                                                st.write(f"- {analysis.replace('_', ' ').title()}")
                                        else:
                                            st.write("No analyses suggested for this sheet.")
                                
                                # Cross-sheet analysis
                                st.subheader("Cross-Sheet Analysis")
                                
                                # Check for common columns across sheets
                                all_columns = {sheet: set(df.columns) for sheet, df in sheets_data.items()}
                                common_columns = set.intersection(*all_columns.values()) if all_columns else set()
                                
                                if common_columns:
                                    st.write(f"**Common Columns Across All Sheets ({len(common_columns)}):**")
                                    for col in sorted(common_columns)[:10]:  # Limit display to first 10
                                        st.write(f"- {col}")
                                    if len(common_columns) > 10:
                                        st.write(f"... and {len(common_columns) - 10} more")
                                    
                                    # Suggest cross-sheet comparison for a selected common column
                                    if len(common_columns) > 0:
                                        selected_column = st.selectbox("Select a common column for cross-sheet comparison:", sorted(common_columns))
                                        
                                        # Check if selected column is numeric
                                        numeric_in_all_sheets = True
                                        for sheet_name, sheet_df in sheets_data.items():
                                            if not pd.api.types.is_numeric_dtype(sheet_df[selected_column]):
                                                numeric_in_all_sheets = False
                                                break
                                        
                                        if numeric_in_all_sheets:
                                            # Create dataframe for comparison
                                            comparison_data = []
                                            for sheet_name, sheet_df in sheets_data.items():
                                                comparison_data.append({
                                                    'Sheet': sheet_name,
                                                    'Mean': sheet_df[selected_column].mean(),
                                                    'Median': sheet_df[selected_column].median(),
                                                    'Std Dev': sheet_df[selected_column].std(),
                                                    'Min': sheet_df[selected_column].min(),
                                                    'Max': sheet_df[selected_column].max(),
                                                    'Count': sheet_df[selected_column].count()
                                                })
                                            
                                            comparison_df = pd.DataFrame(comparison_data)
                                            st.write(f"**Comparison of {selected_column} Across Sheets:**")
                                            st.dataframe(comparison_df)
                                            
                                            # Create a visualization
                                            fig = px.bar(comparison_df, x='Sheet', y='Mean', 
                                                        error_y='Std Dev',
                                                        title=f"Mean of {selected_column} Across Sheets",
                                                        labels={'Mean': f'Mean {selected_column}'})
                                            st.plotly_chart(fig)
                                        else:
                                            st.write(f"Column '{selected_column}' is not numeric in all sheets.")
                                else:
                                    st.write("No common columns found across all sheets.")
                                
                                # Store the multi-sheet analysis results for later use
                                st.session_state.all_sheets_analysis = all_sheets_analysis
                                st.session_state.sheets_data = sheets_data
                                
                                # Add ML model training based on multiple sheets
                                st.subheader("Train ML Model Using Multi-Sheet Data")
                                with st.expander("Cross-Sheet ML Model Training", expanded=True):
                                    st.write("This will analyze relationships between sheets and train a machine learning model")
                                    
                                    # Import the ML functions we need
                                    from ml_utils import train_cross_sheet_ml_model, identify_master_sheet
                                    
                                    # Try to automatically identify the master sheet
                                    master_sheet_name, master_info = identify_master_sheet(sheets_data)
                                    
                                    if master_sheet_name:
                                        st.success(f"Identified master sheet: **{master_sheet_name}** (Score: {master_info['score']})")
                                        
                                        # Display reasons for selection
                                        if 'info' in master_info:
                                            reasons = []
                                            for key, value in master_info['info'].items():
                                                if key == 'master_in_name' and value:
                                                    reasons.append("Contains 'master' in sheet name")
                                                elif key == 'id_columns' and value:
                                                    reasons.append(f"Contains ID columns: {', '.join(value[:3])}" + 
                                                                  (f" and {len(value)-3} more" if len(value) > 3 else ""))
                                                elif key == 'system_columns' and value:
                                                    reasons.append(f"Contains system columns: {', '.join(value[:3])}" + 
                                                                  (f" and {len(value)-3} more" if len(value) > 3 else ""))
                                                elif key == 'compact_format' and value:
                                                    reasons.append("Has compact format (fewer rows, more columns)")
                                                elif key == 'first_sheet' and value:
                                                    reasons.append("Is the first sheet in the workbook")
                                            
                                            if reasons:
                                                st.write("**Selection reasons:**")
                                                for r in reasons:
                                                    st.write(f"- {r}")
                                    else:
                                        st.warning("Could not automatically identify a master sheet")
                                    
                                    # Option to manually select master sheet
                                    allow_manual_selection = st.checkbox("Manually select master sheet", value=not bool(master_sheet_name))
                                    
                                    if allow_manual_selection:
                                        master_sheet_name = st.selectbox(
                                            "Select the master sheet:", 
                                            list(sheets_data.keys()),
                                            index=list(sheets_data.keys()).index(master_sheet_name) if master_sheet_name else 0
                                        )
                                    
                                    # Train the cross-sheet model
                                    if st.button("Train ML Model on All Sheets"):
                                        with st.spinner("Training ML model across sheets..."):
                                            # Call the function to train a model using all sheets
                                            model_results = train_cross_sheet_ml_model(
                                                sheets_data=sheets_data,
                                                master_sheet_name=master_sheet_name
                                            )
                                            
                                            # Store and display results
                                            st.session_state.cross_sheet_model = model_results
                                            
                                            if 'error' in model_results:
                                                st.error(f"Error training model: {model_results['error']}")
                                            else:
                                                st.success("ML model successfully trained!")
                                                
                                                # Display model information
                                                st.write("### Model Information")
                                                if 'model_info' in model_results:
                                                    model_info = model_results['model_info']
                                                    
                                                    st.write(f"**Model Type:** {model_info.get('type', 'Unknown')}")
                                                    st.write(f"**Algorithm:** {model_info.get('algorithm', 'Unknown')}")
                                                    
                                                    # Display target column
                                                    if 'target_column' in model_results:
                                                        target_info = model_results['target_column']
                                                        st.write(f"**Target Column:** {target_info.get('name', 'Unknown')}")
                                                        
                                                        if 'auto_selected' in target_info and target_info['auto_selected']:
                                                            st.write("*Target column was automatically selected*")
                                                            if 'alternatives' in target_info and target_info['alternatives']:
                                                                st.write(f"Alternative targets: {', '.join(target_info['alternatives'])}")
                                                
                                                # Display feature importance
                                                if 'feature_importance' in model_info:
                                                    st.write("### Feature Importance")
                                                    
                                                    # Create a dataframe with feature importance
                                                    importance_data = []
                                                    for feature, importance in model_info['feature_importance'][:10]:  # Show top 10
                                                        importance_data.append({
                                                            'Feature': feature,
                                                            'Importance': importance
                                                        })
                                                    
                                                    importance_df = pd.DataFrame(importance_data)
                                                    
                                                    # Create bar chart
                                                    fig = px.bar(
                                                        importance_df, 
                                                        x='Feature', 
                                                        y='Importance',
                                                        title='Feature Importance'
                                                    )
                                                    
                                                    st.plotly_chart(fig)
                                                    
                                                # Show merged data summary
                                                if 'merged_data_shape' in model_results:
                                                    st.write(f"**Merged Data Shape:** {model_results['merged_data_shape'][0]} rows × {model_results['merged_data_shape'][1]} columns")
                        
                        except Exception as e:
                            st.error(f"Error analyzing all sheets: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())
                        
                        # No need to continue with single-sheet analysis if we're analyzing all sheets
                        st.session_state.analysis_results = {
                            'type': 'multi_sheet',
                            'sheets_data': sheets_data,
                            'all_sheets_analysis': all_sheets_analysis
                        }
                        return
                
                # Continue with regular single-sheet analysis
                # Analyze data structure to get insights
                data_analysis = analyze_data_structure(df)
                
                # Show data characteristics
                st.subheader("Data Structure Analysis")
                
                # Basic data info
                st.write(f"**Rows:** {data_analysis['num_rows']}, **Columns:** {data_analysis['num_cols']}")
                
                # Column types summary
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Numeric Columns:**")
                    if data_analysis['numeric_columns']:
                        for col in data_analysis['numeric_columns']:
                            st.write(f"- {col}")
                    else:
                        st.write("None found")
                
                with col2:
                    st.write("**Categorical Columns:**")
                    if data_analysis['categorical_columns']:
                        for col in data_analysis['categorical_columns']:
                            st.write(f"- {col}")
                    else:
                        st.write("None found")
                
                # Possible grouping columns
                st.write("**Potential Grouping Variables:**")
                if data_analysis['possible_grouping_columns']:
                    for col in data_analysis['possible_grouping_columns']:
                        st.write(f"- {col}")
                else:
                    st.write("None found")
                
                # Pre/post data detection
                if data_analysis.get('has_pre_post_data', False):
                    st.write("**Pre/Post Data Detected!** Paired analysis is possible.")
                
                # Suggested analyses
                st.subheader("Suggested Analyses")
                
                # Store for automated analyses
                auto_analysis_options = []
                
                if 'suggested_analyses' in data_analysis:
                    for analysis in data_analysis['suggested_analyses']:
                        if analysis == 'correlation_analysis':
                            st.write("- **Correlation Analysis**: Examine relationships between numerical variables")
                            auto_analysis_options.append("Correlation Analysis")
                        
                        elif analysis == 'group_comparison':
                            st.write("- **Group Comparison**: Compare metrics across different groups")
                            auto_analysis_options.append("Group Comparison")
                        
                        elif analysis == 'anova':
                            st.write("- **ANOVA**: Test differences between group means")
                            auto_analysis_options.append("ANOVA")
                        
                        elif analysis == 'paired_t_test':
                            st.write("- **Paired t-test**: Compare before/after measurements")
                            auto_analysis_options.append("Paired t-test")
                        
                        elif analysis == 'regression_analysis':
                            st.write("- **Regression Analysis**: Model relationships between variables")
                            auto_analysis_options.append("Regression Analysis")
                        
                        elif analysis == 'clustering':
                            st.write("- **Clustering**: Find natural groupings in your data")
                            auto_analysis_options.append("Clustering")
                
                # AUTOMATED ML-BASED ANALYSIS:
                # This section configures the ML-powered automated analysis
                # which works independently from the manual drag & drop configuration
                if auto_analysis_options:
                    st.subheader("ML-Based Automated Analysis")
                    
                    # Check if we have manual descriptive analysis configuration
                    manual_config_available = 'descriptive_analysis_config' in st.session_state
                    
                    if manual_config_available:
                        st.success("Using configuration from your Manual Descriptive Analysis to enhance the ML analysis")
                        
                        # Display the configuration being used
                        with st.expander("Manual Analysis Configuration"):
                            config = st.session_state.descriptive_analysis_config
                            
                            if config.get('master_columns'):
                                st.write(f"**Master Columns:** {', '.join(config['master_columns'])}")
                            
                            if config.get('age_column'):
                                st.write(f"**Age Column:** {config['age_column']}")
                            
                            if config.get('gender_column'):
                                st.write(f"**Gender Column:** {config['gender_column']}")
                            
                            if config.get('pre_post_pairs'):
                                st.write("**Pre/Post Pairs:**")
                                for pair in config['pre_post_pairs']:
                                    st.write(f"- {pair['pre']} → {pair['post']}")
                    else:
                        st.info("The following analyses are detected and suggested by our ML system. Complete a Manual Descriptive Analysis first to enhance the ML model with your configurations.")
                    
                    # Use previously selected values if available
                    default_selections = st.session_state.auto_analyses.copy() if st.session_state.auto_analyses else []
                    
                    # If no previous selections or current selections not in options, default to first 3
                    if not default_selections or not any(selection in auto_analysis_options for selection in default_selections):
                        default_selections = auto_analysis_options[:min(3, len(auto_analysis_options))]
                    else:
                        # Filter to only include valid options
                        default_selections = [s for s in default_selections if s in auto_analysis_options]
                    
                    # Log before selection    
                    st.write(f"**Debug - Default Selections:** {default_selections}")
                    
                    # Store the selection in a variable first
                    selected_analyses = st.multiselect(
                        "Select automated analyses to run:",
                        auto_analysis_options,
                        default=default_selections
                    )
                    
                    # Then update the session state
                    st.session_state.auto_analyses = selected_analyses.copy()
                    
                    # Log after selection
                    st.write(f"**Debug - Updated Auto Analyses:** {st.session_state.auto_analyses}")
                    
                    # Show options based on selected analyses
                    if "Correlation Analysis" in st.session_state.auto_analyses:
                        num_cols = data_analysis.get('numeric_columns', [])
                        
                        # Check for manual numeric columns from descriptive analysis
                        manual_numeric_columns = []
                        if 'descriptive_analysis_config' in st.session_state and st.session_state.descriptive_analysis_config.get('numeric_columns'):
                            manual_numeric_columns = st.session_state.descriptive_analysis_config.get('numeric_columns', [])
                            if manual_numeric_columns:
                                st.success(f"Using numeric columns from manual analysis for correlation calculations")
                                
                                # Prioritize manual numeric columns
                                # Get intersection of manual columns and actual dataframe columns
                                manual_cols_in_df = [col for col in manual_numeric_columns if col in df.columns]
                                
                                # Add any manual numeric columns that exist in dataframe to our options
                                for col in manual_cols_in_df:
                                    if col not in num_cols:
                                        num_cols.append(col)
                        
                        if len(num_cols) >= 2:
                            # Initialize corr_columns if not set
                            if 'corr_columns' not in st.session_state:
                                st.session_state.corr_columns = []
                            
                            # Use manual columns as default if available
                            default_corr_columns = []
                            if manual_numeric_columns:
                                # Only use columns that exist in the dataframe
                                default_corr_columns = [col for col in manual_numeric_columns if col in df.columns]
                                if default_corr_columns:
                                    st.info("Using your manually selected numeric columns as defaults for correlation analysis")
                            
                            # If no manual columns, use previous selection if available
                            if not default_corr_columns and st.session_state.corr_columns:
                                default_corr_columns = [col for col in st.session_state.corr_columns if col in num_cols]
                            
                            # If still no default, use all numeric columns (up to 10 to avoid performance issues)
                            if not default_corr_columns:
                                default_corr_columns = num_cols[:min(10, len(num_cols))]
                            
                            # Debug info in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Correlation Columns in Session State:** {st.session_state.corr_columns}")
                                st.write(f"**Default Correlation Columns:** {default_corr_columns}")
                                if manual_numeric_columns:
                                    st.write(f"**Numeric Columns from Manual Analysis:** {manual_numeric_columns}")
                            
                            # Multi-select for correlation columns
                            selected_corr_columns = st.multiselect(
                                "Select columns for correlation analysis:",
                                num_cols,
                                default=default_corr_columns
                            )
                            
                            # Update session state
                            st.session_state.corr_columns = selected_corr_columns.copy()
                            
                            # Method selection
                            corr_method = st.selectbox(
                                "Correlation method:",
                                ["pearson", "spearman", "kendall"],
                                index=0
                            )
                            
                            # Alpha selection
                            corr_alpha = st.slider(
                                "Significance level (α):",
                                0.01, 0.10, 0.05, key="corr_alpha"
                            )
                            
                            # Store correlation analysis settings
                            st.session_state.auto_corr_settings = {
                                'columns': st.session_state.corr_columns,
                                'method': corr_method,
                                'alpha': corr_alpha
                            }
                        else:
                            st.warning("Need at least two numeric columns for correlation analysis.")
                    
                    if "Group Comparison" in st.session_state.auto_analyses:
                        # Check for master columns from manual descriptive analysis
                        manual_master_columns = []
                        if 'descriptive_analysis_config' in st.session_state and st.session_state.descriptive_analysis_config.get('master_columns'):
                            manual_master_columns = st.session_state.descriptive_analysis_config['master_columns']
                            st.success(f"Using master columns from manual analysis: {', '.join(manual_master_columns)}")
                        
                        # Get potential grouping columns
                        group_cols = data_analysis.get('possible_grouping_columns', [])
                        
                        # Prioritize master columns from manual analysis
                        if manual_master_columns:
                            # Add manual master columns that exist in the dataframe
                            for col in manual_master_columns:
                                if col in df.columns and col not in group_cols:
                                    group_cols.insert(0, col)  # Add at the beginning
                            
                            # If we have manual master columns, show a message
                            st.info("Group columns include your manually selected master columns plus auto-detected columns")
                        
                        if group_cols:
                            # Debug info in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Group Column in Session State:** {st.session_state.group_col}")
                                if manual_master_columns:
                                    st.write(f"**Master Columns from Manual Analysis:** {manual_master_columns}")
                            
                            # Use manual master column as default if available
                            default_index = 0
                            if manual_master_columns and manual_master_columns[0] in group_cols:
                                default_index = group_cols.index(manual_master_columns[0])
                            elif st.session_state.group_col in group_cols:
                                default_index = group_cols.index(st.session_state.group_col)
                                
                            # Store the selection in a variable first
                            selected_group = st.selectbox(
                                "Select grouping column for comparison:",
                                group_cols,
                                index=default_index
                            )
                            
                            # Then update the session state
                            st.session_state.group_col = selected_group
                            
                            # Debug after selection in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Updated Group Column:** {st.session_state.group_col}")
                    
                    if "Paired t-test" in st.session_state.auto_analyses:
                        # Try to use pre/post pairs from manual descriptive analysis first
                        pre_post_pairs = []
                        
                        if 'descriptive_analysis_config' in st.session_state and st.session_state.descriptive_analysis_config.get('pre_post_pairs'):
                            # Use the pairs from descriptive analysis
                            manual_pairs = st.session_state.descriptive_analysis_config['pre_post_pairs']
                            for pair in manual_pairs:
                                pre_col = pair['pre']
                                post_col = pair['post']
                                if pre_col in df.columns and post_col in df.columns:
                                    pre_post_pairs.append((pre_col, post_col))
                            
                            st.success(f"Using {len(pre_post_pairs)} pre/post pairs from your manual analysis configuration")
                            
                        # Fall back to the original method if no pairs from descriptive analysis
                        elif 'pair_areas' in st.session_state and st.session_state.pair_areas:
                            st.info("Using pre/post pairs from current drag & drop configuration")
                            for pair_area in st.session_state.pair_areas:
                                pre_cols = [col['name'] for col in pair_area.get('pre', [])]
                                post_cols = [col['name'] for col in pair_area.get('post', [])]
                                
                                # Match each pre column with a post column if possible
                                for i in range(min(len(pre_cols), len(post_cols))):
                                    if pre_cols[i] in df.columns and post_cols[i] in df.columns:
                                        pre_post_pairs.append((pre_cols[i], post_cols[i]))
                        
                        # Get all unique pre and post columns for selection
                        if pre_post_pairs:
                            pre_cols = list(set([pair[0] for pair in pre_post_pairs]))
                            post_cols = list(set([pair[1] for pair in pre_post_pairs]))
                        else:
                            pre_cols = []
                            post_cols = []
                        
                        if pre_cols and post_cols:
                            # Debug info in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Pre Column in Session State:** {st.session_state.pre_col}")
                            
                            # Use previously selected pre column if available
                            default_pre_index = 0
                            if st.session_state.pre_col in pre_cols:
                                default_pre_index = pre_cols.index(st.session_state.pre_col)
                                
                            # Store in variable first
                            selected_pre = st.selectbox(
                                "Select Pre measurement column:", 
                                pre_cols,
                                index=default_pre_index
                            )
                            
                            # Update session state
                            st.session_state.pre_col = selected_pre
                            
                            # Debug after selection in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Updated Pre Column:** {st.session_state.pre_col}")
                            
                            # Find matching post column from the pair configuration
                            matching_posts = []
                            for pre, post in pre_post_pairs:
                                if pre == st.session_state.pre_col:
                                    matching_posts.append(post)
                            
                            if matching_posts:
                                default_post = matching_posts[0]
                            else:
                                default_post = post_cols[0] if post_cols else None
                                
                            # If we previously had a post column selected, try to use it
                            if st.session_state.post_col in post_cols:
                                default_post = st.session_state.post_col
                            
                            if default_post:
                                # Debug info in collapsible area
                                with st.expander("Debug Information", expanded=False):
                                    st.write(f"**Post Column in Session State:** {st.session_state.post_col}")
                                    st.write(f"**Default Post Column:** {default_post}")
                                
                                # Find the index of the default post column
                                post_index = 0
                                if default_post in post_cols:
                                    post_index = post_cols.index(default_post)
                                    
                                # Store in variable first
                                selected_post = st.selectbox(
                                    "Select Post measurement column:", 
                                    post_cols,
                                    index=post_index
                                )
                                
                                # Update session state
                                st.session_state.post_col = selected_post
                                
                                # Debug after selection in collapsible area
                                with st.expander("Debug Information", expanded=False):
                                    st.write(f"**Updated Post Column:** {st.session_state.post_col}")
                    
                    if "Regression Analysis" in st.session_state.auto_analyses:
                        num_cols = data_analysis.get('numeric_columns', [])
                        if len(num_cols) >= 2:
                            # Debug target variable
                            st.write(f"**Debug - Target Column in Session State:** {st.session_state.target_col}")
                            
                            # Set default selection based on previously selected target variable
                            default_index = 0
                            if st.session_state.target_col in num_cols:
                                default_index = num_cols.index(st.session_state.target_col)
                            
                            # Store in variable first    
                            selected_target = st.selectbox(
                                "Select target variable (dependent variable):",
                                num_cols,
                                index=default_index
                            )
                            
                            # Update session state
                            st.session_state.target_col = selected_target
                            
                            # Debug after selection
                            st.write(f"**Debug - Updated Target Column:** {st.session_state.target_col}")
                            
                            # Filter out the target column for predictors
                            predictor_options = [col for col in num_cols if col != st.session_state.target_col]
                            
                            # Debug predictor variables
                            st.write(f"**Debug - Predictor Columns in Session State:** {st.session_state.predictor_cols}")
                            
                            # Use previously selected predictors if available
                            default_predictors = [p for p in st.session_state.predictor_cols if p in predictor_options]
                            
                            # If no previous selections, default to first 3
                            if not default_predictors:
                                default_predictors = predictor_options[:min(3, len(predictor_options))]
                            
                            st.write(f"**Debug - Default Predictor Selections:** {default_predictors}")
                            
                            # Store in variable first
                            selected_predictors = st.multiselect(
                                "Select predictor variables:",
                                predictor_options,
                                default=default_predictors
                            )
                            
                            # Update session state
                            st.session_state.predictor_cols = selected_predictors.copy()
                            
                            # Debug after selection
                            st.write(f"**Debug - Updated Predictor Columns:** {st.session_state.predictor_cols}")
                else:
                    st.warning("No automatic analyses could be suggested based on your data structure.")
            
            # Button to proceed to data tables
            if st.button("Continue to Data Tables"):
                st.session_state.current_step = 4
                # Add JavaScript to scroll to top after rerun
                st.markdown(
                    """
                    <script>
                        window.scrollTo(0, 0);
                    </script>
                    """,
                    unsafe_allow_html=True
                )
                st.rerun()
    
    elif st.session_state.current_step == 4:
        # Step 4: Tables of Data
        st.header(emoji_header("4. Tables of Data", "Tables"))
        
        if st.session_state.data is not None and st.session_state.selected_columns:
            df = st.session_state.data
            
            # Generate difference columns based on manually selected pre/post pairs
            # Check if user has manually configured pre/post pairs
            if 'pair_areas' in st.session_state and st.session_state.pair_areas:
                manual_pairs = []
                for pair_area in st.session_state.pair_areas:
                    pre_cols = [col['name'] for col in pair_area.get('pre', [])]
                    post_cols = [col['name'] for col in pair_area.get('post', [])]
                    
                    # Match each pre column with a post column
                    for i in range(min(len(pre_cols), len(post_cols))):
                        if pre_cols[i] in df.columns and post_cols[i] in df.columns:
                            manual_pairs.append((pre_cols[i], post_cols[i]))
                
                if manual_pairs:
                    # Calculate difference columns for manually selected pairs
                    df, difference_info = calculate_manual_paired_differences(df, manual_pairs)
                    
                    if difference_info['summary']['total_difference_columns_added'] > 0:
                        st.info(f"Generated {difference_info['summary']['total_difference_columns_added']} difference columns from {len(manual_pairs)} manually selected pre/post pairs.")
                        
                        # Update selected columns to include difference columns
                        for diff_col in difference_info['difference_columns']:
                            if diff_col not in st.session_state.selected_columns:
                                st.session_state.selected_columns.append(diff_col)
                        
                        # Update the session state data with difference columns
                        st.session_state.data = df
            
            # Show results based on selected analysis type
            if st.session_state.analysis_type == "Descriptive Statistics":
                st.subheader("Descriptive Statistics")
                
                try:
                    # Generate and display descriptive statistics
                    # Get stats for all columns regardless of what was selected for display
                    stats_df = get_descriptive_stats(df)
                    
                    # Convert 'Mean ± Std Dev' to a string column to avoid Arrow conversion errors
                    for col in stats_df.columns:
                        if isinstance(stats_df.loc['Mean ± Std Dev', col], str):
                            # If Mean ± Std Dev is already a string, we need to convert the whole column to strings
                            # for proper display in the dataframe
                            stats_df[col] = stats_df[col].astype(str)
                    
                    # Display the raw data table first
                    st.subheader("Raw Data Table")
                    st.dataframe(df, height=600)
                    
                    st.subheader("Overall Descriptive Statistics")
                    st.dataframe(stats_df)
                    
                    # Master parameter analysis - use configured master parameters from drag & drop UI
                    st.subheader("Master Parameter Analysis")
                    
                    master_columns = []
                    # Debug info moved to a collapsible area for cleaner UI
                    with st.expander("Debug Information (click to expand)", expanded=False):
                        st.write("**Master Configuration:**")
                        if 'master_columns' in st.session_state:
                            st.write(f"Master columns in session state: {st.session_state.master_columns}")
                            
                            # List out each master column for debugging
                            for i, col in enumerate(st.session_state.master_columns):
                                st.write(f"Master column {i+1}: {col}")
                                
                                # Check if column exists in dataframe
                                if col['name'] in df.columns:
                                    st.write(f"✓ Column '{col['name']}' found in dataframe")
                                else:
                                    st.write(f"❌ Column '{col['name']}' NOT found in dataframe")
                        else:
                            st.write("No master columns in session state")
                        
                        if 'age_column' in st.session_state and st.session_state.age_column:
                            st.write(f"Age column: {st.session_state.age_column}")
                            if 'generated_age_group_col' in st.session_state:
                                st.write(f"Generated age group column: {st.session_state.generated_age_group_col}")
                            else:
                                st.write("No generated age group column in session state")
                        else:
                            st.write("No age column selected")
                        
                    # Use the master columns set in the drag & drop UI if available
                    if 'master_columns' in st.session_state and st.session_state.master_columns:
                        # Ensure we're handling all columns, even if they're complex dictionary objects
                        for col in st.session_state.master_columns:
                            # Handle both string and dictionary column formats
                            col_name = col['name'] if isinstance(col, dict) else col
                            
                            if col_name in df.columns:
                                master_columns.append(col_name)
                            else:
                                st.warning(f"Master column '{col_name}' was configured but not found in the dataset.")
                        
                        # Explicitly add the generated age group column if it exists
                        if 'age_column' in st.session_state and st.session_state.age_column:
                            if 'generated_age_group_col' in st.session_state and st.session_state.generated_age_group_col in df.columns:
                                age_group_col = st.session_state.generated_age_group_col
                                if age_group_col not in master_columns:
                                    master_columns.append(age_group_col)
                                    st.info(f"Added generated Age Group column '{age_group_col}' to master parameters")
                    
                    # If no master columns were set in the UI, use System, Gender, Age Group if they exist
                    if not master_columns:
                        # Check if 'System' column exists
                        if 'System' in df.columns:
                            master_columns.append('System')
                        
                        # Only include Gender if it was explicitly defined by the user
                        # Gender auto-detection disabled to respect user configuration
                        
                        # Check if age groups were generated - first check for session state variable
                        if 'generated_age_group_col' in st.session_state and st.session_state.generated_age_group_col in df.columns:
                            age_group_col = st.session_state.generated_age_group_col
                            master_columns.append(age_group_col)
                            st.info(f"Added generated Age Group column '{age_group_col}' to master parameters")
                        # Fallback to standard name
                        elif 'Generated Age Group' in df.columns:
                            master_columns.append('Generated Age Group')
                        
                        # Also check for any column containing "Age Group"
                        age_group_cols = [col for col in df.columns if "Age Group" in col and col not in master_columns]
                        for age_col in age_group_cols:
                            if age_col not in master_columns:
                                master_columns.append(age_col)
                    
                    # Display final master columns in debug expander
                    with st.expander("Debug Information (click to expand)", expanded=False):
                        st.write("**Final Master Columns:**")
                        st.write(f"Master columns for analysis: {master_columns}")
                    
                    if master_columns:
                        st.write(f"Analysis by master parameters: {', '.join(master_columns)}")
                        
                        # For each master parameter, perform analysis - Each one in its own expander
                        for master_param in master_columns:
                            with st.expander(f"Analysis by {master_param}", expanded=False):
                                # Call analyze_master_parameters for this master parameter
                                master_results = analyze_master_parameters(df, [master_param])
                                
                                if not master_results or master_param not in master_results:
                                    st.error(f"Error analyzing master parameter {master_param}.")
                                else:
                                    results = master_results[master_param]
                                    col_type = results.get('type', 'unknown')
                                    st.write(f"Parameter type: {col_type}")
                                    
                                    # Check if there's only one group
                                    groups = results.get('groups', {})
                                    if len(groups) == 1:
                                        group_name = list(groups.keys())[0]
                                        st.info(f"⚠️ Note: The dataset only contains one value for {master_param} ({group_name}).")
                                    
                                    # Sort the groups for age groups to ensure they appear in order from small to large
                                    if 'age group' in master_param.lower() or 'Generated Age Group' in master_param:
                                        # Sort the groups by extracting the start age from each group name
                                        try:
                                            # Extract age numbers from strings like "0-10" or "10-20"
                                            def extract_start_age(group_name):
                                                if isinstance(group_name, str) and '-' in group_name:
                                                    return int(group_name.split('-')[0])
                                                return 0  # Default for sorting non-standard names
                                            
                                            # Sort by the starting age value
                                            sorted_groups = sorted(groups.items(), key=lambda x: extract_start_age(x[0]))
                                            groups = dict(sorted_groups)
                                        except Exception as e:
                                            st.error(f"Error sorting age groups: {e}")
                                    
                                    # Display results for each group
                                    for group_name, group_data in groups.items():
                                        st.write(f"**{master_param} = {group_name} (n={group_data['count']})**")
                                        group_stats = group_data.get('stats')
                                        if group_stats is not None:
                                            # Convert 'Mean ± Std Dev' to string to prevent Arrow conversion errors
                                            for col in group_stats.columns:
                                                if 'Mean ± Std Dev' in group_stats.index and isinstance(group_stats.loc['Mean ± Std Dev', col], str):
                                                    group_stats[col] = group_stats[col].astype(str)
                                            
                                            # Display statistics
                                            st.dataframe(group_stats)
                                        else:
                                            st.write("No statistics available for this group.")
                    
                    # Pre/Post Analysis - use only manually assigned pairs
                    st.subheader("Pre/Post Comparisons")
                    
                    # MANUAL DESCRIPTIVE ANALYSIS:
                    # Use manually configured pairs from the drag & drop UI
                    # This is different from the automated ML analysis which detects pairs automatically
                    pre_post_pairs = []
                    if 'pair_areas' in st.session_state and st.session_state.pair_areas:
                        for pair_area in st.session_state.pair_areas:
                            pre_cols = [col['name'] for col in pair_area.get('pre', [])]
                            post_cols = [col['name'] for col in pair_area.get('post', [])]
                            
                            # Match each pre column with a post column if possible
                            for i in range(min(len(pre_cols), len(post_cols))):
                                if pre_cols[i] in df.columns and post_cols[i] in df.columns:
                                    pre_post_pairs.append((pre_cols[i], post_cols[i]))
                    
                    if pre_post_pairs:
                        st.write(f"Analyzing {len(pre_post_pairs)} manually assigned pre/post pairs...")
                        
                        # Create a dataframe to store the results
                        pre_post_results = []
                        
                        for pre_col, post_col in pre_post_pairs:
                            if pre_col in df.columns and post_col in df.columns:
                                # Ensure columns are numeric
                                if pd.api.types.is_numeric_dtype(df[pre_col]) and pd.api.types.is_numeric_dtype(df[post_col]):
                                    # Calculate difference
                                    diff = df[pre_col] - df[post_col]
                                    
                                    # Check for zero division issues
                                    has_zero = (df[pre_col] == 0).any()
                                    if has_zero:
                                        # Use a safe percentage change calculation
                                        pct_change = pd.Series([
                                            ((a - b) / a * 100) if a != 0 else float('nan')
                                            for a, b in zip(df[pre_col], df[post_col])
                                        ])
                                    else:
                                        pct_change = (diff / df[pre_col]) * 100
                                    
                                    # Calculate statistics
                                    mean_diff = diff.mean()
                                    std_diff = diff.std()
                                    min_diff = diff.min()
                                    max_diff = diff.max()
                                    
                                    # Add to results
                                    pre_post_results.append({
                                        'Pre Column': pre_col,
                                        'Post Column': post_col,
                                        'Mean Diff': round(mean_diff, 3),
                                        'Std Dev': round(std_diff, 3),
                                        'Min Diff': round(min_diff, 3),
                                        'Max Diff': round(max_diff, 3),
                                        'Mean % Change': round(pct_change.mean(), 2) if not pct_change.empty else 'NA'
                                    })
                        
                        if pre_post_results:
                            # Show pre/post analysis results
                            st.write("Pre/Post Difference Analysis:")
                            st.dataframe(pd.DataFrame(pre_post_results))
                            
                            # Show pre/post differences by master parameter if available
                            if master_columns:
                                st.write("Pre/Post Differences by Master Parameter:")
                                
                                for master_param in master_columns:
                                    with st.expander(f"Pre/Post Analysis by {master_param}"):
                                        # Group by master parameter
                                        for group_name, group_df in df.groupby(master_param):
                                            st.write(f"**{master_param} = {group_name} (n={len(group_df)})**")
                                            
                                            # Create a dataframe for this group's pre/post results
                                            group_results = []
                                            
                                            for pre_col, post_col in pre_post_pairs:
                                                if pre_col in group_df.columns and post_col in group_df.columns:
                                                    # Ensure columns are numeric
                                                    if pd.api.types.is_numeric_dtype(group_df[pre_col]) and pd.api.types.is_numeric_dtype(group_df[post_col]):
                                                        # Calculate difference
                                                        diff = group_df[pre_col] - group_df[post_col]
                                                        
                                                        # Check for zero division issues
                                                        has_zero = (group_df[pre_col] == 0).any()
                                                        if has_zero:
                                                            # Use a safe percentage change calculation
                                                            pct_change = pd.Series([
                                                                ((a - b) / a * 100) if a != 0 else float('nan')
                                                                for a, b in zip(group_df[pre_col], group_df[post_col])
                                                            ])
                                                        else:
                                                            pct_change = (diff / group_df[pre_col]) * 100
                                                        
                                                        # Add to results
                                                        group_results.append({
                                                            'Pre Column': pre_col,
                                                            'Post Column': post_col,
                                                            'Mean Diff': round(diff.mean(), 3),
                                                            'Std Dev': round(diff.std(), 3),
                                                            'Min Diff': round(diff.min(), 3),
                                                            'Max Diff': round(diff.max(), 3),
                                                            'Mean % Change': round(pct_change.mean(), 2) if not pct_change.empty else 'NA'
                                                        })
                                            
                                            if group_results:
                                                st.dataframe(pd.DataFrame(group_results))
                                            else:
                                                st.write("No pre/post results available for this group.")
                    else:
                        st.info("No pre/post pairs have been configured. Please use the Excel Drag & Drop configuration to assign pre/post pairs.")
                    
                    # Store results for later
                    st.session_state.analysis_results = {
                        'type': 'descriptive',
                        'stats_df': stats_df,
                        'master_columns': master_columns,
                        'age_column': age_column if 'age_column' in locals() else None,
                        'age_groups': age_groups if 'age_groups' in locals() else None,
                        'pre_post_pairs': pre_post_pairs if 'pre_post_pairs' in locals() else []
                    }
                    
                    # Store descriptive analysis configuration for ML model training
                    # Get all selected columns
                    selected_cols = st.session_state.selected_columns
                    
                    # Identify numeric columns for correlation analysis
                    numeric_columns = []
                    for col in selected_cols:
                        try:
                            if pd.api.types.is_numeric_dtype(df[col]):
                                numeric_columns.append(col)
                        except:
                            pass
                    
                    # Store all configuration in session state
                    st.session_state.descriptive_analysis_config = {
                        'master_columns': master_columns,
                        'age_column': age_column if 'age_column' in locals() else None,
                        'gender_column': gender_column if 'gender_column' in locals() else None,
                        'pre_post_pairs': pre_post_pairs if 'pre_post_pairs' in locals() else [],
                        'all_selected_columns': selected_cols,
                        'numeric_columns': numeric_columns,
                        'analyzed_dataframe': df
                    }
                    
                    # Display additional statistics for all columns
                    # Get all selected columns
                    all_selected_cols = st.session_state.selected_columns
                    
                    # Check if we have master columns from configuration
                    master_columns = []
                    if 'configuration_complete' in st.session_state and st.session_state.configuration_complete:
                        if 'master_columns' in st.session_state:
                            # Get column names from the column objects
                            master_columns = [col['name'] for col in st.session_state.master_columns]
                            
                    # Skip this section as we already performed Master Parameter Analysis in the Tables section
                    # The analysis results from the Tables section can be reused if needed
                    # This avoids duplicating the same analysis and confusing the user
                    
                    # Add a note about where to find the full master parameter analysis
                    if master_columns:
                        st.info("⚠️ Note: Master Parameter Analysis has already been performed in the Tables section above. Please refer to that section for detailed results.")
                    
                    # Calculate statistics for all columns
                    st.subheader("Additional Statistics for All Columns")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("Skewness:")
                        # Create DataFrame with all columns
                        skew_data = []
                        for col in all_selected_cols:
                            try:
                                # Try to calculate skewness - will work for numerical and some other types
                                skew_value = df[col].skew()
                                skew_data.append({'Column': col, 'Skewness': skew_value})
                            except:
                                # For columns where skewness can't be calculated
                                skew_data.append({'Column': col, 'Skewness': float('nan')})
                        
                        skew_df = pd.DataFrame(skew_data)
                        # Convert all values to strings to prevent Arrow conversion errors
                        skew_df['Skewness'] = skew_df['Skewness'].apply(lambda x: "NA" if pd.isna(x) else f"{x:.4f}")
                        st.dataframe(skew_df)
                    
                    with col2:
                        st.write("Kurtosis:")
                        # Create DataFrame with all columns
                        kurt_data = []
                        for col in all_selected_cols:
                            try:
                                # Try to calculate kurtosis - will work for numerical and some other types
                                kurt_value = df[col].kurtosis()
                                kurt_data.append({'Column': col, 'Kurtosis': kurt_value})
                            except:
                                # For columns where kurtosis can't be calculated
                                kurt_data.append({'Column': col, 'Kurtosis': float('nan')})
                        
                        kurt_df = pd.DataFrame(kurt_data)
                        # Convert all values to strings to prevent Arrow conversion errors
                        kurt_df['Kurtosis'] = kurt_df['Kurtosis'].apply(lambda x: "NA" if pd.isna(x) else f"{x:.4f}")
                        st.dataframe(kurt_df)
                    
                    # We've already handled master parameter analysis above
                
                except Exception as e:
                    st.error(f"Error calculating statistics: {str(e)}")
            
            elif st.session_state.analysis_type == "Statistical Tests":
                if 'ttest_settings' in st.session_state and st.session_state.ttest_settings:
                    st.subheader("t-Test Results")
                    settings = st.session_state.ttest_settings
                    
                    try:
                        if settings.get('test_type') == "One-sample t-test":
                            result = perform_ttest(
                                df, 
                                settings['column'], 
                                test_type="one-sample", 
                                mu=settings['mu'], 
                                alpha=settings['alpha']
                            )
                            
                            # Display results
                            st.write(f"**Testing if the mean of {settings['column']} differs from {settings['mu']}**")
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.write(f"Sample mean: {result['sample_mean']:.4f}")
                                st.write(f"t-statistic: {result['t_stat']:.4f}")
                                st.write(f"p-value: {result['p_value']:.4f}")
                            
                            with col2:
                                st.write(f"Degrees of freedom: {result['df']}")
                                st.write(f"95% Confidence Interval: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                            
                            # Interpretation
                            st.subheader("Interpretation")
                            if result['p_value'] < settings['alpha']:
                                st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {settings['alpha']}, we **reject** the null hypothesis.")
                                st.write(f"There is significant evidence that the true mean of {settings['column']} differs from {settings['mu']}.")
                            else:
                                st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {settings['alpha']}, we **fail to reject** the null hypothesis.")
                                st.write(f"There is insufficient evidence that the true mean of {settings['column']} differs from {settings['mu']}.")
                            
                            # Store results for later
                            st.session_state.analysis_results = {
                                'type': 'ttest',
                                'test_type': 'one-sample',
                                'result': result
                            }
                        
                        elif settings.get('test_type') == "Two-sample t-test":
                            if settings.get('method') == "Compare two columns":
                                result = perform_ttest(
                                    df, 
                                    settings['col1'], 
                                    test_type="two-sample-cols", 
                                    col2=settings['col2'], 
                                    equal_var=settings['equal_var'], 
                                    alpha=settings['alpha']
                                )
                                
                                # Display results
                                st.write(f"**Comparing means of {settings['col1']} and {settings['col2']}**")
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.write(f"Mean of {settings['col1']}: {result['mean1']:.4f}")
                                    st.write(f"Mean of {settings['col2']}: {result['mean2']:.4f}")
                                    st.write(f"Mean difference: {result['mean_diff']:.4f}")
                                
                                with col2:
                                    st.write(f"t-statistic: {result['t_stat']:.4f}")
                                    st.write(f"p-value: {result['p_value']:.4f}")
                                    st.write(f"Degrees of freedom: {result['df']:.1f}")
                                
                                st.write(f"95% Confidence Interval for difference: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                                
                                # Interpretation
                                st.subheader("Interpretation")
                                test_name = "Welch's t-test" if not settings['equal_var'] else "Student's t-test"
                                st.write(f"Using {test_name} ({'unequal' if not settings['equal_var'] else 'equal'} variances):")
                                
                                if result['p_value'] < settings['alpha']:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {settings['alpha']}, we **reject** the null hypothesis.")
                                    st.write(f"There is significant evidence that the means of {settings['col1']} and {settings['col2']} differ.")
                                else:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {settings['alpha']}, we **fail to reject** the null hypothesis.")
                                    st.write(f"There is insufficient evidence that the means of {settings['col1']} and {settings['col2']} differ.")
                                
                                # Store results for later
                                st.session_state.analysis_results = {
                                    'type': 'ttest',
                                    'test_type': 'two-sample-cols',
                                    'result': result
                                }
                            
                            elif settings.get('method') == "Compare groups within a column":
                                result = perform_ttest(
                                    df, 
                                    settings['num_col'], 
                                    test_type="two-sample-groups", 
                                    group_col=settings['group_col'],
                                    group1=settings['group1'],
                                    group2=settings['group2'],
                                    equal_var=settings['equal_var'], 
                                    alpha=settings['alpha']
                                )
                                
                                # Display results
                                st.write(f"**Comparing {settings['num_col']} between {settings['group1']} and {settings['group2']} groups**")
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.write(f"Mean for {settings['group1']}: {result['mean1']:.4f}")
                                    st.write(f"Mean for {settings['group2']}: {result['mean2']:.4f}")
                                    st.write(f"Mean difference: {result['mean_diff']:.4f}")
                                
                                with col2:
                                    st.write(f"t-statistic: {result['t_stat']:.4f}")
                                    st.write(f"p-value: {result['p_value']:.4f}")
                                    st.write(f"Degrees of freedom: {result['df']:.1f}")
                                
                                st.write(f"95% Confidence Interval for difference: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                                
                                # Interpretation
                                st.subheader("Interpretation")
                                test_name = "Welch's t-test" if not settings['equal_var'] else "Student's t-test"
                                st.write(f"Using {test_name} ({'unequal' if not settings['equal_var'] else 'equal'} variances):")
                                
                                if result['p_value'] < settings['alpha']:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {settings['alpha']}, we **reject** the null hypothesis.")
                                    st.write(f"There is significant evidence that the mean {settings['num_col']} differs between {settings['group1']} and {settings['group2']} groups.")
                                else:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {settings['alpha']}, we **fail to reject** the null hypothesis.")
                                    st.write(f"There is insufficient evidence that the mean {settings['num_col']} differs between {settings['group1']} and {settings['group2']} groups.")
                                
                                # Store results for later
                                st.session_state.analysis_results = {
                                    'type': 'ttest',
                                    'test_type': 'two-sample-groups',
                                    'result': result
                                }
                    
                    except Exception as e:
                        st.error(f"Error performing t-test: {str(e)}")
                
                elif 'anova_settings' in st.session_state and st.session_state.anova_settings:
                    st.subheader("ANOVA Results")
                    settings = st.session_state.anova_settings
                    
                    try:
                        if 'num_col' in settings and 'cat_col' in settings:
                            result = perform_anova(
                                df, 
                                settings['num_col'], 
                                settings['cat_col'], 
                                alpha=settings['alpha']
                            )
                            
                            # Display ANOVA table
                            st.write(f"**One-way ANOVA: Effect of {settings['cat_col']} on {settings['num_col']}**")
                            st.dataframe(result['anova_table'])
                            
                            # Display group statistics
                            st.subheader("Group Statistics")
                            st.dataframe(result['group_stats'])
                            
                            # Interpretation
                            st.subheader("Interpretation")
                            if result['p_value'] < settings['alpha']:
                                st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {settings['alpha']}, we **reject** the null hypothesis.")
                                st.write(f"There is significant evidence that the mean of {settings['num_col']} differs across at least one of the {settings['cat_col']} groups.")
                            else:
                                st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {settings['alpha']}, we **fail to reject** the null hypothesis.")
                                st.write(f"There is insufficient evidence that the mean of {settings['num_col']} differs across the {settings['cat_col']} groups.")
                            
                            # Store results for later
                            st.session_state.analysis_results = {
                                'type': 'anova',
                                'result': result
                            }
                    
                    except Exception as e:
                        st.error(f"Error performing ANOVA: {str(e)}")
                
                elif 'chi2_settings' in st.session_state and st.session_state.chi2_settings:
                    st.subheader("Chi-Square Test Results")
                    settings = st.session_state.chi2_settings
                    
                    try:
                        # Run Chi-Square test based on stored settings
                        result = perform_chi_square_test(
                            df,
                            settings['col1'],
                            settings['col2'],
                            alpha=settings['alpha']
                        )
                        
                        # Display Chi-Square results
                        st.write(f"### Chi-Square Test of Independence")
                        st.write(f"**Categorical Variables:** {settings['col1']} × {settings['col2']}")
                        st.write(f"**Null Hypothesis (H₀):** The two categorical variables are independent (no association)")
                        st.write(f"**Alternative Hypothesis (H₁):** The two categorical variables are not independent (there is an association)")
                        
                        # Display warning about expected frequencies if applicable
                        if result['warning']:
                            st.warning(result['warning'])
                        
                        # Display contingency table
                        st.write("**Contingency Table (Observed Frequencies):**")
                        st.dataframe(result['formatted_table'])
                        
                        # Display Chi-Square test statistics
                        st.write(f"**Results:**")
                        st.write(f"- Chi-square statistic: {result['chi2']:.4f}")
                        st.write(f"- Degrees of Freedom: {result['dof']}")
                        st.write(f"- p-value: {result['p_value']:.4f}")
                        st.write(f"- Cramer's V (effect size): {result['cramers_v']:.4f} ({result['effect_size']} effect)")
                        st.write(f"- Sample size: {result['n']}")
                        
                        # Visualization of observed vs expected
                        st.write("**Visualization:**")
                        
                        # Create a heatmap of the contingency table
                        import plotly.graph_objects as go
                        import plotly.express as px
                        
                        # Get contingency table
                        cont_table = result['contingency_table']
                        
                        # Create heatmap
                        fig = px.imshow(
                            cont_table.values,
                            labels=dict(x=settings['col2'], y=settings['col1'], color="Frequency"),
                            x=cont_table.columns,
                            y=cont_table.index,
                            text_auto=True,
                            aspect="auto",
                            color_continuous_scale="Blues"
                        )
                        fig.update_layout(
                            title=f"Contingency Table Heatmap: {settings['col1']} × {settings['col2']}",
                        )
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Statistical significance
                        if result['p_value'] < settings['alpha']:
                            st.success(f"**Finding:** We reject the null hypothesis (p < {settings['alpha']}). There is a significant association between {settings['col1']} and {settings['col2']}.")
                            
                            # Describe strength of association
                            if result['effect_size'] == "Negligible":
                                st.info(f"The association is statistically significant but negligible in strength (Cramer's V = {result['cramers_v']:.4f}).")
                            elif result['effect_size'] == "Weak":
                                st.info(f"The association is statistically significant with weak strength (Cramer's V = {result['cramers_v']:.4f}).")
                            elif result['effect_size'] == "Moderate":
                                st.info(f"The association is statistically significant with moderate strength (Cramer's V = {result['cramers_v']:.4f}).")
                            else:  # Strong
                                st.info(f"The association is statistically significant with strong strength (Cramer's V = {result['cramers_v']:.4f}).")
                                
                        else:
                            st.info(f"**Finding:** We fail to reject the null hypothesis (p > {settings['alpha']}). There is not sufficient evidence of an association between {settings['col1']} and {settings['col2']}.")
                        
                    except Exception as e:
                        st.error(f"Error performing Chi-Square test: {str(e)}")
            
            elif st.session_state.analysis_type == "Correlation Analysis":
                if 'corr_settings' in st.session_state and st.session_state.corr_settings:
                    st.subheader("Correlation Analysis Results")
                    settings = st.session_state.corr_settings
                    
                    try:
                        if 'columns' in settings and len(settings['columns']) >= 2:
                            result = perform_correlation_analysis(
                                df, 
                                settings['columns'], 
                                method=settings['method'], 
                                alpha=settings['alpha']
                            )
                            
                            # Display correlation matrix
                            st.subheader(f"{settings['method'].capitalize()} Correlation Matrix")
                            st.dataframe(result['corr_matrix'])
                            
                            # Display significant correlations
                            st.subheader("Significant Correlations")
                            if not result['significant_corrs'].empty:
                                st.dataframe(result['significant_corrs'])
                            else:
                                st.write("No significant correlations found at the specified significance level.")
                            
                            # Store results for later
                            st.session_state.analysis_results = {
                                'type': 'correlation',
                                'result': result
                            }
                    
                    except Exception as e:
                        st.error(f"Error performing correlation analysis: {str(e)}")
            
            # P-Value Analysis Results
            if 'pvalue_settings' in st.session_state and st.session_state.pvalue_settings:
                st.header("P-Value Analysis Results")
                settings = st.session_state.pvalue_settings
                
                try:
                    if 'columns' in settings and len(settings['columns']) >= 1:
                        # Perform overall p-value analysis
                        st.subheader("Overall P-Values for Each Parameter")
                        overall_results = perform_pvalue_analysis(df, settings)
                        
                        if 'overall_pvalues' in overall_results:
                            st.dataframe(overall_results['overall_pvalues'])
                            
                            # Create chart for overall p-values
                            if len(overall_results['overall_pvalues']) > 0:
                                chart_id = f"pvalue_overall_{int(time.time())}"
                                fig_overall = create_pvalue_chart(overall_results['overall_pvalues'], "Overall P-Values by Parameter")
                                
                                col1, col2 = st.columns([5, 1])
                                with col1:
                                    st.plotly_chart(fig_overall, use_container_width=True)
                                with col2:
                                    # Heart button for favoriting
                                    is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                    heart_color = "❤️" if is_favorited else "🤍"
                                    if st.button(
                                        heart_color, 
                                        key=f"heart_overall_{chart_id}",
                                        help="Add to favorites"
                                    ):
                                        toggle_chart_favorite(chart_id, fig_overall, "Overall P-Values by Parameter", "P-Value Chart")
                                        st.rerun()
                        
                        # Perform system-wise p-value analysis if master columns are available
                        if settings.get('master_columns') and len(settings['master_columns']) > 0:
                            st.subheader("System-Wise P-Values")
                            
                            if 'system_pvalues' in overall_results:
                                for master_col, master_data in overall_results['system_pvalues'].items():
                                    st.write(f"**Analysis by {master_col}:**")
                                    
                                    for system, system_data in master_data.items():
                                        st.write(f"**{master_col} = {system}:**")
                                        st.dataframe(system_data)
                                        
                                        # Create chart for each system
                                        if len(system_data) > 0:
                                            chart_id = f"pvalue_system_{master_col}_{system}_{int(time.time())}"
                                            fig_system = create_pvalue_chart(system_data, f"P-Values for {master_col} = {system}")
                                            
                                            col1, col2 = st.columns([5, 1])
                                            with col1:
                                                st.plotly_chart(fig_system, use_container_width=True)
                                            with col2:
                                                # Heart button for favoriting
                                                is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                                heart_color = "❤️" if is_favorited else "🤍"
                                                if st.button(
                                                    heart_color, 
                                                    key=f"heart_system_{chart_id}",
                                                    help="Add to favorites"
                                                ):
                                                    toggle_chart_favorite(chart_id, fig_system, f"P-Values for {master_col} = {system}", "P-Value Chart")
                                                    st.rerun()
                        
                        # Store results for later
                        st.session_state.analysis_results = {
                            'type': 'pvalue',
                            'result': overall_results
                        }
                
                except Exception as e:
                    st.error(f"Error performing p-value analysis: {str(e)}")
            
            # ML Automated Analysis Results
            if 'auto_analyses' in st.session_state and st.session_state.auto_analyses:
                st.header("ML-Based Automated Analysis Results")
                
                # Display Correlation Analysis Results if configured
                if "Correlation Analysis" in st.session_state.auto_analyses and 'auto_corr_settings' in st.session_state:
                    st.subheader("Automated Correlation Analysis")
                    
                    auto_corr_settings = st.session_state.auto_corr_settings
                    
                    try:
                        if 'columns' in auto_corr_settings and len(auto_corr_settings['columns']) >= 2:
                            # Show debug info in collapsible area
                            with st.expander("Debug Information", expanded=False):
                                st.write(f"**Correlation Settings:** {auto_corr_settings}")
                                    
                            # Prepare dataframe for correlation analysis
                            # Ensure we're working with numeric values only
                            corr_df = df.copy()
                            
                            # Check for columns with "Mean ± Std Dev" values and convert them
                            for col in auto_corr_settings['columns']:
                                if col in corr_df.columns:
                                    # If it's not numeric, try to convert it
                                    if not pd.api.types.is_numeric_dtype(corr_df[col]):
                                        try:
                                            # Try to convert to numeric, ignoring errors
                                            corr_df[col] = pd.to_numeric(corr_df[col], errors='coerce')
                                        except Exception as e:
                                            st.warning(f"Column {col} could not be converted to numeric for correlation analysis: {str(e)}")
                            
                            # Perform correlation analysis
                            corr_result = perform_correlation_analysis(
                                corr_df, 
                                auto_corr_settings['columns'], 
                                method=auto_corr_settings['method'], 
                                alpha=auto_corr_settings['alpha']
                            )
                            
                            # Display correlation matrix
                            st.subheader(f"{auto_corr_settings['method'].capitalize()} Correlation Matrix")
                            st.dataframe(corr_result['corr_matrix'])
                            
                            # Display heatmap
                            st.subheader("Correlation Heatmap")
                            fig = create_correlation_heatmap(
                                corr_df, 
                                auto_corr_settings['columns'], 
                                method=auto_corr_settings['method']
                            )
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Display significant correlations
                            st.subheader("Significant Correlations")
                            if not corr_result['significant_corrs'].empty:
                                st.dataframe(corr_result['significant_corrs'])
                                
                                # Show explanation
                                st.info(f"The table above shows correlations with p-values less than {auto_corr_settings['alpha']}, indicating statistically significant relationships between variables.")
                            else:
                                st.info(f"No significant correlations found at the significance level α = {auto_corr_settings['alpha']}.")
                            
                            # Store results for later use
                            st.session_state.auto_correlation_results = corr_result
                        else:
                            st.warning("Need at least two columns for correlation analysis. Please configure the analysis with more columns.")
                    
                    except Exception as e:
                        st.error(f"Error performing automated correlation analysis: {str(e)}")
                
                # Add horizontal line to separate from manual analysis results
                st.markdown("---")
            
            # Button to proceed to visualizations
            if st.button("Continue to Visualizations"):
                st.session_state.current_step = 5
                # Add JavaScript to scroll to top after rerun
                st.markdown(
                    """
                    <script>
                        window.scrollTo(0, 0);
                    </script>
                    """,
                    unsafe_allow_html=True
                )
                st.rerun()
    
    elif st.session_state.current_step == 5:
        # Step 5: Graph/Chart Preparation
        st.markdown('<div id="visualizations"></div>', unsafe_allow_html=True)
        st.header(emoji_header("5. Graph/Chart Preparation", "Charts"))
        
        if st.session_state.data is not None and st.session_state.selected_columns:
            df = st.session_state.data
            
            # Visualization type selector
            st.subheader("Select Visualization Type")
            
            # Check if we have master columns from configuration
            master_columns = []
            if 'configuration_complete' in st.session_state and st.session_state.configuration_complete:
                if 'master_columns' in st.session_state:
                    # Get column names from the column objects
                    master_columns = [col['name'] for col in st.session_state.master_columns]
            
            viz_options = [
                "Histogram", 
                "Box Plot", 
                "Scatter Plot", 
                "Correlation Heatmap", 
                "Distribution Plot",
                "Bar Chart",
                "Line Chart",
                "Frequency Analysis"  # Added this new option
            ]
            
            # Add Master Analysis as an option if we have master columns
            if master_columns:
                viz_options.insert(0, "Master Parameter Analysis")
            
            st.session_state.visualization_type = st.selectbox(
                "Choose the type of visualization you want to create:",
                viz_options,
                index=viz_options.index(st.session_state.visualization_type) if st.session_state.visualization_type in viz_options else 0
            )
            
            # Create appropriate visualization based on user selection
            if st.session_state.visualization_type == "Master Parameter Analysis":
                st.subheader("Master Parameter Analysis Visualization")
                
                # Check if we have master columns
                if not master_columns:
                    st.warning("No master columns available. Please go back to the configuration and assign master columns.")
                else:
                    # Let user select which master column to analyze
                    selected_master = st.selectbox("Select master parameter for analysis:", master_columns)
                    
                    # Analyze data grouped by the selected master parameter
                    master_results = analyze_master_parameters(df, [selected_master])
                    
                    if not master_results or selected_master not in master_results:
                        st.error(f"Error analyzing master parameter {selected_master}.")
                    else:
                        results = master_results[selected_master]
                        col_type = results.get('type', 'unknown')
                        st.write(f"Column type: {col_type}")
                        
                        # Check if there's only one group
                        groups = results.get('groups', {})
                        if len(groups) == 1:
                            group_name = list(groups.keys())[0]
                            st.info(f"⚠️ Note: The dataset only contains one value for {selected_master} ({group_name}). " +
                                   f"For comparative analysis, upload data with multiple {selected_master} values.")
                        
                        # Chart type selection
                        chart_type = st.radio(
                            "Select chart type for master parameter analysis:",
                            ["Group Statistics Comparison", "Distribution by Groups", "Pre-Post Analysis by Group"]
                        )
                        
                        # Select a numeric column to visualize across groups
                        num_cols = [col for col in df.columns if pd.api.types.is_numeric_dtype(df[col]) and col != selected_master]
                        if not num_cols:
                            st.warning("No numerical columns available for visualization.")
                        else:
                            if chart_type in ["Group Statistics Comparison", "Distribution by Groups"]:
                                selected_numeric = st.selectbox("Select numerical column to visualize across groups:", num_cols)
                            elif chart_type == "Pre-Post Analysis by Group":
                                # For pre/post analysis
                                pre_post_pairs = st.session_state.get('pre_post_pairs', [])
                                
                                # Let user select pre/post columns
                                col1, col2 = st.columns(2)
                                with col1:
                                    pre_col = st.selectbox("Select Pre column:", num_cols, key="pre_col_selector")
                                with col2:
                                    post_col = st.selectbox("Select Post column:", num_cols, key="post_col_selector")
                            
                            # Handle different chart types
                            if chart_type in ["Group Statistics Comparison", "Distribution by Groups"]:
                                # Create summary dataframe for visualization
                                summary_data = []
                                for group_name, group_data in groups.items():
                                    group_stats = group_data.get('stats')
                                    if group_stats is not None and selected_numeric in group_stats.columns:
                                        # Extract statistics - try to get actual numeric values
                                        try:
                                            mean_val = float(group_stats.loc['Mean', selected_numeric])
                                            std_val = float(group_stats.loc['Std Dev', selected_numeric]) if 'Std Dev' in group_stats.index else 0
                                            median_val = float(group_stats.loc['Median', selected_numeric]) if 'Median' in group_stats.index else mean_val
                                            min_val = float(group_stats.loc['Min', selected_numeric]) if 'Min' in group_stats.index else 0
                                            max_val = float(group_stats.loc['Max', selected_numeric]) if 'Max' in group_stats.index else 0
                                            
                                            if pd.notna(mean_val):
                                                summary_data.append({
                                                    'Group': str(group_name),
                                                    'Mean': mean_val,
                                                    'Std Dev': std_val,
                                                    'Median': median_val,
                                                    'Min': min_val,
                                                    'Max': max_val,
                                                    'Count': group_data['count']
                                                })
                                        except (ValueError, TypeError):
                                            # Skip if we can't extract a numeric value
                                            pass
                            
                            elif chart_type == "Pre-Post Analysis by Group":
                                # Handle pre-post analysis by group
                                if 'pre_col' not in locals() or 'post_col' not in locals():
                                    st.warning("Please select Pre and Post columns for analysis.")
                                else:
                                    # Create a dataframe for pre-post comparison
                                    pre_post_data = []
                                    
                                    # Calculate pre-post stats for each group
                                    for group_name, group_data in groups.items():
                                        group_df = group_data.get('data')
                                        if group_df is not None and pre_col in group_df.columns and post_col in group_df.columns:
                                            pre_mean = group_df[pre_col].mean()
                                            post_mean = group_df[post_col].mean()
                                            pre_std = group_df[pre_col].std()
                                            post_std = group_df[post_col].std()
                                            count = len(group_df)
                                            # Calculate percent change
                                            if pre_mean != 0:
                                                pct_change = ((post_mean - pre_mean) / pre_mean) * 100
                                            else:
                                                pct_change = 0
                                                
                                            # Calculate p-value for pre vs post
                                            try:
                                                from scipy import stats
                                                t_stat, p_value = stats.ttest_rel(group_df[pre_col], group_df[post_col])
                                            except:
                                                p_value = float('nan')
                                            
                                            pre_post_data.append({
                                                'Group': str(group_name),
                                                'Pre Mean': pre_mean,
                                                'Pre StdDev': pre_std,
                                                'Post Mean': post_mean,
                                                'Post StdDev': post_std,
                                                'Change %': pct_change,
                                                'P-value': p_value,
                                                'Count': count
                                            })
                                    
                                    summary_data = pre_post_data
                            
                            if not summary_data:
                                st.warning(f"Could not extract valid statistics for {selected_numeric} across {selected_master} groups.")
                            else:
                                summary_df = pd.DataFrame(summary_data)
                                
                                # Visualization type
                                viz_subtype = st.radio(
                                    "Select visualization type:",
                                    ["Bar Chart", "Box Plot"]
                                )
                                
                                if viz_subtype == "Bar Chart":
                                    import plotly.express as px
                                    import plotly.graph_objects as go
                                    
                                    # Let user select stat to display based on chart type
                                    if chart_type in ["Group Statistics Comparison", "Distribution by Groups"]:
                                        stat_options = ["Mean", "Mean with StdDev", "Median", "Min/Max"]
                                        stat_display = st.selectbox("Choose statistic to display:", stat_options)
                                    elif chart_type == "Pre-Post Analysis by Group":
                                        stat_options = ["Pre-Post Means", "Percent Reduction", "Pre vs Post with P-values"]
                                        stat_display = st.selectbox("Choose pre-post analysis display:", stat_options)
                                    
                                    if stat_display == "Mean":
                                        # Create bar chart comparing means across groups
                                        fig = px.bar(
                                            summary_df,
                                            x='Group',
                                            y='Mean',
                                            title=f"Mean {selected_numeric} by {selected_master}",
                                            labels={'Group': selected_master, 'Mean': f'Mean {selected_numeric}'},
                                            text='Count',  # Show count on bars
                                            color='Group',
                                            color_discrete_sequence=px.colors.qualitative.Plotly
                                        )
                                        
                                        # Add count labels above bars
                                        fig.update_traces(texttemplate='n=%{text}', textposition='outside')
                                    
                                    elif stat_display == "Mean with StdDev":
                                        # Create bar chart with error bars and Excel-like styling
                                        fig = px.bar(
                                            summary_df,
                                            x='Group',
                                            y='Mean',
                                            error_y='Std Dev',
                                            title=f"Mean ± StdDev of {selected_numeric} by {selected_master}",
                                            labels={'Group': selected_master, 'Mean': f'Mean {selected_numeric}'},
                                            color='Group',
                                            color_discrete_sequence=px.colors.qualitative.Plotly,
                                            template="plotly_white"  # Clean white template as base
                                        )
                                        
                                        # Apply Excel theme to the figure
                                        from visualization_utils import apply_excel_theme
                                        fig = apply_excel_theme(fig)
                                        
                                        # Add formatted Mean ± StdDev as text
                                        for i, row in summary_df.iterrows():
                                            fig.add_annotation(
                                                x=row['Group'],
                                                y=row['Mean'] + row['Std Dev'] + 0.2,  # Position above error bar
                                                text=f"{row['Mean']:.2f} ± {row['Std Dev']:.2f} (n={row['Count']})",
                                                showarrow=False,
                                                font=dict(family="Times New Roman, Times, serif", size=10)
                                            )
                                    
                                    elif stat_display == "Median":
                                        # Create bar chart comparing medians across groups with Excel-like styling
                                        fig = px.bar(
                                            summary_df,
                                            x='Group',
                                            y='Median',
                                            title=f"Median {selected_numeric} by {selected_master}",
                                            labels={'Group': selected_master, 'Median': f'Median {selected_numeric}'},
                                            text='Count',
                                            color='Group',
                                            color_discrete_sequence=px.colors.qualitative.Plotly,
                                            template="plotly_white"  # Clean white template as base
                                        )
                                        
                                        # Apply Excel theme to the figure
                                        from visualization_utils import apply_excel_theme
                                        fig = apply_excel_theme(fig)
                                        
                                        fig.update_traces(texttemplate='n=%{text}', textposition='outside')
                                    
                                    elif stat_display == "Min/Max":
                                        # Create a figure with min/max range
                                        fig = go.Figure()
                                        
                                        # Add a bar for each group showing min/max range
                                        for i, row in summary_df.iterrows():
                                            fig.add_trace(go.Bar(
                                                x=[row['Group']],
                                                y=[row['Max'] - row['Min']],
                                                base=[row['Min']],
                                                name=row['Group'],
                                                text=[f"Min: {row['Min']:.2f}<br>Max: {row['Max']:.2f}<br>n={row['Count']}"],
                                                hoverinfo='text'
                                            ))
                                        
                                        fig.update_layout(
                                            title=f"Min/Max Range of {selected_numeric} by {selected_master}",
                                            xaxis_title=selected_master,
                                            yaxis_title=f"{selected_numeric} Range"
                                        )
                                    
                                    # Pre-Post analysis display options
                                    elif stat_display == "Pre-Post Means":
                                        # Create a grouped bar chart for pre and post means
                                        fig = go.Figure()
                                        
                                        # Get pre and post column names
                                        pre_name = pre_col
                                        post_name = post_col
                                        
                                        # Create the grouped bar chart - pre values
                                        fig.add_trace(go.Bar(
                                            x=summary_df['Group'],
                                            y=[row['Pre Mean'] for i, row in summary_df.iterrows()],
                                            name=f"{pre_name}",
                                            text=[f"{row['Pre Mean']:.2f}<br>n={row['Count']}" for i, row in summary_df.iterrows()],
                                            textposition='auto',
                                            marker_color='#2ca02c'  # Green
                                        ))
                                        
                                        # Add post values
                                        fig.add_trace(go.Bar(
                                            x=summary_df['Group'],
                                            y=[row['Post Mean'] for i, row in summary_df.iterrows()],
                                            name=f"{post_name}",
                                            text=[f"{row['Post Mean']:.2f}<br>n={row['Count']}" for i, row in summary_df.iterrows()],
                                            textposition='auto',
                                            marker_color='#d62728'  # Red
                                        ))
                                        
                                        # Update layout
                                        fig.update_layout(
                                            title=f"Average Pre and Post {pre_name.replace('Pre', '')} by {selected_master}",
                                            xaxis_title=selected_master,
                                            yaxis_title=f"Value",
                                            barmode='group'
                                        )
                                        
                                    elif stat_display == "Percent Reduction":
                                        # Create a bar chart showing percent reduction
                                        fig = go.Figure()
                                        
                                        # Add percent change bars
                                        fig.add_trace(go.Bar(
                                            x=summary_df['Group'],
                                            y=[-row['Change %'] for i, row in summary_df.iterrows()],  # Negate to show reduction as positive
                                            text=[f"{abs(row['Change %']):.2f}%" for i, row in summary_df.iterrows()],
                                            textposition='outside',
                                            marker_color='#1f77b4'  # Blue
                                        ))
                                        
                                        # Update layout with Excel-like styling
                                        fig.update_layout(
                                            title=f"Percentage Reduction in {pre_col.replace('Pre', '')} by {selected_master}",
                                            xaxis_title=selected_master,
                                            yaxis_title="Reduction (%)",
                                            font=dict(family="Times New Roman, Times, serif"),
                                            title_font=dict(family="Times New Roman, Times, serif", size=14),
                                            legend=dict(font=dict(family="Times New Roman, Times, serif")),
                                            template="plotly_white"  # Clean white template as base
                                        )
                                        
                                        # Apply Excel theme for consistent styling
                                        from visualization_utils import apply_excel_theme
                                        fig = apply_excel_theme(fig)
                                        
                                    elif stat_display == "Pre vs Post with P-values":
                                        # Create a grouped bar chart for pre and post with p-values annotation
                                        fig = go.Figure()
                                        
                                        # Get pre and post column names
                                        pre_name = pre_col
                                        post_name = post_col
                                        
                                        # Create the grouped bar chart - pre values
                                        fig.add_trace(go.Bar(
                                            x=summary_df['Group'],
                                            y=[row['Pre Mean'] for i, row in summary_df.iterrows()],
                                            name=f"{pre_name}",
                                            error_y=dict(
                                                type='data',
                                                array=[row['Pre StdDev'] for i, row in summary_df.iterrows()],
                                                visible=True
                                            ),
                                            marker_color='#2ca02c'  # Green
                                        ))
                                        
                                        # Add post values
                                        fig.add_trace(go.Bar(
                                            x=summary_df['Group'],
                                            y=[row['Post Mean'] for i, row in summary_df.iterrows()],
                                            name=f"{post_name}",
                                            error_y=dict(
                                                type='data',
                                                array=[row['Post StdDev'] for i, row in summary_df.iterrows()],
                                                visible=True
                                            ),
                                            marker_color='#d62728'  # Red
                                        ))
                                        
                                        # Add p-value annotations
                                        for i, row in summary_df.iterrows():
                                            p_val = row['P-value']
                                            sig_text = ""
                                            if not pd.isna(p_val):
                                                if p_val < 0.001:
                                                    sig_text = "p<0.001***"
                                                elif p_val < 0.01:
                                                    sig_text = f"p={p_val:.3f}**"
                                                elif p_val < 0.05:
                                                    sig_text = f"p={p_val:.3f}*"
                                                else:
                                                    sig_text = f"p={p_val:.3f}"
                                            
                                            fig.add_annotation(
                                                x=row['Group'],
                                                y=max(row['Pre Mean'], row['Post Mean']) + 
                                                   max(row['Pre StdDev'], row['Post StdDev']) + 0.5,
                                                text=sig_text,
                                                showarrow=False,
                                                font=dict(family="Times New Roman, Times, serif", size=10)
                                            )
                                        
                                        # Update layout with Excel-like styling
                                        fig.update_layout(
                                            title=f"Pre vs Post {pre_name.replace('Pre', '')} with P-values by {selected_master}",
                                            xaxis_title=selected_master,
                                            yaxis_title=f"Value",
                                            barmode='group',
                                            font=dict(family="Times New Roman, Times, serif"),
                                            title_font=dict(family="Times New Roman, Times, serif", size=14),
                                            legend=dict(font=dict(family="Times New Roman, Times, serif")),
                                            template="plotly_white"  # Clean white template as base
                                        )
                                        
                                        # Apply Excel theme for consistent styling
                                        from visualization_utils import apply_excel_theme
                                        fig = apply_excel_theme(fig)
                                    
                                    st.plotly_chart(fig, use_container_width=True)
                                    
                                    # Store visualization for output
                                    st.session_state.visualization = {
                                        'type': 'master_bar_chart',
                                        'figure': fig,
                                        'settings': {
                                            'master_column': selected_master,
                                            'numeric_column': selected_numeric
                                        }
                                    }
                                    
                                elif viz_subtype == "Box Plot":
                                    # For box plot, we need the actual data
                                    # Create a box plot grouped by master parameter
                                    fig = create_boxplot(df, selected_numeric, selected_master)
                                    st.plotly_chart(fig, use_container_width=True)
                                    
                                    # Store visualization for output
                                    st.session_state.visualization = {
                                        'type': 'master_box_plot',
                                        'figure': fig,
                                        'settings': {
                                            'master_column': selected_master,
                                            'numeric_column': selected_numeric
                                        }
                                    }
            
            elif st.session_state.visualization_type == "Histogram":
                st.subheader("Histogram")
                # Select column for histogram
                all_cols = st.session_state.selected_columns
                if all_cols:
                    hist_col = st.selectbox("Select column for histogram:", all_cols)
                    bins = st.slider("Number of bins:", min_value=5, max_value=100, value=20)
                    
                    # Create histogram
                    fig = create_histogram(df, hist_col, bins)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Create unique chart ID
                    chart_id = f"hist_{hist_col}_{bins}"
                    chart_title = f"Histogram of {hist_col} ({bins} bins)"
                    
                    # Check if this chart is in favorites
                    is_favorite = chart_id in st.session_state.get('chart_favorites', {})
                    
                    # Create heart button with active state based on favorite status
                    col1, col2 = st.columns([9, 1])
                    with col2:
                        heart_color = "red" if is_favorite else "gray"
                        heart_icon = "❤️" if is_favorite else "🤍"
                        if st.button(heart_icon, key=f"fav_btn_{chart_id}", help="Add/remove from favorites"):
                            toggle_chart_favorite(chart_id, fig, chart_title, "histogram")
                            st.rerun()
                    
                    # Store visualization for output
                    st.session_state.visualization = {
                        'type': 'histogram',
                        'figure': fig,
                        'settings': {
                            'column': hist_col,
                            'bins': bins
                        }
                    }
                else:
                    st.warning("No columns available for histogram.")
            
            elif st.session_state.visualization_type == "Box Plot":
                st.subheader("Box Plot")
                # Select column for box plot
                all_cols = st.session_state.selected_columns
                if all_cols:
                    box_col = st.selectbox("Select column for box plot:", all_cols)
                    
                    # Optional grouping
                    group_by = None
                    use_grouping = st.checkbox("Group by another variable")
                    if use_grouping:
                        # Filter out the selected column
                        grouping_options = [col for col in all_cols if col != box_col]
                        if grouping_options:
                            group_by = st.selectbox("Select grouping variable:", grouping_options)
                    
                    # Create box plot
                    fig = create_boxplot(df, box_col, group_by)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Store visualization for output
                    st.session_state.visualization = {
                        'type': 'boxplot',
                        'figure': fig,
                        'settings': {
                            'column': box_col,
                            'group_by': group_by
                        }
                    }
                else:
                    st.warning("No columns available for box plot.")
            
            elif st.session_state.visualization_type == "Scatter Plot":
                st.subheader("Scatter Plot")
                # Select columns for scatter plot
                all_cols = st.session_state.selected_columns
                if len(all_cols) >= 2:
                    x_col = st.selectbox("Select X-axis column:", all_cols, index=0)
                    # Filter out the X column from Y options
                    y_options = [col for col in all_cols if col != x_col]
                    y_col = st.selectbox("Select Y-axis column:", y_options, index=0)
                    
                    # Optional color grouping
                    color_by = None
                    # Filter out X and Y columns 
                    color_options = [col for col in all_cols if col != x_col and col != y_col]
                    if color_options:
                        use_color = st.checkbox("Color by another variable")
                        if use_color:
                            color_by = st.selectbox("Select color variable:", color_options)
                    
                    # Create scatter plot
                    fig = create_scatterplot(df, x_col, y_col, color_by)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Create heart button for scatter plot
                    chart_title = f"Scatter Plot: {y_col} vs {x_col}"
                    if color_by:
                        chart_title += f" (colored by {color_by})"
                    
                    # Create unique chart ID for scatter plot
                    chart_id = f"scatter_{x_col}_{y_col}_{color_by}" if color_by else f"scatter_{x_col}_{y_col}"
                    
                    # Check if this chart is in favorites
                    is_favorite = chart_id in st.session_state.get('chart_favorites', {})
                    
                    # Create heart button with active state based on favorite status
                    col1, col2 = st.columns([9, 1])
                    with col2:
                        heart_icon = "❤️" if is_favorite else "🤍"
                        if st.button(heart_icon, key=f"fav_btn_{chart_id}", help="Add/remove from favorites"):
                            toggle_chart_favorite(chart_id, fig, chart_title, "scatter_plot")
                            st.rerun()
                    
                    # Store visualization for output
                    st.session_state.visualization = {
                        'type': 'scatterplot',
                        'figure': fig,
                        'settings': {
                            'x_col': x_col,
                            'y_col': y_col,
                            'color_by': color_by
                        }
                    }
                else:
                    st.warning("At least two columns are required for a scatter plot.")
            
            elif st.session_state.visualization_type == "Correlation Heatmap":
                st.subheader("Correlation Heatmap")
                # Select columns for correlation
                all_cols = st.session_state.selected_columns
                if len(all_cols) >= 2:
                    corr_cols = st.multiselect(
                        "Select columns for correlation heatmap:",
                        all_cols,
                        default=all_cols[:min(5, len(all_cols))]
                    )
                    
                    if len(corr_cols) >= 2:
                        corr_method = st.selectbox(
                            "Correlation method:",
                            ["pearson", "spearman", "kendall"]
                        )
                        
                        # Create correlation heatmap
                        fig = create_correlation_heatmap(df, corr_cols, corr_method)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Create heart button for correlation heatmap
                        chart_title = f"Correlation Heatmap ({corr_method} method)"
                        
                        # Create unique chart ID for correlation heatmap
                        chart_id = f"corr_heatmap_{corr_method}"
                        
                        # Check if this chart is in favorites
                        is_favorite = chart_id in st.session_state.get('chart_favorites', {})
                        
                        # Create heart button with active state based on favorite status
                        col1, col2 = st.columns([9, 1])
                        with col2:
                            heart_icon = "❤️" if is_favorite else "🤍"
                            if st.button(heart_icon, key=f"fav_btn_{chart_id}", help="Add/remove from favorites"):
                                toggle_chart_favorite(chart_id, fig, chart_title, "correlation_heatmap")
                                st.rerun()
                        
                        # Store visualization for output
                        st.session_state.visualization = {
                            'type': 'correlation_heatmap',
                            'figure': fig,
                            'settings': {
                                'columns': corr_cols,
                                'method': corr_method
                            }
                        }
                    else:
                        st.info("Please select at least two columns for correlation heatmap.")
                else:
                    st.warning("At least two columns are required for a correlation heatmap.")
            
            elif st.session_state.visualization_type == "Distribution Plot":
                st.subheader("Distribution Plot")
                # Select column for distribution plot
                all_cols = st.session_state.selected_columns
                if all_cols:
                    dist_col = st.selectbox("Select column for distribution:", all_cols)
                    
                    # Create distribution plot
                    fig = create_distribution_plot(df, dist_col)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Heart button for distribution plot
                    chart_title = f"Distribution Plot: {dist_col}"
                    chart_id = f"dist_{dist_col}_{int(time.time())}"
                    
                    # Check if this chart is in favorites
                    is_favorite = chart_id in st.session_state.get('chart_favorites', {})
                    
                    # Create heart button with active state based on favorite status
                    col1, col2 = st.columns([9, 1])
                    with col2:
                        heart_icon = "❤️" if is_favorite else "🤍"
                        if st.button(heart_icon, key=f"fav_btn_{chart_id}", help="Add/remove from favorites"):
                            toggle_chart_favorite(chart_id, fig, chart_title, "distribution_plot")
                            st.rerun()
                    
                    # Store visualization for output
                    st.session_state.visualization = {
                        'type': 'distribution_plot',
                        'figure': fig,
                        'settings': {
                            'column': dist_col
                        }
                    }
                else:
                    st.warning("No columns available for distribution plot.")
            
            elif st.session_state.visualization_type == "Bar Chart":
                st.subheader("Bar Chart")
                
                # Option for categorical data frequency or numeric summary by category
                chart_type = st.radio(
                    "Select bar chart type:",
                    ["Categorical Frequency", "Numeric Summary by Category"]
                )
                
                if chart_type == "Categorical Frequency":
                    cat_cols = [col for col in st.session_state.selected_columns if not pd.api.types.is_numeric_dtype(df[col])]
                    if cat_cols:
                        cat_col = st.selectbox("Select categorical column:", cat_cols)
                        # Get value counts and create a bar chart
                        value_counts = df[cat_col].value_counts().reset_index()
                        value_counts.columns = [cat_col, 'Count']
                        
                        # Limit display to top N categories if there are too many
                        top_n = st.slider("Show top N categories:", min_value=5, max_value=50, value=10)
                        if len(value_counts) > top_n:
                            value_counts = value_counts.head(top_n)
                            st.info(f"Showing top {top_n} categories only.")
                        
                        # Create bar chart with Excel-like styling
                        fig = px.bar(
                            value_counts, 
                            x=cat_col, 
                            y='Count',
                            title=f"Frequency of {cat_col} Categories",
                            labels={cat_col: cat_col, 'Count': 'Frequency'},
                            template="plotly_white",  # Clean white template as base
                            color_discrete_sequence=px.colors.qualitative.Plotly  # Excel-like colors
                        )
                        
                        # Add Excel-like data labels on top of bars
                        fig.update_traces(
                            texttemplate='%{y}', 
                            textposition='outside',
                            textfont=dict(family="Times New Roman", size=10)
                        )
                        
                        # Apply Excel theme to the figure
                        from visualization_utils import apply_excel_theme
                        fig = apply_excel_theme(fig)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Store visualization for output
                        st.session_state.visualization = {
                            'type': 'bar_chart_frequency',
                            'figure': fig,
                            'settings': {
                                'column': cat_col,
                                'top_n': top_n
                            }
                        }
                    else:
                        st.warning("No categorical columns available for frequency bar chart.")
                
                elif chart_type == "Numeric Summary by Category":
                    cat_cols = [col for col in st.session_state.selected_columns if not pd.api.types.is_numeric_dtype(df[col])]
                    num_cols = [col for col in st.session_state.selected_columns if pd.api.types.is_numeric_dtype(df[col])]
                    
                    if cat_cols and num_cols:
                        cat_col = st.selectbox("Select categorical column (X-axis):", cat_cols)
                        num_col = st.selectbox("Select numerical column to summarize:", num_cols)
                        
                        # Select summary statistic
                        summary_stat = st.selectbox(
                            "Select summary statistic:",
                            ["Mean", "Median", "Sum", "Min", "Max"]
                        )
                        
                        # Group by category and calculate the selected statistic
                        if summary_stat == "Mean":
                            summary_df = df.groupby(cat_col)[num_col].mean().reset_index()
                            stat_name = "Mean"
                        elif summary_stat == "Median":
                            summary_df = df.groupby(cat_col)[num_col].median().reset_index()
                            stat_name = "Median"
                        elif summary_stat == "Sum":
                            summary_df = df.groupby(cat_col)[num_col].sum().reset_index()
                            stat_name = "Sum"
                        elif summary_stat == "Min":
                            summary_df = df.groupby(cat_col)[num_col].min().reset_index()
                            stat_name = "Min"
                        elif summary_stat == "Max":
                            summary_df = df.groupby(cat_col)[num_col].max().reset_index()
                            stat_name = "Max"
                        
                        # Limit display to top N categories if there are too many
                        if len(summary_df) > 15:
                            top_n = st.slider("Show top N categories by value:", min_value=5, max_value=30, value=15)
                            summary_df = summary_df.sort_values(num_col, ascending=False).head(top_n)
                            st.info(f"Showing top {top_n} categories by {stat_name.lower()} value.")
                        
                        # Create bar chart with Excel-like styling
                        fig = px.bar(
                            summary_df, 
                            x=cat_col, 
                            y=num_col,
                            title=f"{stat_name} of {num_col} by {cat_col}",
                            labels={cat_col: cat_col, num_col: f"{stat_name} of {num_col}"},
                            template="plotly_white",  # Clean white template as base
                            color_discrete_sequence=px.colors.qualitative.Plotly  # Excel-like colors
                        )
                        
                        # Add Excel-like data labels on top of bars
                        fig.update_traces(
                            texttemplate='%{y:.2f}',  # Format to 2 decimal places
                            textposition='outside',
                            textfont=dict(family="Times New Roman", size=10)
                        )
                        
                        # Apply Excel theme to the figure
                        from visualization_utils import apply_excel_theme
                        fig = apply_excel_theme(fig)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Store visualization for output
                        st.session_state.visualization = {
                            'type': 'bar_chart_summary',
                            'figure': fig,
                            'settings': {
                                'cat_col': cat_col,
                                'num_col': num_col,
                                'summary_stat': summary_stat
                            }
                        }
                    else:
                        if not cat_cols:
                            st.warning("No categorical columns available for grouping.")
                        if not num_cols:
                            st.warning("No numerical columns available to summarize.")
            
            elif st.session_state.visualization_type == "Line Chart":
                st.subheader("Line Chart")
                
                # Check if there's any datetime column that could be used for a time series
                has_datetime = any(pd.api.types.is_datetime64_dtype(df[col]) for col in st.session_state.selected_columns)
                
                if has_datetime:
                    # Time series line chart
                    datetime_cols = [col for col in st.session_state.selected_columns if pd.api.types.is_datetime64_dtype(df[col])]
                    time_col = st.selectbox("Select time/date column (X-axis):", datetime_cols)
                    
                    other_cols = [col for col in st.session_state.selected_columns if col != time_col]
                    if other_cols:
                        y_cols = st.multiselect("Select column(s) to plot:", other_cols, default=[other_cols[0]])
                        
                        if y_cols:
                            # Create a copy of relevant data, sorted by time
                            plot_df = df[[time_col] + y_cols].sort_values(time_col).copy()
                            
                            # Create line chart with Excel-like styling
                            fig = px.line(
                                plot_df, 
                                x=time_col, 
                                y=y_cols,
                                title=f"Time Series of {', '.join(y_cols)}",
                                labels={col: col for col in [time_col] + y_cols},
                                template="plotly_white",  # Clean white template as base
                                color_discrete_sequence=px.colors.qualitative.Plotly  # Excel-like colors
                            )
                            
                            # Apply Excel theme to the figure
                            from visualization_utils import apply_excel_theme
                            fig = apply_excel_theme(fig)
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Store visualization for output
                            st.session_state.visualization = {
                                'type': 'line_chart_time',
                                'figure': fig,
                                'settings': {
                                    'time_col': time_col,
                                    'y_cols': y_cols
                                }
                            }
                        else:
                            st.info("Please select at least one column to plot.")
                    else:
                        st.warning("No columns available for line chart.")
                
                else:
                    # Regular line chart (not time series)
                    all_cols = st.session_state.selected_columns
                    if len(all_cols) >= 2:
                        x_col = st.selectbox("Select X-axis column:", all_cols, index=0)
                        # Filter out the X column from Y options
                        y_options = [col for col in all_cols if col != x_col]
                        y_cols = st.multiselect("Select Y-axis column(s):", y_options, default=[y_options[0]])
                        
                        if y_cols:
                            # Create a copy of relevant data, sorted by X
                            plot_df = df[[x_col] + y_cols].sort_values(x_col).copy()
                            
                            # Create line chart with Excel-like styling
                            fig = px.line(
                                plot_df, 
                                x=x_col, 
                                y=y_cols,
                                title=f"Line Chart of {', '.join(y_cols)} vs {x_col}",
                                labels={col: col for col in [x_col] + y_cols},
                                template="plotly_white",  # Clean white template as base
                                color_discrete_sequence=px.colors.qualitative.Plotly  # Excel-like colors
                            )
                            
                            # Apply Excel theme to the figure
                            from visualization_utils import apply_excel_theme
                            fig = apply_excel_theme(fig)
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Store visualization for output
                            st.session_state.visualization = {
                                'type': 'line_chart',
                                'figure': fig,
                                'settings': {
                                    'x_col': x_col,
                                    'y_cols': y_cols
                                }
                            }
                        else:
                            st.info("Please select at least one Y-axis column.")
                    else:
                        st.warning("At least two columns are required for a line chart.")
            
            elif st.session_state.visualization_type == "Frequency Analysis":
                st.subheader("Frequency Analysis with Pie Charts")
                
                st.info("This visualization shows frequency distribution with both count and percentage pie charts.")
                
                # Get categorical columns
                categorical_cols = [col for col in df.columns if not pd.api.types.is_numeric_dtype(df[col]) or df[col].nunique() < 10]
                
                if not categorical_cols:
                    st.warning("No suitable categorical columns found in the data for frequency analysis.")
                else:
                    # Column selection
                    selected_col = st.selectbox(
                        "Select a categorical column to analyze:", 
                        categorical_cols, 
                        key="freq_analysis_col"
                    )
                    
                    # Sorting options
                    sort_options = ["Count (descending)", "Alphabetical"]
                    sort_by = st.radio("Sort results by:", sort_options)
                    
                    # Display options
                    chart_options = ["Percentage pie chart", "Count pie chart", "Both pie charts", "Bar chart", "All chart types"]
                    display_option = st.radio("Select display option:", chart_options)
                    
                    # Map display option to chart_type parameter
                    chart_type_mapping = {
                        "Percentage pie chart": "pie", 
                        "Count pie chart": "pie",
                        "Both pie charts": "pie",
                        "Bar chart": "bar",
                        "All chart types": "both"
                    }
                    
                    chart_type = chart_type_mapping[display_option]
                    
                    # Calculate value counts and create pie chart
                    from visualization_utils import create_frequency_analysis
                    
                    # Convert sorting option to parameter value
                    sort_param = "count" if sort_by == "Count (descending)" else "alphabetical"
                    
                    # Generate frequency analysis and pie chart(s)
                    value_counts_df, figures = create_frequency_analysis(
                        df, 
                        selected_col, 
                        include_percentage=True, 
                        sort_by=sort_param,
                        chart_type=chart_type
                    )
                    
                    # Display the frequency table in Excel-like format
                    st.subheader(f"Frequency Distribution of {selected_col}")
                    
                    # Add a total row to the dataframe
                    total_row = pd.DataFrame({
                        selected_col: ['Total'],
                        'Count': [value_counts_df['Count'].sum()],
                        'Percentage': ['100.0%']
                    })
                    
                    # Combine with original dataframe
                    display_df = pd.concat([value_counts_df, total_row], ignore_index=True)
                    
                    # Apply Excel-like styling with conditional formatting
                    def highlight_total_row(row):
                        """Highlight the total row with gray background"""
                        if row.name == len(display_df) - 1:  # If it's the last row (total)
                            return ['background-color: #E0E0E0'] * len(row)
                        else:
                            return [''] * len(row)
                    
                    # Apply the styling
                    styled_df = display_df.style.apply(highlight_total_row, axis=1)
                    
                    # Apply Excel-like header styling
                    styled_df = styled_df.set_table_styles([
                        {'selector': 'thead th', 
                         'props': [('background-color', '#F8F9FA'), 
                                  ('color', '#495057'),
                                  ('font-weight', 'bold'),
                                  ('border', '1px solid #DFE0E1'),
                                  ('text-align', 'center')]}
                    ])
                    
                    # Display the styled dataframe
                    st.dataframe(styled_df, use_container_width=True)
                    
                    # Display the generated charts based on selection
                    if isinstance(figures, list):
                        # If we have multiple figures
                        if display_option == "Both pie charts" or display_option == "All chart types":
                            col1, col2 = st.columns(2)
                            
                            # Display percentage pie chart in first column
                            with col1:
                                chart_title = f"Percentage Distribution of {selected_col}"
                                st.subheader(chart_title)
                                
                                chart_id = f"freq_percent_{selected_col}_{int(time.time())}"
                                
                                # Create chart display with heart button
                                chart_col, heart_col = st.columns([5, 1])
                                with chart_col:
                                    st.plotly_chart(figures[0], use_container_width=True)
                                with heart_col:
                                    is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                    heart_color = "❤️" if is_favorited else "🤍"
                                    if st.button(
                                        heart_color, 
                                        key=f"heart_percent_{chart_id}",
                                        help="Add to favorites"
                                    ):
                                        toggle_chart_favorite(chart_id, figures[0], chart_title, "percentage_pie")
                                        st.rerun()
                            
                            # Display count pie chart in second column
                            with col2:
                                chart_title = f"Count Distribution of {selected_col}"
                                st.subheader(chart_title)
                                
                                chart_id = f"freq_count_{selected_col}_{int(time.time())}"
                                
                                # Create chart display with heart button
                                chart_col, heart_col = st.columns([5, 1])
                                with chart_col:
                                    st.plotly_chart(figures[1], use_container_width=True)
                                with heart_col:
                                    is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                    heart_color = "❤️" if is_favorited else "🤍"
                                    if st.button(
                                        heart_color, 
                                        key=f"heart_count_{chart_id}",
                                        help="Add to favorites"
                                    ):
                                        toggle_chart_favorite(chart_id, figures[1], chart_title, "count_pie")
                                        st.rerun()
                            
                            # If we have a bar chart, display it full width
                            if len(figures) > 2:
                                chart_title = f"Bar Chart Distribution of {selected_col}"
                                st.subheader(chart_title)
                                
                                chart_id = f"freq_bar_{selected_col}_{int(time.time())}"
                                
                                # Create chart display with heart button
                                chart_col, heart_col = st.columns([5, 1])
                                with chart_col:
                                    st.plotly_chart(figures[2], use_container_width=True)
                                with heart_col:
                                    is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                    heart_color = "❤️" if is_favorited else "🤍"
                                    if st.button(
                                        heart_color, 
                                        key=f"heart_bar_{chart_id}",
                                        help="Add to favorites"
                                    ):
                                        toggle_chart_favorite(chart_id, figures[2], chart_title, "bar_chart")
                                        st.rerun()
                        else:
                            # Display just one chart type
                            for i, fig in enumerate(figures):
                                chart_type_name = "Percentage" if i == 0 else "Count" if i == 1 else "Bar"
                                chart_title = f"{chart_type_name} Distribution of {selected_col}"
                                st.subheader(chart_title)
                                
                                chart_id = f"freq_multi_{i}_{selected_col}_{int(time.time())}"
                                chart_type_key = "percentage_pie" if i == 0 else "count_pie" if i == 1 else "bar_chart"
                                
                                # Create chart display with heart button
                                chart_col, heart_col = st.columns([5, 1])
                                with chart_col:
                                    st.plotly_chart(fig, use_container_width=True)
                                with heart_col:
                                    is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                                    heart_color = "❤️" if is_favorited else "🤍"
                                    if st.button(
                                        heart_color, 
                                        key=f"heart_multi_{chart_id}",
                                        help="Add to favorites"
                                    ):
                                        toggle_chart_favorite(chart_id, fig, chart_title, chart_type_key)
                                        st.rerun()
                    else:
                        # If we have a single figure
                        chart_type_name = "Percentage" if display_option == "Percentage pie chart" else "Count" if display_option == "Count pie chart" else "Bar"
                        chart_title = f"{chart_type_name} Distribution of {selected_col}"
                        st.subheader(chart_title)
                        
                        chart_id = f"freq_single_{selected_col}_{int(time.time())}"
                        chart_type_key = "percentage_pie" if display_option == "Percentage pie chart" else "count_pie" if display_option == "Count pie chart" else "bar_chart"
                        
                        # Create chart display with heart button
                        chart_col, heart_col = st.columns([5, 1])
                        with chart_col:
                            st.plotly_chart(figures, use_container_width=True)
                        with heart_col:
                            is_favorited = chart_id in st.session_state.get('chart_favorites', {})
                            heart_color = "❤️" if is_favorited else "🤍"
                            if st.button(
                                heart_color, 
                                key=f"heart_single_{chart_id}",
                                help="Add to favorites"
                            ):
                                toggle_chart_favorite(chart_id, figures, chart_title, chart_type_key)
                                st.rerun()
                    
                    # Store visualization for output
                    st.session_state.visualization = {
                        'type': 'frequency_analysis',
                        'figure': figures[0] if isinstance(figures, list) else figures,
                        'settings': {
                            'column': selected_col,
                            'sort_by': sort_param,
                            'chart_type': chart_type
                        }
                    }
            
            # Visualization customization options
            with st.expander("Customize Visualization"):
                st.write("Customize your visualization with the following options:")
                
                # Title customization
                custom_title = st.text_input("Chart Title:", "")
                
                # Color scheme selection
                color_schemes = ["blues", "reds", "greens", "purples", "oranges", "viridis", "plasma", "inferno", "magma", "cividis"]
                color_scheme = st.selectbox("Color Scheme:", color_schemes)
                
                # Apply customizations if a visualization exists
                if 'visualization' in st.session_state and st.button("Apply Customizations"):
                    if 'figure' in st.session_state.visualization:
                        fig = st.session_state.visualization['figure']
                        
                        # Update title if provided
                        if custom_title:
                            fig.update_layout(title=custom_title)
                        
                        # Update color scheme
                        fig.update_layout(coloraxis=dict(colorscale=color_scheme))
                        
                        # Update the visualization in session state
                        st.session_state.visualization['figure'] = fig
                        st.session_state.visualization['customizations'] = {
                            'title': custom_title,
                            'color_scheme': color_scheme
                        }
                        
                        # Display the updated visualization
                        st.success("Customizations applied!")
                        st.plotly_chart(fig, use_container_width=True)
            
            # Button to proceed to output
            if st.button("Continue to Output"):
                st.session_state.current_step = 6
                # Add JavaScript to scroll to top after rerun
                st.markdown(
                    """
                    <script>
                        window.scrollTo(0, 0);
                    </script>
                    """,
                    unsafe_allow_html=True
                )
                st.rerun()
    
    elif st.session_state.current_step == 6:
        # Step 6: Reference Data
        st.markdown('<div id="reference-data"></div>', unsafe_allow_html=True)
        st.header(emoji_header("6. Reference Data", "Research"))
        
        # Import the reference data UI function
        from web_data_utils import reference_data_ui
        
        # Show the reference data interface
        reference_data_ui()
        
    elif st.session_state.current_step == 7:
        # Step 7: Output
        st.markdown('<div id="output"></div>', unsafe_allow_html=True)
        st.header(emoji_header("7. Output", "Output"))
        
        if st.session_state.data is not None:
            # Create tabs for output section
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Charts & Export", "📋 Data Exploration", "📈 Visualizations", "📊 Statistics", "🔗 Correlations"])
            
            # Tab 1: Charts and Export
            with tab1:
                # Enhanced Favorites section with business styling
                st.markdown("""
                <div class="business-section-header">
                    <h2>📊 Favorited Charts Collection</h2>
                    <p>Manage and export your selected charts and visualizations</p>
                </div>
                """, unsafe_allow_html=True)
            
                # Check if chart_favorites exists in session state
                if 'chart_favorites' not in st.session_state:
                    st.session_state.chart_favorites = {}
                
                if len(st.session_state.chart_favorites) > 0:
                    # Enhanced metrics display
                    st.markdown(f"""
                    <div class="metrics-card">
                        <div class="metric-item">
                            <div class="metric-number">{len(st.session_state.chart_favorites)}</div>
                            <div class="metric-label">Charts Selected</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                    # Enhanced action buttons
                    col1, col2, col3 = st.columns([2, 1, 1])
                    
                    with col1:
                        if st.button("📄 Export All Charts to Microsoft Word", type="primary", use_container_width=True):
                            try:
                                with st.spinner("Generating Word document..."):
                                    doc_bytes = export_charts_to_word()
                                    # Create download link with better styling
                                    b64_word = base64.b64encode(doc_bytes.getvalue()).decode()
                                    st.markdown(f"""
                                    <div class="download-container">
                                        <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_word}" 
                                           download="statistical_analysis_charts.docx" 
                                           class="download-button">
                                            📥 Download Word Document
                                        </a>
                                    </div>
                                    """, unsafe_allow_html=True)
                                    st.success("Word document generated successfully! Click the download button above.")
                                    show_confetti()
                            except Exception as e:
                                st.error(f"Error creating Word document: {str(e)}")
                    
                    with col2:
                        if st.button("👁️ Preview Charts", use_container_width=True):
                            st.session_state.show_preview = not st.session_state.get('show_preview', False)
                    
                    with col3:
                        if st.button("🗑️ Clear All", use_container_width=True):
                            st.session_state.chart_favorites = {}
                            st.rerun()
                    
                    # Enhanced chart preview section
                    if st.session_state.get('show_preview', False):
                        st.markdown("### 📋 Chart Preview")
                        with st.expander("View All Favorited Charts", expanded=True):
                            for i, (chart_id, chart_item) in enumerate(st.session_state.chart_favorites.items()):
                                st.markdown(f"""
                                <div class="chart-preview-item">
                                    <h4>{i+1}. {chart_item['title']}</h4>
                                    <span class="chart-type-badge">{chart_item.get('chart_type', 'Chart')}</span>
                                </div>
                                """, unsafe_allow_html=True)
                                
                                col_chart, col_remove = st.columns([5, 1])
                                with col_chart:
                                    st.plotly_chart(chart_item['figure'], use_container_width=True)
                                with col_remove:
                                    if st.button("❌", key=f"remove_fav_{chart_id}", help="Remove from favorites"):
                                        remove_from_favorites(chart_id)
                                        st.rerun()
                                
                                if i < len(st.session_state.chart_favorites) - 1:
                                    st.divider()
                else:
                    # Enhanced empty state
                    st.markdown("""
                    <div class="empty-state">
                        <div class="empty-state-icon">📈</div>
                        <h3>No Charts Selected</h3>
                        <p>Browse charts in the <strong>Graph/Chart Preparation</strong> section and click the heart ❤️ icon to add charts to your favorites for export.</p>
                    </div>
                    """, unsafe_allow_html=True)
            
                st.divider()
                
                # HTML Report section
                st.subheader("Generate Analysis Report")
            
                # Options for report content
                include_summary = st.checkbox("Include Data Summary", value=True)
                include_stats = st.checkbox("Include Descriptive Statistics", value=True)
                include_viz = st.checkbox("Include Visualizations", value=True)
                include_corr = st.checkbox("Include Correlation Analysis", value=True)
                
                if st.button("Generate HTML Report"):
                    try:
                        # Generate the HTML report
                        html_report = generate_report(
                            st.session_state.data[st.session_state.selected_columns] if st.session_state.selected_columns else st.session_state.data,
                            filename=st.session_state.filename,
                            include_summary=include_summary,
                            include_stats=include_stats,
                            include_viz=include_viz,
                            include_corr=include_corr
                        )
                        
                        # Create a download link
                        b64 = base64.b64encode(html_report.encode("utf-8")).decode()
                        href = f'<a href="data:text/html;base64,{b64}" download="statistical_analysis_report.html">Download HTML Report</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        
                        # Show a preview
                        with st.expander("Preview Report"):
                            st.components.v1.html(html_report, height=500, scrolling=True)
                        
                        # Show confetti and success message with emoji
                        show_confetti()
                        success_emoji = random.choice(EMOJI_DICT["Success"])
                        st.success(f"{success_emoji} Report generated successfully! {success_emoji} Click the link above to download.")
                    
                    except Exception as e:
                        st.error(f"Error generating report: {str(e)}")
                
                # Export data option
                st.subheader("Export Analysis Data")
            
                export_format = st.selectbox(
                    "Select export format:",
                    ["CSV", "Excel"]
                )
                
                if st.button("Export Data"):
                    try:
                        export_df = st.session_state.data[st.session_state.selected_columns] if st.session_state.selected_columns else st.session_state.data
                        
                        if export_format == "CSV":
                            csv = export_df.to_csv(index=False)
                            b64 = base64.b64encode(csv.encode()).decode()
                            href = f'<a href="data:file/csv;base64,{b64}" download="analysis_data.csv">Download CSV</a>'
                            st.markdown(href, unsafe_allow_html=True)
                        
                        elif export_format == "Excel":
                            output = io.BytesIO()
                            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                export_df.to_excel(writer, index=False, sheet_name='Data')
                                
                                # Add a sheet for statistics if available
                                if 'analysis_results' in st.session_state:
                                    result = st.session_state.analysis_results
                                    if result['type'] == 'descriptive' and 'stats_df' in result:
                                        result['stats_df'].to_excel(writer, sheet_name='Statistics')
                                    elif result['type'] in ['ttest', 'anova', 'correlation'] and 'result' in result:
                                        # Create a summary sheet with the results
                                        summary_df = pd.DataFrame({'Key': ['Analysis Type'], 'Value': [result['type']]})
                                        summary_df.to_excel(writer, sheet_name='Analysis Results', index=False)
                            
                            b64 = base64.b64encode(output.getvalue()).decode()
                            href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="analysis_data.xlsx">Download Excel</a>'
                            st.markdown(href, unsafe_allow_html=True)
                        
                        st.success(f"Data exported successfully as {export_format}! Click the link above to download.")
                    
                    except Exception as e:
                        st.error(f"Error exporting data: {str(e)}")
            
                # Save visualization option if available
                if 'visualization' in st.session_state and 'figure' in st.session_state.visualization:
                    st.subheader("Export Visualization")
                
                    viz_format = st.selectbox(
                        "Select image format:",
                        ["PNG", "JPEG", "SVG", "PDF"]
                    )
                    
                    if st.button("Export Visualization"):
                        try:
                            fig = st.session_state.visualization['figure']
                            
                            # Get the image bytes in the selected format
                            if viz_format == "PNG":
                                img_bytes = fig.to_image(format="png")
                                mime_type = "image/png"
                                file_ext = "png"
                            elif viz_format == "JPEG":
                                img_bytes = fig.to_image(format="jpeg")
                                mime_type = "image/jpeg"
                                file_ext = "jpg"
                            elif viz_format == "SVG":
                                img_bytes = fig.to_image(format="svg")
                                mime_type = "image/svg+xml"
                                file_ext = "svg"
                            elif viz_format == "PDF":
                                img_bytes = fig.to_image(format="pdf")
                                mime_type = "application/pdf"
                                file_ext = "pdf"
                            
                            # Create a download link
                            b64 = base64.b64encode(img_bytes).decode()
                            href = f'<a href="data:{mime_type};base64,{b64}" download="visualization.{file_ext}">Download {viz_format}</a>'
                            st.markdown(href, unsafe_allow_html=True)
                            
                            st.success(f"Visualization exported as {viz_format}! Click the link above to download.")
                        
                        except Exception as e:
                            st.error(f"Error exporting visualization: {str(e)}")
                

                
                # Start over option
                st.divider()
                st.subheader("Start a New Analysis")
                if st.button("Start Over"):
                    # Reset session state
                    st.session_state.data = None
                    st.session_state.filename = None
                    st.session_state.current_step = 1
                    st.session_state.selected_columns = []
                    st.session_state.analysis_type = None
                    st.session_state.visualization_type = None
                    st.session_state.ai_generated_summary = None
                    if 'analysis_results' in st.session_state:
                        del st.session_state.analysis_results
                    if 'visualization' in st.session_state:
                        del st.session_state.visualization
                    if 'ttest_settings' in st.session_state:
                        del st.session_state.ttest_settings
                    if 'anova_settings' in st.session_state:
                        del st.session_state.anova_settings
                    if 'corr_settings' in st.session_state:
                        del st.session_state.corr_settings
                    
                    st.success("Session reset. Ready for a new analysis!")
                    st.rerun()
            
            # Tab 2: Data Exploration
            with tab2:
                st.header("Data Exploration")
                
                # Display full dataframe with pagination
                st.subheader("Full Dataset")
                df = st.session_state.data
                st.dataframe(df)
                
                # Display column information
                st.subheader("Column Information")
                col1, col2 = st.columns(2)
            
            with col1:
                st.write("Numerical Columns:")
                numerical_cols = df.select_dtypes(include=np.number).columns.tolist()
                if numerical_cols:
                    st.write(", ".join(numerical_cols))
                else:
                    st.write("No numerical columns found")
            
            with col2:
                st.write("Categorical Columns:")
                categorical_cols = df.select_dtypes(exclude=np.number).columns.tolist()
                if categorical_cols:
                    st.write(", ".join(categorical_cols))
                else:
                    st.write("No categorical columns found")
            
            # Check for missing values
            st.subheader("Missing Values")
            missing_values = df.isnull().sum()
            if missing_values.sum() > 0:
                st.write("Columns with missing values:")
                missing_df = pd.DataFrame({
                    'Column': missing_values.index,
                    'Missing Values': missing_values.values,
                    'Percentage': (missing_values.values / len(df) * 100).round(2)
                })
                missing_df = missing_df[missing_df['Missing Values'] > 0].sort_values('Missing Values', ascending=False)
                st.table(missing_df)
            else:
                st.write("No missing values found in the dataset.")
            
            # Data types
            st.subheader("Data Types")
            dtypes_df = pd.DataFrame({
                'Column': df.dtypes.index,
                'Data Type': df.dtypes.values
            })
            st.table(dtypes_df)
            
            # Age Grouping (now handled in drag_drop_ui.py after clicking Complete Configuration)
            st.subheader("Age Grouping")
            
            # Display information about age group generation
            if 'age_column' in st.session_state and st.session_state.age_column:
                st.info("Age groups will be generated based on your selected age column after clicking 'Complete Configuration' in the Drag & Drop interface.")
            else:
                st.info("To generate age groups, please assign an age column in the Drag & Drop interface.")
            
            # Check for our generated age group column from session state first
            age_group_col = None
            if 'generated_age_group_col' in st.session_state and st.session_state.generated_age_group_col in df.columns:
                # Use our explicitly generated age group column from drag & drop interface
                age_group_col = st.session_state.generated_age_group_col
                st.success(f"Using the age groups generated from your configuration: '{age_group_col}'")
            else:
                # Fall back to checking for existing columns in the dataframe
                existing_age_group_cols = [col for col in df.columns if 'age group' in col.lower() or 'Age group' in col]
                generated_age_group_cols = [col for col in df.columns if 'Generated Age Group' in col]
                
                # Combine both types of age group columns
                all_age_group_cols = existing_age_group_cols + generated_age_group_cols
                
                if all_age_group_cols:
                    # Still prioritize any column named 'Generated Age Group' if it exists
                    if 'Generated Age Group' in all_age_group_cols:
                        age_group_col = 'Generated Age Group'
                    else:
                        age_group_col = all_age_group_cols[0]
                        st.info(f"Using existing age group column: '{age_group_col}' from the Excel file. To create custom age groups, assign an age column in the configuration.")
            
            # Display age group distribution if we have an age column and an age group column
            if age_group_col is not None and 'age_column' in st.session_state and st.session_state.age_column:
                age_column = st.session_state.age_column['name']
                age_type = st.session_state.age_type if 'age_type' in st.session_state else "years"
                
                st.info(f"Age groups have been created based on {age_column} (measured in {age_type}).")
                
                # Display age group distribution
                age_group_counts = df[age_group_col].value_counts().sort_index()
                st.write("**Age Group Distribution:**")
                st.bar_chart(age_group_counts)
                
                # Show dataframe with the age group column next to age column
                st.write("**Dataset with Age Groups:**")
                st.dataframe(df[[age_column, age_group_col]].head())
                
                # Option to download with age groups
                csv = df.to_csv(index=False)
                st.download_button(
                    "Download Data with Age Groups",
                    csv,
                    "data_with_age_groups.csv",
                    "text/csv",
                    key='download-csv-age-groups'
                )
            elif age_group_col is None:
                st.info("Age groups will be created when you click 'Complete Configuration' in the Drag & Drop interface.")
            
        # Tab 2: Descriptive Statistics
        with tab2:
            st.header("Descriptive Statistics")
            
            # Calculate statistics for ALL columns by default
            all_cols = df.columns.tolist()
            
            # Still allow selection for display, but default to all columns
            selected_cols = st.multiselect(
                "Select columns for analysis (defaults to all columns):", 
                all_cols,
                default=all_cols
            )
            
            if selected_cols:
                try:
                    # Generate and display descriptive statistics
                    stats_df = get_descriptive_stats(df, selected_cols)
                    
                    # Convert 'Mean ± Std Dev' to a string column to avoid Arrow conversion errors
                    for col in stats_df.columns:
                        if isinstance(stats_df.loc['Mean ± Std Dev', col], str):
                            # If Mean ± Std Dev is already a string, we need to convert the whole column to strings
                            # for proper display in the dataframe
                            stats_df[col] = stats_df[col].astype(str)
                    
                    st.subheader("Summary Statistics")
                    st.dataframe(stats_df)
                    
                    # Display additional statistics for all columns
                    # Calculate statistics for all columns
                    st.subheader("Additional Statistics for All Columns")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("Skewness:")
                        # Create DataFrame with all columns
                        skew_data = []
                        for col in selected_cols:
                            try:
                                # Try to calculate skewness - will work for numerical and some other types
                                skew_value = df[col].skew()
                                skew_data.append({'Column': col, 'Skewness': skew_value})
                            except:
                                # For columns where skewness can't be calculated
                                skew_data.append({'Column': col, 'Skewness': "NA"})
                        
                        skew_df = pd.DataFrame(skew_data)
                        # Convert to string to avoid Arrow serialization issues
                        skew_df['Skewness'] = skew_df['Skewness'].astype(str)
                        st.table(skew_df)
                    
                    with col2:
                        st.write("Kurtosis:")
                        # Only calculate kurtosis for numerical columns
                        numerical_selected = [col for col in selected_cols if pd.api.types.is_numeric_dtype(df[col])]
                        if numerical_selected:
                            kurt_data = []
                            for col in numerical_selected:
                                try:
                                    kurt_value = df[col].kurtosis()
                                    kurt_data.append({'Column': col, 'Kurtosis': kurt_value})
                                except:
                                    kurt_data.append({'Column': col, 'Kurtosis': "NA"})
                            
                            kurt_df = pd.DataFrame(kurt_data)
                            kurt_df['Kurtosis'] = kurt_df['Kurtosis'].astype(str)
                            st.table(kurt_df)
                        else:
                            st.info("No numerical columns selected for kurtosis calculation.")
                    
                    # Display frequency tables for categorical columns
                    categorical_selected = [col for col in selected_cols if not pd.api.types.is_numeric_dtype(df[col])]
                    if categorical_selected:
                        st.subheader("Frequency Tables for Categorical Columns")
                        for col in categorical_selected:
                            st.write(f"**{col}:**")
                            freq_table = df[col].value_counts().reset_index()
                            freq_table.columns = [col, 'Count']
                            freq_table['Percentage'] = (freq_table['Count'] / len(df) * 100).round(2)
                            st.table(freq_table)
                
                except Exception as e:
                    st.error(f"Error calculating statistics: {str(e)}")
            else:
                st.info("Please select at least one column for analysis.")
        
        # Tab 3: Visualizations
        with tab3:
            st.header("Data Visualization")
            
            # Visualization type selector
            # Check if we have multi-sheet data
            is_multi_sheet = 'sheets_data' in st.session_state and isinstance(st.session_state.get('sheets_data'), dict) and len(st.session_state.get('sheets_data', {})) > 1
            
            if is_multi_sheet:
                viz_type = st.selectbox(
                    "Select visualization type:",
                    ["Histogram", "Box Plot", "Scatter Plot", "Correlation Heatmap", "Distribution Plot", "Multi-Sheet Relationships"]
                )
            else:
                viz_type = st.selectbox(
                    "Select visualization type:",
                    ["Histogram", "Box Plot", "Scatter Plot", "Correlation Heatmap", "Distribution Plot"]
                )
            
            if viz_type == "Histogram":
                st.subheader("Histogram")
                # Select column for histogram
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                if num_cols:
                    hist_col = st.selectbox("Select column for histogram:", num_cols)
                    bins = st.slider("Number of bins:", min_value=5, max_value=100, value=20)
                    
                    # Create histogram
                    fig = create_histogram(df, hist_col, bins)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("No numerical columns available for histogram.")
            
            elif viz_type == "Box Plot":
                st.subheader("Box Plot")
                # Select column for box plot
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                if num_cols:
                    box_col = st.selectbox("Select numerical column for box plot:", num_cols)
                    
                    # Optional grouping
                    cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
                    group_by = None
                    if cat_cols:
                        use_grouping = st.checkbox("Group by categorical variable")
                        if use_grouping:
                            group_by = st.selectbox("Select grouping variable:", cat_cols)
                    
                    # Create box plot
                    fig = create_boxplot(df, box_col, group_by)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("No numerical columns available for box plot.")
            
            elif viz_type == "Scatter Plot":
                st.subheader("Scatter Plot")
                # Select columns for scatter plot
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                if len(num_cols) >= 2:
                    x_col = st.selectbox("Select X-axis column:", num_cols, index=0)
                    y_col = st.selectbox("Select Y-axis column:", num_cols, index=min(1, len(num_cols)-1))
                    
                    # Optional color grouping
                    cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
                    color_by = None
                    if cat_cols:
                        use_color = st.checkbox("Color by categorical variable")
                        if use_color:
                            color_by = st.selectbox("Select color variable:", cat_cols)
                    
                    # Create scatter plot
                    fig = create_scatterplot(df, x_col, y_col, color_by)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("At least two numerical columns are required for a scatter plot.")
            
            elif viz_type == "Correlation Heatmap":
                st.subheader("Correlation Heatmap")
                # Select columns for correlation
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                if len(num_cols) >= 2:
                    corr_cols = st.multiselect(
                        "Select columns for correlation analysis:",
                        num_cols,
                        default=num_cols[:min(5, len(num_cols))]
                    )
                    
                    if len(corr_cols) >= 2:
                        corr_method = st.selectbox(
                            "Correlation method:",
                            ["pearson", "spearman", "kendall"]
                        )
                        
                        # Create correlation heatmap
                        fig = create_correlation_heatmap(df, corr_cols, corr_method)
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.info("Please select at least two columns for correlation analysis.")
                else:
                    st.warning("At least two numerical columns are required for a correlation heatmap.")
            
            elif viz_type == "Distribution Plot":
                st.subheader("Distribution Plot")
                # Select column for distribution plot
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                if num_cols:
                    dist_col = st.selectbox("Select column for distribution:", num_cols)
                    
                    # Create distribution plot
                    fig = create_distribution_plot(df, dist_col)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("No numerical columns available for distribution plot.")
            
            elif viz_type == "Multi-Sheet Relationships":
                st.subheader("Excel Workbook Sheet Relationships")
                
                if 'sheets_data' not in st.session_state or not isinstance(st.session_state.get('sheets_data'), dict):
                    st.warning("No multi-sheet data available. Please upload an Excel file with multiple sheets and select 'Analyze all sheets' in the Analysis tab.")
                else:
                    sheets_data = st.session_state.sheets_data
                    
                    # Import the needed functions
                    from ml_utils import analyze_sheet_relationships, identify_master_sheet
                    import plotly.graph_objects as go
                    
                    # Analyze sheet relationships
                    relationships = analyze_sheet_relationships(sheets_data)
                    
                    # Show sheet info
                    st.write(f"**Total Sheets:** {len(sheets_data)}")
                    
                    # Identify master sheet
                    master_sheet_name, master_info = identify_master_sheet(sheets_data)
                    if master_sheet_name:
                        st.success(f"**Master Sheet:** {master_sheet_name}")
                    
                    # Create visualization of sheet relationships
                    st.write("### Sheet Relationship Network")
                    st.write("This visualization shows how sheets are related to each other based on common columns.")
                    
                    # Create a network visualization using plotly
                    try:
                        sheet_names = list(sheets_data.keys())
                        
                        # Create edges based on sheet similarities
                        edges = []
                        edge_weights = []
                        sheet_similarities = relationships.get("sheet_similarities", {})
                        
                        for sheet1, similarities in sheet_similarities.items():
                            for sheet2, similarity in similarities.items():
                                # Only add edges with significant similarity
                                if similarity > 0.1:  # Threshold for showing a connection
                                    edges.append((sheet1, sheet2))
                                    edge_weights.append(similarity)
                        
                        # Create node data
                        node_x = []
                        node_y = []
                        node_text = []
                        node_size = []
                        node_color = []
                        
                        # Simple circular layout
                        import math
                        radius = 1
                        for i, sheet in enumerate(sheet_names):
                            angle = (2 * math.pi * i) / len(sheet_names)
                            x = radius * math.cos(angle)
                            y = radius * math.sin(angle)
                            
                            node_x.append(x)
                            node_y.append(y)
                            node_text.append(f"Sheet: {sheet}<br>Rows: {sheets_data[sheet].shape[0]}<br>Columns: {sheets_data[sheet].shape[1]}")
                            
                            # Master sheet gets larger marker
                            if sheet == master_sheet_name:
                                node_size.append(20)
                                node_color.append("red")
                            else:
                                node_size.append(15)
                                node_color.append("blue")
                        
                        # Create edge traces
                        edge_traces = []
                        
                        for i, (sheet1, sheet2) in enumerate(edges):
                            idx1 = sheet_names.index(sheet1)
                            idx2 = sheet_names.index(sheet2)
                            
                            x0, y0 = node_x[idx1], node_y[idx1]
                            x1, y1 = node_x[idx2], node_y[idx2]
                            
                            weight = edge_weights[i]
                            width = weight * 5  # Scale the line width based on similarity
                            
                            edge_trace = go.Scatter(
                                x=[x0, x1, None],
                                y=[y0, y1, None],
                                line=dict(width=width, color='#888'),
                                hoverinfo='none',
                                mode='lines'
                            )
                            
                            edge_traces.append(edge_trace)
                        
                        # Create node trace
                        node_trace = go.Scatter(
                            x=node_x,
                            y=node_y,
                            text=node_text,
                            mode='markers',
                            hoverinfo='text',
                            marker=dict(
                                color=node_color,
                                size=node_size,
                                line=dict(width=2, color='#000')
                            )
                        )
                        
                        # Create the figure
                        fig = go.Figure(
                            data=edge_traces + [node_trace],
                            layout=go.Layout(
                                title='Sheet Relationship Network',
                                showlegend=False,
                                hovermode='closest',
                                margin=dict(b=20, l=5, r=5, t=40),
                                xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                                yaxis=dict(showgrid=False, zeroline=False, showticklabels=False)
                            )
                        )
                        
                        st.plotly_chart(fig, use_container_width=True)
                    
                    except Exception as e:
                        st.error(f"Error creating visualization: {str(e)}")
                        import traceback
                        st.code(traceback.format_exc())
                    
                    # Display common columns between sheets
                    st.write("### Common Columns Between Sheets")
                    common_columns = relationships.get("common_columns", {})
                    
                    if common_columns:
                        common_data = []
                        for key, columns in common_columns.items():
                            sheet1, sheet2 = key.split("__")
                            common_data.append({
                                "Sheet 1": sheet1,
                                "Sheet 2": sheet2,
                                "Common Columns": ", ".join(columns[:5]) + ("..." if len(columns) > 5 else ""),
                                "Count": len(columns)
                            })
                        
                        common_df = pd.DataFrame(common_data)
                        st.dataframe(common_df)
                    else:
                        st.info("No common columns found between sheets.")
                    
                    # Display potential foreign key relationships
                    st.write("### Potential Foreign Key Relationships")
                    foreign_keys = relationships.get("potential_foreign_keys", [])
                    
                    if foreign_keys:
                        foreign_df = pd.DataFrame(foreign_keys)
                        st.dataframe(foreign_df)
                    else:
                        st.info("No potential foreign key relationships detected.")
        
        # Tab 4: Statistical Tests
        with tab4:
            st.header("Statistical Tests")
            
            test_type = st.selectbox(
                "Select statistical test:",
                ["t-Test", "ANOVA", "Correlation Analysis"]
            )
            
            if test_type == "t-Test":
                st.subheader("t-Test")
                
                # Get all columns for selection
                all_cols = df.columns.tolist()
                
                # Also keep track of numerical columns for validation
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                
                if len(all_cols) > 0:
                    # Test type
                    ttest_type = st.radio(
                        "Select t-test type:",
                        ["One-sample t-test", "Two-sample t-test"]
                    )
                    
                    if ttest_type == "One-sample t-test":
                        col = st.selectbox("Select column for test:", all_cols)
                        mu = st.number_input("Population mean (μ₀):", value=0.0)
                        alpha = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                        
                        if st.button("Perform One-sample t-test"):
                            try:
                                result = perform_ttest(df, col, test_type="one-sample", mu=mu, alpha=alpha)
                                
                                # Display results
                                st.write("### Results")
                                st.write(f"Sample mean: {result['sample_mean']:.4f}")
                                st.write(f"t-statistic: {result['t_stat']:.4f}")
                                st.write(f"p-value: {result['p_value']:.4f}")
                                st.write(f"Degrees of freedom: {result['df']}")
                                
                                # Interpretation
                                st.write("### Interpretation")
                                if result['p_value'] < alpha:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {alpha}, we reject the null hypothesis.")
                                    st.write(f"This suggests that the true mean is significantly different from {mu}.")
                                else:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {alpha}, we fail to reject the null hypothesis.")
                                    st.write(f"This suggests that there is not enough evidence to conclude that the true mean is different from {mu}.")
                                
                                # Confidence interval
                                st.write(f"95% Confidence Interval: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                            
                            except Exception as e:
                                st.error(f"Error performing t-test: {str(e)}")
                    
                    elif ttest_type == "Two-sample t-test":
                        # Two options: compare two numerical columns OR compare one numerical column grouped by categories
                        comparison_type = st.radio(
                            "Select comparison method:",
                            ["Compare two columns", "Compare groups within a column"]
                        )
                        
                        if comparison_type == "Compare two columns":
                            col1 = st.selectbox("Select first column:", all_cols, index=0)
                            col2 = st.selectbox("Select second column:", all_cols, index=min(1, len(all_cols)-1))
                            equal_var = st.checkbox("Assume equal variance", value=False)
                            alpha = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                            
                            if st.button("Perform Two-sample t-test (columns)"):
                                try:
                                    result = perform_ttest(df, col1, col2=col2, test_type="two-sample-cols", 
                                                        equal_var=equal_var, alpha=alpha)
                                    
                                    # Display results
                                    st.write("### Results")
                                    st.write(f"Mean of {col1}: {result['mean1']:.4f}")
                                    st.write(f"Mean of {col2}: {result['mean2']:.4f}")
                                    st.write(f"Mean difference: {result['mean_diff']:.4f}")
                                    st.write(f"t-statistic: {result['t_stat']:.4f}")
                                    st.write(f"p-value: {result['p_value']:.4f}")
                                    st.write(f"Degrees of freedom: {result['df']:.1f}")
                                    
                                    # Interpretation
                                    st.write("### Interpretation")
                                    if result['p_value'] < alpha:
                                        st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {alpha}, we reject the null hypothesis.")
                                        st.write(f"This suggests that there is a significant difference between the means of {col1} and {col2}.")
                                    else:
                                        st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {alpha}, we fail to reject the null hypothesis.")
                                        st.write(f"This suggests that there is not enough evidence to conclude that the means of {col1} and {col2} are significantly different.")
                                    
                                    # Confidence interval
                                    st.write(f"95% Confidence Interval of the difference: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                                
                                except Exception as e:
                                    st.error(f"Error performing t-test: {str(e)}")
                        
                        else:  # Compare groups within a column
                            cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
                            if not cat_cols:
                                st.warning("No categorical columns available for grouping.")
                            else:
                                num_col = st.selectbox("Select column to test:", all_cols)
                                group_col = st.selectbox("Select categorical column for grouping:", cat_cols)
                                
                                # Get unique values in the categorical column
                                unique_vals = df[group_col].unique()
                                if len(unique_vals) < 2:
                                    st.warning(f"Column '{group_col}' must have at least 2 unique values for comparison.")
                                else:
                                    if len(unique_vals) > 2:
                                        st.info(f"Column '{group_col}' has {len(unique_vals)} unique values. t-test requires exactly 2 groups.")
                                    
                                    group1 = st.selectbox("Select first group:", unique_vals, index=0)
                                    group2 = st.selectbox("Select second group:", unique_vals, index=min(1, len(unique_vals)-1))
                                    
                                    equal_var = st.checkbox("Assume equal variance", value=False)
                                    alpha = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                                    
                                    if st.button("Perform Two-sample t-test (groups)"):
                                        try:
                                            result = perform_ttest(df, num_col, test_type="two-sample-groups",
                                                                group_col=group_col, group1=group1, group2=group2,
                                                                equal_var=equal_var, alpha=alpha)
                                            
                                            # Display results
                                            st.write("### Results")
                                            st.write(f"Mean of {num_col} for {group_col}={group1}: {result['mean1']:.4f}")
                                            st.write(f"Mean of {num_col} for {group_col}={group2}: {result['mean2']:.4f}")
                                            st.write(f"Mean difference: {result['mean_diff']:.4f}")
                                            st.write(f"t-statistic: {result['t_stat']:.4f}")
                                            st.write(f"p-value: {result['p_value']:.4f}")
                                            st.write(f"Degrees of freedom: {result['df']:.1f}")
                                            
                                            # Interpretation
                                            st.write("### Interpretation")
                                            if result['p_value'] < alpha:
                                                st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {alpha}, we reject the null hypothesis.")
                                                st.write(f"This suggests that there is a significant difference in {num_col} between {group_col}={group1} and {group_col}={group2}.")
                                            else:
                                                st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {alpha}, we fail to reject the null hypothesis.")
                                                st.write(f"This suggests that there is not enough evidence to conclude that {num_col} differs significantly between {group_col}={group1} and {group_col}={group2}.")
                                            
                                            # Confidence interval
                                            st.write(f"95% Confidence Interval of the difference: ({result['ci_lower']:.4f}, {result['ci_upper']:.4f})")
                                        
                                        except Exception as e:
                                            st.error(f"Error performing t-test: {str(e)}")
                else:
                    st.warning("No numerical columns available for t-test.")
            
            elif test_type == "ANOVA":
                st.subheader("Analysis of Variance (ANOVA)")
                
                # Get all columns for selection
                all_cols = df.columns.tolist()
                
                # Need at least one numerical column (dependent variable) and one categorical column (factor)
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
                
                if not num_cols:
                    st.warning("No numerical columns available for ANOVA.")
                elif not cat_cols:
                    st.warning("No categorical columns available for ANOVA.")
                else:
                    # Allow any column type for selection, not just numerical
                    num_col = st.selectbox("Select column (dependent variable):", all_cols)
                    cat_col = st.selectbox("Select categorical column (factor):", cat_cols)
                    
                    # Check the number of groups in the categorical column
                    unique_groups = df[cat_col].nunique()
                    if unique_groups < 2:
                        st.warning(f"Column '{cat_col}' must have at least 2 groups for ANOVA.")
                    else:
                        alpha = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                        
                        if st.button("Perform ANOVA"):
                            try:
                                result = perform_anova(df, num_col, cat_col, alpha)
                                
                                # Display results
                                st.write("### Results")
                                st.write(f"F-statistic: {result['f_stat']:.4f}")
                                st.write(f"p-value: {result['p_value']:.4f}")
                                
                                # ANOVA table
                                st.write("### ANOVA Table")
                                st.dataframe(result['anova_table'])
                                
                                # Group statistics
                                st.write("### Group Statistics")
                                st.dataframe(result['group_stats'])
                                
                                # Interpretation
                                st.write("### Interpretation")
                                if result['p_value'] < alpha:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is less than the significance level of {alpha}, we reject the null hypothesis.")
                                    st.write(f"This suggests that at least one group in {cat_col} has a significantly different mean {num_col} from the others.")
                                else:
                                    st.write(f"With a p-value of {result['p_value']:.4f}, which is greater than the significance level of {alpha}, we fail to reject the null hypothesis.")
                                    st.write(f"This suggests that there is not enough evidence to conclude that any group in {cat_col} has a significantly different mean {num_col} from the others.")
                                
                                # Visualization of group means
                                st.write("### Visualization of Group Means")
                                fig = create_boxplot(df, num_col, cat_col)
                                st.plotly_chart(fig, use_container_width=True)
                            
                            except Exception as e:
                                st.error(f"Error performing ANOVA: {str(e)}")
            
            elif test_type == "Correlation Analysis":
                st.subheader("Correlation Analysis")
                
                # Get all columns for selection
                all_cols = df.columns.tolist()
                
                # Also keep track of numerical columns for default selection
                num_cols = df.select_dtypes(include=np.number).columns.tolist()
                
                if len(num_cols) < 2:
                    st.warning("At least two numerical columns are recommended for correlation analysis.")
                
                # Allow any column for selection, not just numerical
                corr_cols = st.multiselect(
                    "Select columns for correlation analysis:",
                    all_cols,
                    default=num_cols[:min(3, len(num_cols))] if num_cols else all_cols[:min(3, len(all_cols))]
                )
                
                if len(corr_cols) < 2:
                    st.info("Please select at least two columns for correlation analysis.")
                else:
                    corr_method = st.selectbox(
                        "Correlation method:",
                        ["pearson", "spearman", "kendall"]
                    )
                    alpha = st.slider("Significance level (α):", 0.01, 0.10, 0.05)
                    
                    if st.button("Perform Correlation Analysis"):
                        try:
                            result = perform_correlation_analysis(df, corr_cols, method=corr_method, alpha=alpha)
                            
                            # Display correlation matrix
                            st.write("### Correlation Matrix")
                            st.dataframe(result['corr_matrix'])
                            
                            # Display p-values matrix
                            st.write("### P-values Matrix")
                            st.dataframe(result['p_matrix'])
                            
                            # Display correlation heatmap
                            st.write("### Correlation Heatmap")
                            fig = create_correlation_heatmap(df, corr_cols, corr_method)
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Display significant correlations
                            st.write("### Significant Correlations")
                            if not result['significant_corrs'].empty:
                                st.dataframe(result['significant_corrs'])
                            else:
                                st.write("No significant correlations found at α = {}.".format(alpha))
                            
                            # Scatter plot matrix
                            st.write("### Scatter Plot Matrix")
                            st.write("Generating scatter plot matrix for selected variables...")
                            for i in range(len(corr_cols)):
                                for j in range(i+1, len(corr_cols)):
                                    x_col = corr_cols[i]
                                    y_col = corr_cols[j]
                                    st.write(f"**{x_col}** vs **{y_col}**")
                                    st.write(f"Correlation ({corr_method}): {result['corr_matrix'].loc[x_col, y_col]:.4f}, p-value: {result['p_matrix'].loc[x_col, y_col]:.4f}")
                                    fig = create_scatterplot(df, x_col, y_col)
                                    st.plotly_chart(fig, use_container_width=True)
                        
                        except Exception as e:
                            st.error(f"Error performing correlation analysis: {str(e)}")
        
        # Tab 5: Export Results
        with tab5:
            st.header("Export Results")
            
            export_type = st.selectbox(
                "Select export format:",
                ["CSV", "Excel", "HTML Report", "JSON"]
            )
            
            if export_type == "CSV":
                st.subheader("Export to CSV")
                
                export_options = st.multiselect(
                    "Select data to export:",
                    ["Original Data", "Descriptive Statistics"]
                )
                
                if "Original Data" in export_options:
                    # Generate CSV download link for original data
                    csv = df.to_csv(index=False)
                    b64 = base64.b64encode(csv.encode()).decode()
                    href = f'<a href="data:file/csv;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_data.csv">Download Original Data as CSV</a>'
                    st.markdown(href, unsafe_allow_html=True)
                
                if "Descriptive Statistics" in export_options:
                    # Generate statistics for all columns, not just numerical ones
                    all_cols = df.columns.tolist()
                    if all_cols:
                        stats_df = get_descriptive_stats(df, all_cols)
                        
                        # Convert any problematic values to strings for proper CSV export
                        for col in stats_df.columns:
                            if 'Mean ± Std Dev' in stats_df.index and isinstance(stats_df.loc['Mean ± Std Dev', col], str):
                                stats_df[col] = stats_df[col].astype(str)
                                
                        csv = stats_df.to_csv(index=True)
                        b64 = base64.b64encode(csv.encode()).decode()
                        href = f'<a href="data:file/csv;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_stats.csv">Download Descriptive Statistics as CSV</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    else:
                        st.warning("No columns available for statistics.")
            
            elif export_type == "Excel":
                st.subheader("Export to Excel")
                
                export_options = st.multiselect(
                    "Select data to export:",
                    ["Original Data", "Descriptive Statistics"]
                )
                
                if export_options:
                    # Create a BytesIO object to hold the Excel file
                    output = io.BytesIO()
                    
                    # Create Excel writer
                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                        if "Original Data" in export_options:
                            df.to_excel(writer, sheet_name='Data', index=False)
                        
                        if "Descriptive Statistics" in export_options:
                            # Get all columns, not just numerical ones
                            all_cols = df.columns.tolist()
                            if all_cols:
                                stats_df = get_descriptive_stats(df, all_cols)
                                
                                # Convert 'Mean ± Std Dev' to a string column to avoid Arrow conversion errors
                                for col in stats_df.columns:
                                    if isinstance(stats_df.loc['Mean ± Std Dev', col], str):
                                        # If Mean ± Std Dev is already a string, we need to convert the whole column to strings
                                        stats_df[col] = stats_df[col].astype(str)
                                
                                stats_df.to_excel(writer, sheet_name='Statistics')
                    
                    # Create download link
                    output.seek(0)
                    b64 = base64.b64encode(output.read()).decode()
                    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_analysis.xlsx">Download Excel file</a>'
                    st.markdown(href, unsafe_allow_html=True)
            
            elif export_type == "HTML Report":
                st.subheader("Generate HTML Report")
                
                include_sections = st.multiselect(
                    "Select sections to include in the report:",
                    ["Data Summary", "Descriptive Statistics", "Visualizations", "Correlation Analysis"]
                )
                
                if include_sections:
                    # Generate HTML report
                    if st.button("Generate HTML Report"):
                        report_html = generate_report(
                            df, 
                            filename=st.session_state.filename,
                            include_summary="Data Summary" in include_sections,
                            include_stats="Descriptive Statistics" in include_sections,
                            include_viz="Visualizations" in include_sections,
                            include_corr="Correlation Analysis" in include_sections
                        )
                        
                        # Create download link
                        b64 = base64.b64encode(report_html.encode()).decode()
                        href = f'<a href="data:text/html;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_report.html">Download HTML Report</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        
                        # Preview
                        with st.expander("Preview Report"):
                            st.components.v1.html(report_html, height=600, scrolling=True)
            
            elif export_type == "JSON":
                st.subheader("Export to JSON")
                
                export_options = st.multiselect(
                    "Select data to export:",
                    ["Original Data", "Descriptive Statistics"]
                )
                
                if "Original Data" in export_options:
                    # Generate JSON download link for original data
                    json_str = df.to_json(orient="records", date_format="iso")
                    b64 = base64.b64encode(json_str.encode()).decode()
                    href = f'<a href="data:application/json;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_data.json">Download Original Data as JSON</a>'
                    st.markdown(href, unsafe_allow_html=True)
                
                if "Descriptive Statistics" in export_options:
                    # Generate statistics for all columns, not just numerical ones
                    all_cols = df.columns.tolist()
                    if all_cols:
                        stats_df = get_descriptive_stats(df, all_cols)
                        
                        # Convert 'Mean ± Std Dev' to a string column to avoid Arrow conversion errors
                        for col in stats_df.columns:
                            if isinstance(stats_df.loc['Mean ± Std Dev', col], str):
                                # If Mean ± Std Dev is already a string, we need to convert the whole column to strings
                                stats_df[col] = stats_df[col].astype(str)
                        
                        # Convert to JSON with required formatting
                        json_str = stats_df.to_json(orient="split")
                        b64 = base64.b64encode(json_str.encode()).decode()
                        href = f'<a href="data:application/json;base64,{b64}" download="{st.session_state.filename.split(".")[0]}_stats.json">Download Descriptive Statistics as JSON</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    else:
                        st.warning("No columns available for statistics.")

    elif st.session_state.current_step == 8:
        # Step 8: AI-Powered Analysis Summary
        st.markdown('<div id="ai-analysis"></div>', unsafe_allow_html=True)
        st.header(emoji_header("8. AI-Powered Analysis Summary", "Research"))
        
        if st.session_state.data is not None:
            st.markdown("""
            <div class="business-section-header">
                <h2>🤖 AI-Powered Analysis Summary</h2>
                <p>Generate comprehensive narrative summaries of your statistical analyses using advanced AI models</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Display current AI configuration
            selected_model = st.session_state.get('selected_gemini_model', 'gemini-1.5-flash')
            st.info(f"Current AI Model: {selected_model}")
            
            # AI Analysis Features
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Analysis Speed", "⚡ Fast" if "flash" in selected_model else "🔬 Detailed")
            with col2:
                st.metric("Model Type", "Gemini Flash" if "flash" in selected_model else "Gemini Pro")
            with col3:
                st.metric("Output Quality", "High" if "pro" in selected_model else "Optimized")
            
            # Check if API key is configured
            status = check_gemini_api_status()
            
            if not status['ready']:
                st.warning("Configure your Gemini API key in the sidebar to enable AI-powered analysis summaries.")
                st.info("Get your free Gemini API key from Google AI Studio: https://aistudio.google.com/app/apikey")
            else:
                st.success("AI Analysis Ready - Generate narrative summaries of your statistical results")
                
                # Get available analysis results
                available_analyses = get_available_analysis_results()
                
                if not available_analyses:
                    st.info("No analysis results available yet. Complete some statistical analyses in Steps 3-4 to generate AI summaries.")
                else:
                    st.markdown("**Available Analysis Results:**")
                    
                    # Analysis selection
                    selected_analyses = st.multiselect(
                        "Select analyses to include in AI summary:",
                        available_analyses,
                        default=available_analyses,
                        help="Choose which completed analyses you want included in the AI-generated narrative summary"
                    )
                    
                    # Generate AI Summary button
                    if st.button("Generate AI Summary", type="primary", use_container_width=True):
                        if not selected_analyses:
                            st.warning("Please select at least one analysis to include in the summary.")
                        else:
                            # Clear previous summary
                            st.session_state.ai_generated_summary = None
                            
                            with st.spinner("Generating AI-powered analysis summary..."):
                                # Format data for AI
                                consolidated_text = format_data_for_ai(selected_analyses)
                                
                                if consolidated_text:
                                    # Generate AI summary
                                    summary = get_gemini_summary(consolidated_text)
                                    st.session_state.ai_generated_summary = summary
                                else:
                                    st.session_state.ai_generated_summary = "Error: No valid analysis data found for the selected analyses."
                    
                    # Display AI Summary
                    if st.session_state.ai_generated_summary:
                        if st.session_state.ai_generated_summary.startswith("Error:"):
                            st.error(st.session_state.ai_generated_summary)
                        else:
                            st.success("AI Analysis Complete!")
                            
                            # Display the summary with formatting
                            st.markdown("### AI-Generated Analysis Summary")
                            st.markdown(st.session_state.ai_generated_summary)
                            
                            # Option to copy or download the summary
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.text_area(
                                    "Copy Summary:",
                                    value=st.session_state.ai_generated_summary,
                                    height=200,
                                    help="Select all text and copy to clipboard"
                                )
                            
                            with col2:
                                # Download as text file
                                summary_bytes = st.session_state.ai_generated_summary.encode('utf-8')
                                st.download_button(
                                    "Download Summary",
                                    data=summary_bytes,
                                    file_name=f"ai_analysis_summary_{st.session_state.get('filename', 'data')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                                
                                # Download as Word document
                                try:
                                    from docx import Document
                                    doc = Document()
                                    doc.add_heading('AI-Generated Analysis Summary', 0)
                                    doc.add_paragraph(st.session_state.ai_generated_summary)
                                    
                                    doc_buffer = io.BytesIO()
                                    doc.save(doc_buffer)
                                    doc_buffer.seek(0)
                                    
                                    st.download_button(
                                        "Download as Word",
                                        data=doc_buffer.getvalue(),
                                        file_name=f"ai_analysis_summary_{st.session_state.get('filename', 'data')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True
                                    )
                                except Exception as e:
                                    st.error(f"Error creating Word document: {str(e)}")
        else:
            st.warning("Please upload data in steps 1-2 to access AI analysis features.")
    
    else:
        # Display a message when no data is loaded
        st.info("Please upload a data file or use the sample dataset to get started.")

if __name__ == "__main__":
    main()
